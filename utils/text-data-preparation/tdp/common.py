"""
Shared helpers for the text-data-preparation pipeline.

The original 21_cleaning_reports.py and 18_text_report_comparison.py each
carried their own python-docx table-building code; the single generic
comparison-docx builder here replaces both.
"""

from __future__ import annotations

import re
from dataclasses import dataclass
from pathlib import Path
from typing import List, Optional

# ── CSV helpers ────────────────────────────────────────────────────────────
REPORT_COL_CANDIDATES = [
    "report", "report_text", "radrpt", "text", "report_body", "body",
    "findings", "narrative", "clinical_text", "raw_report",
    "radiology_report", "report_content",
]


def detect_report_column(headers: List[str]) -> str:
    lower_map = {h.lower().strip(): h for h in headers}
    for candidate in REPORT_COL_CANDIDATES:
        if candidate in lower_map:
            return lower_map[candidate]
    raise ValueError(
        f"Could not auto-detect the report column. Available: {headers}. "
        "Set report_column in the pipeline config."
    )


def normalize_text(text) -> str:
    """Whitespace normalization that preserves report content."""
    if text is None:
        return ""
    text = str(text).replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


# ── Generic comparison DOCX builder ────────────────────────────────────────
@dataclass
class DocxColumn:
    key: str            # dict key in each row
    label: str          # header text
    width_inches: float
    bold: bool = False


def _set_cell_bg(cell, hex_color: str) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    cell._tc.get_or_add_tcPr().append(shd)


def _set_cell_margins(cell, top=60, bottom=60, left=100, right=100) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    tc_mar = OxmlElement("w:tcMar")
    for side, val in (("top", top), ("bottom", bottom),
                      ("left", left), ("right", right)):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"), str(val))
        el.set(qn("w:type"), "dxa")
        tc_mar.append(el)
    cell._tc.get_or_add_tcPr().append(tc_mar)


def _set_col_width(cell, width_inches: float) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    tc_w = OxmlElement("w:tcW")
    tc_w.set(qn("w:w"), str(int(width_inches * 1440)))
    tc_w.set(qn("w:type"), "dxa")
    cell._tc.get_or_add_tcPr().append(tc_w)


def _write_multiline(cell, text: str, size: int = 9, bold: bool = False) -> None:
    from docx.shared import Pt
    cell.paragraphs[0].clear()
    lines = str(text).split("\n")
    for i, line in enumerate(lines):
        para = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        para.paragraph_format.space_after = Pt(2) if i < len(lines) - 1 else Pt(0)
        run = para.add_run(line.strip())
        run.font.name = "Calibri"
        run.font.size = Pt(size)
        run.font.bold = bold


def build_comparison_docx(
    rows: List[dict],
    columns: List[DocxColumn],
    title: str,
    subtitle: str,
    out_path: Path,
    landscape: bool = False,
) -> Path:
    """
    Fixed-width zebra-striped comparison table with a repeating header row.
    One table row per input dict; cell text taken from row[col.key].
    """
    from docx import Document
    from docx.enum.table import WD_ALIGN_VERTICAL
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor

    doc = Document()
    section = doc.sections[0]
    if landscape:
        section.page_width, section.page_height = Inches(11), Inches(8.5)
        section.orientation = 1
    else:
        section.page_width, section.page_height = Inches(8.5), Inches(11)
    for attr in ("left_margin", "right_margin", "top_margin", "bottom_margin"):
        setattr(section, attr, Inches(0.5))

    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.paragraph_format.space_after = Pt(4)
    t = title_para.add_run(title)
    t.font.name = "Calibri"
    t.font.size = Pt(14)
    t.font.bold = True

    sub_para = doc.add_paragraph()
    sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_para.paragraph_format.space_after = Pt(10)
    s = sub_para.add_run(subtitle)
    s.font.name = "Calibri"
    s.font.size = Pt(9)
    s.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    table = doc.add_table(rows=1, cols=len(columns))
    table.style = "Table Grid"

    # Lock table to fixed total width
    tbl_pr = table._tbl.find(qn("w:tblPr"))
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        table._tbl.insert(0, tbl_pr)
    tbl_w = OxmlElement("w:tblW")
    tbl_w.set(qn("w:w"), str(int(sum(c.width_inches for c in columns) * 1440)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_pr.append(tbl_w)

    # Header row — repeats on every page
    hdr = table.rows[0]
    tr_pr = hdr._tr.get_or_add_trPr()
    tr_pr.append(OxmlElement("w:tblHeader"))
    for cell, col in zip(hdr.cells, columns):
        _set_cell_bg(cell, "1F3864")
        _set_cell_margins(cell)
        _set_col_width(cell, col.width_inches)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(col.label)
        run.font.name = "Calibri"
        run.font.size = Pt(10)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # Data rows — zebra striping
    for i, row_data in enumerate(rows):
        bg: Optional[str] = "EEF3F8" if i % 2 == 0 else None
        row = table.add_row()
        for cell, col in zip(row.cells, columns):
            if bg:
                _set_cell_bg(cell, bg)
            _set_cell_margins(cell)
            _set_col_width(cell, col.width_inches)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            _write_multiline(cell, row_data.get(col.key, ""), bold=col.bold)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    return out_path
