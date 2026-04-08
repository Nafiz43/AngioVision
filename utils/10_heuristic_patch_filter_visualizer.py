#!/usr/bin/env python3
"""
patch_grid_patch_size_sweep.py

For each patch size in [8, 16, 32, 64, 128, 256]:
  - Sweeps informative_percent 0 → 100 in steps of 5 (21 values)
  - Processes 10 hardcoded frames in parallel per threshold
  - Arranges 10 images as 2 rows × 5 cols with a threshold label banner on top
  - Saves a separate .docx to /data/Deep_Angiography/DICOM-metadata-stats/
      patch_size_008_threshold_sweep.docx
      patch_size_016_threshold_sweep.docx
      patch_size_032_threshold_sweep.docx
      patch_size_064_threshold_sweep.docx
      patch_size_128_threshold_sweep.docx

Run (no arguments needed):
    python3 patch_grid_patch_size_sweep.py

Dependencies:
    pip install opencv-python numpy tqdm python-docx --break-system-packages
"""

from __future__ import annotations

import io
import math
import subprocess
import sys
from pathlib import Path
from concurrent.futures import ProcessPoolExecutor, as_completed

import cv2
import numpy as np
from tqdm import tqdm


# ══════════════════════════════════════════════════════
#  ALL HARDCODED CONFIGURATION — nothing to pass in
# ══════════════════════════════════════════════════════

HARDCODED_FRAME_PATHS = [
    # Original 5
    "/data/Deep_Angiography/DICOM_Sequence_Processed/03_DSA 3 LD/2.16.840.1.113883.3.16.90729317817737798748724286231292096160/frames/2.16.840.1.113883.3.16.90729317817737798748724286231292096160_frame_0005.png",
    "/data/Deep_Angiography/DICOM_Sequence_Processed/gzeAvoZBK5/2.16.840.1.113883.3.16.27887747233506821641442574601125740969/frames/2.16.840.1.113883.3.16.27887747233506821641442574601125740969_frame_0016.png",
    "/data/Deep_Angiography/DICOM_Sequence_Processed/received/2.16.840.1.113883.3.16.206533095497994768887122049484591788029/frames/2.16.840.1.113883.3.16.206533095497994768887122049484591788029_frame_0007.png",
    "/data/Deep_Angiography/DICOM_Sequence_Processed/14_DSA 3/2.16.840.1.113883.3.16.229884264127170088524043314078208850582/frames/2.16.840.1.113883.3.16.229884264127170088524043314078208850582_frame_0008.png",
    "/data/Deep_Angiography/DICOM_Sequence_Processed/DICOM/2.16.840.1.113883.3.16.229149697742230994696452930366730877174/frames/2.16.840.1.113883.3.16.229149697742230994696452930366730877174_frame_0011.png",
    # New 5
    "/data/Deep_Angiography/DICOM_Sequence_Processed/01_DSA 6/2.16.840.1.113883.3.16.25137877796933303971610529580184712113/frames/2.16.840.1.113883.3.16.25137877796933303971610529580184712113_frame_0090.png",
    "/data/Deep_Angiography/DICOM_Sequence_Processed/01_DSA 6/2.16.840.1.113883.3.16.204846168798754550014642121905769244969/frames/2.16.840.1.113883.3.16.204846168798754550014642121905769244969_frame_0017.png",
    "/data/Deep_Angiography/DICOM_Sequence_Processed/01_DSA 6/2.16.840.1.113883.3.16.322765241157606197270633442737278789160/frames/2.16.840.1.113883.3.16.322765241157606197270633442737278789160_frame_0025.png",
    "/data/Deep_Angiography/DICOM_Sequence_Processed/1e8th1tcdp/2.16.840.1.113883.3.16.164546581401405817413395995977118384897/frames/2.16.840.1.113883.3.16.164546581401405817413395995977118384897_frame_0022.png",
    "/data/Deep_Angiography/DICOM_Sequence_Processed/1QKmyc87mw/2.16.840.1.113883.3.16.122439881467805637368480778603593518413/frames/2.16.840.1.113883.3.16.122439881467805637368480778603593518413_frame_0016.png",
]

PATCH_SIZES      = [8, 16, 32, 64, 128]
THRESHOLD_VALUES = list(range(0, 101, 5))   # 0,5,10,...,100  (21 values)
OUTPUT_BASE_DIR  = Path("/data/Deep_Angiography/DICOM-metadata-stats/patch-sweeps")

# Processing knobs
TARGET_HEIGHT           = 512
MAX_WORKERS             = 10
DARK_PIXEL_THRESHOLD    = 20
DARK_PIXEL_RATIO_REJECT = 0.85
MIN_VARIANCE            = 10.0
MIN_EDGE_DENSITY        = 0.01
MIN_ENTROPY             = 2.0
ALPHA                   = 0.18   # lower = more transparent fill

WEIGHT_VARIANCE     = 1.0
WEIGHT_ENTROPY      = 1.0
WEIGHT_EDGE_DENSITY = 1.5

GREEN         = (0, 160, 0)     # muted green  (was 0,255,0)
RED           = (0, 0, 180)     # muted red    (was 0,0,255)
TEXT_COLOR    = (255, 255, 255)
LINE_THICKNESS = 1
FONT          = cv2.FONT_HERSHEY_SIMPLEX
HEADER_HEIGHT = 52   # pixel height of the label banner baked into each composite


# ══════════════════════════════════════════════════════
#  AUTO-INSTALL python-docx IF MISSING
# ══════════════════════════════════════════════════════

def ensure_python_docx() -> None:
    try:
        import docx  # noqa: F401
    except ImportError:
        print("python-docx not found — installing...")
        subprocess.check_call([
            sys.executable, "-m", "pip", "install",
            "python-docx", "--break-system-packages", "-q"
        ])
        print("python-docx installed.")


# ══════════════════════════════════════════════════════
#  PATCH ANALYSIS HELPERS
# ══════════════════════════════════════════════════════

def compute_entropy(gray_patch: np.ndarray) -> float:
    hist = cv2.calcHist([gray_patch], [0], None, [256], [0, 256]).ravel()
    s = hist.sum()
    if s == 0:
        return 0.0
    prob = hist / s
    prob = prob[prob > 0]
    return float(-np.sum(prob * np.log2(prob)))


def compute_edge_density(gray_patch: np.ndarray) -> float:
    edges = cv2.Canny(gray_patch, 50, 150)
    return float(np.count_nonzero(edges)) / float(edges.size)


def dark_pixel_ratio(gray_patch: np.ndarray, dark_threshold: int) -> float:
    return float(np.mean(gray_patch <= dark_threshold))


def normalize_array(values: np.ndarray) -> np.ndarray:
    if len(values) == 0:
        return values
    vmin, vmax = np.min(values), np.max(values)
    if math.isclose(float(vmin), float(vmax)):
        return np.zeros_like(values, dtype=np.float32)
    return ((values - vmin) / (vmax - vmin)).astype(np.float32)


def classify_patches(image_bgr: np.ndarray,
                     patch_size: int,
                     informative_percent: float) -> list[dict]:
    gray = cv2.cvtColor(image_bgr, cv2.COLOR_BGR2GRAY)
    h, w = gray.shape
    patches: list[dict] = []

    for y in range(0, h, patch_size):
        for x in range(0, w, patch_size):
            patch = gray[y:y + patch_size, x:x + patch_size]
            if patch.shape[0] != patch_size or patch.shape[1] != patch_size:
                continue
            variance     = float(np.var(patch))
            entropy      = compute_entropy(patch)
            edge_density = compute_edge_density(patch)
            dark_ratio   = dark_pixel_ratio(patch, DARK_PIXEL_THRESHOLD)
            trivial = (
                dark_ratio      >= DARK_PIXEL_RATIO_REJECT
                or variance     <  MIN_VARIANCE
                or edge_density <  MIN_EDGE_DENSITY
                or entropy      <  MIN_ENTROPY
            )
            patches.append({
                "x": x, "y": y, "w": patch_size, "h": patch_size,
                "variance": variance, "entropy": entropy,
                "edge_density": edge_density, "dark_ratio": dark_ratio,
                "trivial": trivial, "score": None, "label": None,
            })

    valid_idx = [i for i, p in enumerate(patches) if not p["trivial"]]
    if valid_idx:
        var_arr  = np.array([patches[i]["variance"]     for i in valid_idx], dtype=np.float32)
        ent_arr  = np.array([patches[i]["entropy"]      for i in valid_idx], dtype=np.float32)
        edge_arr = np.array([patches[i]["edge_density"] for i in valid_idx], dtype=np.float32)
        scores   = (WEIGHT_VARIANCE     * normalize_array(var_arr)
                  + WEIGHT_ENTROPY      * normalize_array(ent_arr)
                  + WEIGHT_EDGE_DENSITY * normalize_array(edge_arr))
        for idx, score in zip(valid_idx, scores):
            patches[idx]["score"] = float(score)

        n_info   = max(1, int(round(len(valid_idx) * informative_percent / 100.0)))
        ranked   = sorted(valid_idx, key=lambda i: patches[i]["score"], reverse=True)
        info_set = set(ranked[:n_info])
        for i, p in enumerate(patches):
            p["label"] = "informative" if (not p["trivial"] and i in info_set) else "uninformative"
    else:
        for p in patches:
            p["label"] = "uninformative"
    return patches


def draw_patch_overlay(image_bgr: np.ndarray, patches: list[dict]) -> np.ndarray:
    base    = image_bgr.copy()
    overlay = image_bgr.copy()
    for p in patches:
        color = GREEN if p["label"] == "informative" else RED
        cv2.rectangle(overlay, (p["x"], p["y"]),
                      (p["x"]+p["w"], p["y"]+p["h"]), color, -1)
    blended = cv2.addWeighted(overlay, ALPHA, base, 1.0 - ALPHA, 0)
    for p in patches:
        color = GREEN if p["label"] == "informative" else RED
        cv2.rectangle(blended, (p["x"], p["y"]),
                      (p["x"]+p["w"], p["y"]+p["h"]), color, LINE_THICKNESS)
    return blended


def add_legend(image_bgr: np.ndarray) -> np.ndarray:
    out     = image_bgr.copy()
    overlay = out.copy()
    cv2.rectangle(overlay, (10, 10), (240, 78), (0, 0, 0), -1)
    out = cv2.addWeighted(overlay, 0.55, out, 0.45, 0)
    cv2.rectangle(out, (18, 18), (32, 32), GREEN, -1)
    cv2.putText(out, "Informative",   (38, 30), FONT, 0.42, TEXT_COLOR, 1, cv2.LINE_AA)
    cv2.rectangle(out, (18, 42), (32, 56), RED, -1)
    cv2.putText(out, "Uninformative", (38, 54), FONT, 0.42, TEXT_COLOR, 1, cv2.LINE_AA)
    return out


def resize_to_height(image_bgr: np.ndarray, target_height: int) -> np.ndarray:
    h, w = image_bgr.shape[:2]
    if h == target_height:
        return image_bgr
    scale = target_height / float(h)
    tw    = max(1, int(round(w * scale)))
    return cv2.resize(image_bgr, (tw, target_height), interpolation=cv2.INTER_AREA)


# ══════════════════════════════════════════════════════
#  WORKER  (runs in subprocess pool)
# ══════════════════════════════════════════════════════

def _process_frame(task: dict) -> dict:
    img = cv2.imread(task["frame_path"])
    if img is None:
        return {"ok": False, "idx": task["idx"], "error": task["frame_path"]}
    patches = classify_patches(img, task["patch_size"], task["threshold"])
    out     = draw_patch_overlay(img, patches)
    out     = add_legend(out)
    out     = resize_to_height(out, TARGET_HEIGHT)
    return {"ok": True, "idx": task["idx"], "image": out}


# ══════════════════════════════════════════════════════
#  BUILD COMPOSITE  (10 frames → banner + 2 rows × 5)
# ══════════════════════════════════════════════════════

def build_composite(images: list[np.ndarray],
                    threshold: int,
                    patch_size: int) -> np.ndarray:
    def pad_row(row: list[np.ndarray], n: int = 5) -> np.ndarray:
        h, w = row[0].shape[:2]
        while len(row) < n:
            row.append(np.zeros((h, w, 3), dtype=np.uint8))
        return np.hstack(row)

    row1 = pad_row(list(images[:5]))
    row2 = pad_row(list(images[5:10]))

    tw = max(row1.shape[1], row2.shape[1])

    def wpad(img: np.ndarray) -> np.ndarray:
        diff = tw - img.shape[1]
        if diff > 0:
            return np.hstack([img, np.zeros((img.shape[0], diff, 3), dtype=np.uint8)])
        return img

    grid   = np.vstack([wpad(row1), wpad(row2)])
    header = np.full((HEADER_HEIGHT, grid.shape[1], 3), 30, dtype=np.uint8)
    label  = f"informative_percent = {threshold}%   |   patch_size = {patch_size}px"
    cv2.putText(header, label, (16, 34), FONT, 0.85, (220, 220, 220), 2, cv2.LINE_AA)

    return np.vstack([header, grid])


# ══════════════════════════════════════════════════════
#  PROCESS ONE PATCH SIZE → list[(threshold, Path)]
# ══════════════════════════════════════════════════════

def run_patch_size(patch_size: int,
                   frame_paths: list[Path],
                   tmp_dir: Path) -> list[tuple[int, Path]]:
    ps_dir = tmp_dir / f"patch_{patch_size:03d}"
    ps_dir.mkdir(parents=True, exist_ok=True)

    results: list[tuple[int, Path]] = []

    for threshold in tqdm(THRESHOLD_VALUES,
                          desc=f"  patch={patch_size:3d}px  thresholds",
                          leave=False):
        tasks = [
            {"idx": i, "frame_path": str(fp),
             "patch_size": patch_size, "threshold": threshold}
            for i, fp in enumerate(frame_paths)
        ]

        workers    = max(1, min(MAX_WORKERS, len(tasks)))
        frame_imgs = [None] * len(tasks)

        with ProcessPoolExecutor(max_workers=workers) as ex:
            futures = {ex.submit(_process_frame, t): t["idx"] for t in tasks}
            for fut in as_completed(futures):
                r = fut.result()
                if not r["ok"]:
                    print(f"[WARN] Cannot read: {r['error']}")
                    continue
                frame_imgs[r["idx"]] = r["image"]

        images = [img for img in frame_imgs if img is not None]
        if not images:
            print(f"[WARN] No images for patch={patch_size} thr={threshold}, skipping")
            continue

        composite = build_composite(images, threshold, patch_size)
        out_path  = ps_dir / f"threshold_{threshold:03d}.png"
        if not cv2.imwrite(str(out_path), composite):
            raise IOError(f"Failed to write {out_path}")
        results.append((threshold, out_path))

    return results


# ══════════════════════════════════════════════════════
#  DOCX GENERATION  (pure python-docx, no npm/node)
# ══════════════════════════════════════════════════════

def generate_docx_for_patch_size(patch_size: int,
                                  threshold_image_paths: list[tuple[int, Path]],
                                  output_docx: Path) -> None:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # ── Page layout: A4 landscape, tight margins to maximise image width ──
    section               = doc.sections[0]
    section.page_width    = Cm(29.7)
    section.page_height   = Cm(21.0)
    section.left_margin   = Cm(1.0)
    section.right_margin  = Cm(1.0)
    section.top_margin    = Cm(1.2)
    section.bottom_margin = Cm(1.2)

    # Compute content width in inches directly (avoids int subtraction losing .inches)
    PAGE_WIDTH_IN   = 29.7 / 2.54   # cm → inches
    LEFT_MARGIN_IN  = 1.0  / 2.54
    RIGHT_MARGIN_IN = 1.0  / 2.54
    content_width_inches = PAGE_WIDTH_IN - LEFT_MARGIN_IN - RIGHT_MARGIN_IN

    # ── Document title ──
    title_para = doc.add_heading(
        f"Patch Size = {patch_size}px  —  Threshold Sweep (informative_percent 0→100, step 5)",
        level=1,
    )
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title_para.runs:
        run.font.size      = Pt(18)
        run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub.add_run(
        "10 frames per threshold  |  Row 1: frames 1–5   Row 2: frames 6–10"
    )
    sub_run.font.size      = Pt(9)
    sub_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_paragraph()  # spacer

    # ── One entry per threshold ──
    for threshold, img_path in threshold_image_paths:
        # Section heading
        heading = doc.add_heading(f"informative_percent = {threshold}%", level=2)
        heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in heading.runs:
            run.font.size      = Pt(13)
            run.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)

        # Load composite image
        img = cv2.imread(str(img_path))
        if img is None:
            doc.add_paragraph(f"[Image not found: {img_path}]")
            continue
        img_h, img_w = img.shape[:2]
        aspect = img_h / img_w

        # Encode to PNG bytes in memory (avoids temp files)
        ok, buf = cv2.imencode(".png", img)
        if not ok:
            doc.add_paragraph(f"[Could not encode: {img_path}]")
            continue
        img_stream = io.BytesIO(buf.tobytes())

        # Insert image at full content width, correct aspect ratio
        img_para           = doc.add_paragraph()
        img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run                = img_para.add_run()
        run.add_picture(
            img_stream,
            width=Inches(content_width_inches),
            height=Inches(content_width_inches * aspect),
        )

        doc.add_paragraph()  # spacer between entries

    doc.save(str(output_docx))
    print(f"  ✓  {output_docx.name}")


# ══════════════════════════════════════════════════════
#  MAIN
# ══════════════════════════════════════════════════════

def main() -> None:
    ensure_python_docx()

    # Validate frames
    frame_paths = [Path(p) for p in HARDCODED_FRAME_PATHS]
    missing = [p for p in frame_paths if not p.exists()]
    if missing:
        for m in missing:
            print(f"[ERROR] Frame not found: {m}")
        sys.exit(1)

    OUTPUT_BASE_DIR.mkdir(parents=True, exist_ok=True)
    tmp_dir = OUTPUT_BASE_DIR / "_tmp_composites"
    tmp_dir.mkdir(parents=True, exist_ok=True)

    print(f"\n{'='*60}")
    print(f" Patch sizes : {PATCH_SIZES}")
    print(f" Thresholds  : {THRESHOLD_VALUES}")
    print(f" Frames      : {len(frame_paths)}")
    print(f" Output dir  : {OUTPUT_BASE_DIR}")
    print(f"{'='*60}\n")

    for patch_size in PATCH_SIZES:
        print(f"\n{'─'*50}")
        print(f" Processing patch_size = {patch_size}px")
        print(f"{'─'*50}")

        threshold_image_paths = run_patch_size(patch_size, frame_paths, tmp_dir)

        docx_path = OUTPUT_BASE_DIR / f"patch_size_{patch_size:03d}_threshold_sweep.docx"
        print(f"  Generating DOCX → {docx_path.name} ...")
        generate_docx_for_patch_size(patch_size, threshold_image_paths, docx_path)

    print(f"\n{'='*60}")
    print(f" All done!  DOCX files saved to:\n  {OUTPUT_BASE_DIR}")
    print(f"{'='*60}\n")


if __name__ == "__main__":
    main()