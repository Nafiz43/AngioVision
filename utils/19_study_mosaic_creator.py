import os
import pandas as pd
from pathlib import Path
from docx import Document
from docx.shared import Inches

# -----------------------------
# Paths
# -----------------------------
BASE_DIR = Path("/data/Deep_Angiography/Validation_Data/Validation_Data_2026_03_23")

SEQ_ROOT = BASE_DIR / "DICOM_Sequence_Processed"

STUDY_CSV = BASE_DIR / "consolidated_metadata_ALL_Sequences.csv"
RADRPT_CSV = BASE_DIR / "Validation_Studies_2026_03_23.csv"
MODEL_CSV = Path("/data/Deep_Angiography/AngioVision/fine-tuning/output/report_generation/30_4_16_32/epoch_30_generated_reports.csv")

OUTPUT_DIR = Path("/data/Deep_Angiography/AngioVision/fine-tuning/output")
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

# -----------------------------
# Load CSVs
# -----------------------------
seq_df = pd.read_csv(STUDY_CSV)
rpt_df = pd.read_csv(RADRPT_CSV)
model_df = pd.read_csv(MODEL_CSV)

seq_df.columns = [c.strip() for c in seq_df.columns]
rpt_df.columns = [c.strip() for c in rpt_df.columns]
model_df.columns = [c.strip() for c in model_df.columns]

# -----------------------------
# Build lookup tables
# -----------------------------
# Study → Accession
study_to_acc = dict(zip(seq_df["StudyInstanceUID"], seq_df["Anon Acc #"]))

# Accession → radrpt
acc_to_radrpt = dict(zip(rpt_df["accnum"], rpt_df["radrpt"]))

# Accession → generated report
acc_to_model = dict(zip(model_df["accession"], model_df["generated_report"]))

# -----------------------------
# Fast index of mosaics (IMPORTANT optimization)
# -----------------------------
print("[INFO] Indexing mosaic paths...")

mosaic_index = {}
for p in SEQ_ROOT.rglob("mosaic.png"):
    # key assumption: SOPInstanceUID appears in folder path
    mosaic_index[str(p)] = p

# -----------------------------
# Better helper: match SOP → mosaic
# -----------------------------
def find_mosaic_by_sop(sop_uid: str):
    for path_str, path_obj in mosaic_index.items():
        if sop_uid in path_str:
            return path_obj
    return None

# -----------------------------
# Build SOP mapping per study
# -----------------------------
grouped = seq_df.groupby("StudyInstanceUID")

# -----------------------------
# DOCX builder
# -----------------------------
def create_doc(study_id, sop_list):
    doc = Document()

    acc = study_to_acc.get(study_id, "UNKNOWN")

    # -------------------------
    # Header
    # -------------------------
    doc.add_heading(f"Anon Acc #: {acc}", level=1)

    # -------------------------
    # Ground truth report
    # -------------------------
    radrpt = acc_to_radrpt.get(acc, "NOT FOUND")
    doc.add_heading("Radiology Report (Ground Truth)", level=2)
    doc.add_paragraph(str(radrpt))

    # -------------------------
    # Model report
    # -------------------------
    model_rpt = acc_to_model.get(acc, "NOT FOUND")
    doc.add_heading("Model Generated Report", level=2)
    doc.add_paragraph(str(model_rpt))

    # -------------------------
    # Mosaics
    # -------------------------
    doc.add_heading("Angiography Sequences", level=2)

    added = False

    for sop in sop_list:
        if pd.isna(sop):
            continue

        mosaic_path = find_mosaic_by_sop(str(sop))

        if mosaic_path is None or not mosaic_path.exists():
            continue

        doc.add_paragraph(f"Sequence: {sop}")

        try:
            doc.add_picture(str(mosaic_path), width=Inches(6.5))
            added = True
        except Exception as e:
            print(f"[WARN] Failed image {mosaic_path}: {e}")

    if added:
        out_file = OUTPUT_DIR / f"{acc}.docx"
        doc.save(out_file)
        print(f"[OK] Saved {out_file}")
    else:
        print(f"[SKIP] No mosaics for study {study_id}")

# -----------------------------
# Main loop
# -----------------------------
def main():
    for study_id, group in grouped:
        sop_values = group["SOPInstanceUIDs"].dropna().tolist()

        sop_list = []
        for v in sop_values:
            sop_list.extend([x.strip() for x in str(v).split(",")])

        create_doc(study_id, sop_list)

if __name__ == "__main__":
    main()