"""
generate_report_comparison.py

Reads Report_List_v01_01_augmented.csv and produces a Word (.docx) file
with a 3-column fixed-width table:

    Acc ID  |  Original  |  Augmented

Each (Original, Augmented) pair gets its own row.
If one original has N augmented variants → N rows, with Acc ID and Original
text repeated each time so the reviewer always sees both columns at once.
No horizontal scrolling required.

Requirements:
    pip install pandas python-docx

Usage:
    python generate_report_comparison.py
"""

import pandas as pd
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# ── CONFIG ────────────────────────────────────────────────────────────────────
INPUT_CSV   = "/data/Deep_Angiography/Reports/Report_List_v01_01_augmented.csv"
OUTPUT_DOCX = os.path.join(os.path.dirname(INPUT_CSV), "report_comparison.docx")

# Column widths in inches (page=8.5", margins=0.5" each → content=7.5")
COL_W_ID  = 1.1   # Acc ID
COL_W_ORI = 3.2   # Original
COL_W_AUG = 3.2   # Augmented
# ─────────────────────────────────────────────────────────────────────────────


# ── Helpers ───────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    """Fill a table cell with a solid background colour (hex, e.g. '1F3864')."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def set_cell_margins(cell, top=60, bottom=60, left=100, right=100):
    """Set inner cell padding (twips)."""
    tc    = cell._tc
    tcPr  = tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for side, val in (("top", top), ("bottom", bottom),
                      ("left", left), ("right", right)):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"),    str(val))
        el.set(qn("w:type"), "dxa")
        tcMar.append(el)
    tcPr.append(tcMar)


def set_col_width(cell, width_inches: float):
    """Force a cell to a fixed width."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW  = OxmlElement("w:tcW")
    tcW.set(qn("w:w"),    str(int(width_inches * 1440)))
    tcW.set(qn("w:type"), "dxa")
    tcPr.append(tcW)


def add_header_cell(cell, text: str, width_inches: float):
    set_cell_bg(cell, "1F3864")
    set_cell_margins(cell)
    set_col_width(cell, width_inches)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run  = para.add_run(text)
    run.font.name      = "Calibri"
    run.font.size      = Pt(10)
    run.font.bold      = True
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)


def add_data_cell(cell, text: str, width_inches: float,
                  bold=False, bg_hex: str = None):
    if bg_hex:
        set_cell_bg(cell, bg_hex)
    set_cell_margins(cell)
    set_col_width(cell, width_inches)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

    cell.paragraphs[0].clear()
    lines = str(text).split("\n")
    for i, line in enumerate(lines):
        para = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        para.paragraph_format.space_after = Pt(2) if i < len(lines) - 1 else Pt(0)
        run = para.add_run(line.strip())
        run.font.name = "Calibri"
        run.font.size = Pt(9)
        run.font.bold = bold


def enable_repeat_header(row):
    """Make table header row repeat on every page."""
    trPr = row._tr.get_or_add_trPr()
    tblHeader = OxmlElement("w:tblHeader")
    trPr.append(tblHeader)


def set_table_fixed_width(table):
    """Lock table to fixed column widths."""
    tbl   = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblW = OxmlElement("w:tblW")
    total = int((COL_W_ID + COL_W_ORI + COL_W_AUG) * 1440)
    tblW.set(qn("w:w"),    str(total))
    tblW.set(qn("w:type"), "dxa")
    tblPr.append(tblW)


# ── Data loading ──────────────────────────────────────────────────────────────

def build_rows(df: pd.DataFrame) -> list:
    """
    Return list of dicts with keys: acc_id, original, augmented, aug_label.
    One entry per (original, augmented) pair — original repeats for each variant.
    """
    df.columns       = df.columns.str.strip()
    df["Type"]       = df["Type"].str.strip()
    df["radrpt"]     = df["radrpt"].str.strip()
    df["Anon Acc #"] = df["Anon Acc #"].str.strip()

    originals = df[df["Type"] == "Original"]
    augmented = (df[df["Type"] != "Original"]
                 .sort_values(["Anon Acc #", "Type"]))

    rows = []
    for _, orig in originals.iterrows():
        anon_id   = orig["Anon Acc #"]
        orig_text = orig["radrpt"]
        aug_rows  = augmented[augmented["Anon Acc #"] == anon_id]

        if aug_rows.empty:
            rows.append({"acc_id": anon_id, "original": orig_text,
                         "aug_label": "—", "augmented": "—"})
        else:
            for _, aug in aug_rows.iterrows():
                rows.append({
                    "acc_id":    anon_id,
                    "original":  orig_text,
                    "aug_label": aug["Type"],
                    "augmented": aug["radrpt"],
                })
    return rows


# ── Document builder ──────────────────────────────────────────────────────────

def build_docx(rows: list, output_path: str):
    doc = Document()

    # Page setup: US Letter, 0.5" margins
    section = doc.sections[0]
    section.page_width   = Inches(8.5)
    section.page_height  = Inches(11)
    section.left_margin  = section.right_margin  = Inches(0.5)
    section.top_margin   = section.bottom_margin = Inches(0.5)

    # Title
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.paragraph_format.space_after = Pt(4)
    t = title_para.add_run(
        "Radiology Report Comparison — Original vs Augmented"
    )
    t.font.name = "Calibri"; t.font.size = Pt(14); t.font.bold = True

    sub_para = doc.add_paragraph()
    sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_para.paragraph_format.space_after = Pt(10)
    s = sub_para.add_run(
        f"{len(rows)} row(s)  ·  source: {os.path.basename(INPUT_CSV)}"
    )
    s.font.name = "Calibri"
    s.font.size = Pt(9)
    s.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # Table
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    set_table_fixed_width(table)

    # Header row
    hdr = table.rows[0]
    enable_repeat_header(hdr)
    add_header_cell(hdr.cells[0], "Acc ID",    COL_W_ID)
    add_header_cell(hdr.cells[1], "Original",  COL_W_ORI)
    add_header_cell(hdr.cells[2], "Augmented", COL_W_AUG)

    # Data rows
    for i, r in enumerate(rows):
        bg = "EEF3F8" if i % 2 == 0 else None
        row = table.add_row()
        add_data_cell(row.cells[0], r["acc_id"],    COL_W_ID,  bold=True, bg_hex=bg)
        add_data_cell(row.cells[1], r["original"],  COL_W_ORI, bg_hex=bg)
        add_data_cell(row.cells[2], r["augmented"], COL_W_AUG, bg_hex=bg)

    doc.save(output_path)
    print(f"Saved → {output_path}")


# ── Main ──────────────────────────────────────────────────────────────────────

def main():
    print(f"Reading: {INPUT_CSV}")
    df = pd.read_csv(INPUT_CSV, dtype=str)
    print(f"  → {len(df)} rows loaded")

    rows = build_rows(df)
    print(f"  → {len(rows)} Original/Augmented pairs")

    build_docx(rows, OUTPUT_DOCX)


if __name__ == "__main__":
    main()