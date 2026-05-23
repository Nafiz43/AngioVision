#!/usr/bin/env python3
"""
Angiographic Report Cleaning Pipeline
=====================================

A comprehensive, robust pipeline to clean angiographic reports for downstream
model training (e.g., visual-to-report generation).

Cleaning steps (in order):
    1. Unicode / encoding normalization
    2. PHI removal via Microsoft Presidio (+ custom medical recognizers)
    3. Attestation / signature section removal (trailing block)
    4. Header metadata removal (patient/MRN/DOB/accession/physician/location/phone
       block at the top — not needed for model training)
    5. HISTORY section removal
    6. Medical abbreviation expansion (before casing so acronyms are recognizable)
    7. Dictation shorthand fixes (w/, h/o, s/p, y/o, etc.)
    8. Punctuation, symbol, and list-marker cleanup
    9. Whitespace collapse and section-header normalization (blank line between sections)
   10. Sentence-case conversion (lowercase except block-letter words, capitalize sentence starts)

Input:
    Reads reports from a CSV file. Configure CSV_PATH, REPORT_COLUMN, and
    NUM_REPORTS_TO_CLEAN and NUM_REPORTS_IN_DOCX below.

Output:
    A .docx file with side-by-side tables: Original | Cleaned for each report.

Dependencies:
    pip install presidio-analyzer presidio-anonymizer python-docx
    python -m spacy download en_core_web_lg

Usage:
    python clean_angio_report.py

Author: built for Nafiz (UC Davis DECAL Lab)
"""

from __future__ import annotations

import csv
import os
import re
import unicodedata
from concurrent.futures import ProcessPoolExecutor
from pathlib import Path
from typing import Iterable
from tqdm import tqdm

# =============================================================================
# CONFIGURATION — edit these variables
# =============================================================================
CSV_PATH       = "/data/Deep_Angiography/Reports/Report_List_v01_01_merged_raw.csv"
NUM_REPORTS_TO_CLEAN = None    # number of reports to clean (None = all)
NUM_REPORTS_IN_DOCX  = 20     # number of reports to show in side-by-side DOCX
REPORT_COLUMN  = "radrpt"     # column name holding report text; None = auto-detect
OUTPUT_DOCX    = "/data/Deep_Angiography/Reports/report_comparison.docx"
OUTPUT_CSV     = "/data/Deep_Angiography/Reports/Report_List_v01_01_cleaned.csv"

# Parallelism: None = use os.cpu_count(); otherwise set an integer.
# Presidio + spaCy are CPU-bound and release the GIL poorly, so processes win over threads.
NUM_WORKERS    = None
# Chunk size for ProcessPoolExecutor.map — larger = less IPC overhead, smaller = smoother progress bar.
CHUNK_SIZE     = 16

# ---------- Presidio ----------
from presidio_analyzer import AnalyzerEngine, PatternRecognizer, Pattern
from presidio_anonymizer import AnonymizerEngine
from presidio_anonymizer.entities import OperatorConfig

# ---------- python-docx ----------
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


# =============================================================================
# 1. HARDCODED SAMPLE REPORT  (replace with your own for quick testing)
# =============================================================================
SAMPLE_REPORT = """\
PATIENT: DOE, JOHN A.          MRN: 01234567       DOB: 03/14/1958
ACCESSION #: A20240312-9981    DATE OF PROCEDURE: 03/12/2024
REFERRING PHYSICIAN: Dr. Sarah Mitchell, MD
ATTENDING: Dr. Robert J. Hansen, MD
LOCATION: UC Davis Medical Center, Sacramento, CA
PHONE: (916) 555-0142

HISTORY: Patient is a 66 year old male with known coronary artery disease,
hypertension, type 2 diabetes mellitus, and hyperlipidemia. He has a 40
pack-year smoking history. Previous cardiac catheterization in 2019 showed
non-obstructive disease. He also has a history of HCC involving Segment 5
treated with TACE in 2022.

PROCEDURE: Left heart catheterization with selective coronary angiography and
left ventriculography. DSA performed for vascular evaluation.

INDICATIONS: 66 y/o M w/ h/o HTN, DM2, HLD, and 40 pk-yr smoking hx, p/w
progressive exertional CP over the past 3 wks. Recent stress test was
equivocal. Pt referred for diagnostic cath. History of AKI and prior HCC
involving Segment 5 and Segment 8 of the liver.

HEMODYNAMICS:
    LVEDP: 18 mmHg
    Ao: 142/78 mmHg
    MAP: 99 mmHg
    LVEF (by LV gram): 45–50%
    EBL: 50 mL

CORONARY FINDINGS:
    LM:  Normal. No significant disease.
    LAD: 70–80% stenosis in the proximal segment, TIMI 3 flow.
         D1:  50% ostial lesion.
    LCX: Dominant vessel. 40% mid-segment plaque. No AV shunting.
         OM1: Non-obstructive.
    RCA: Non-dominant. Diffuse 30–40% disease throughout.
    GDA: Patent, no stenosis.
    No thrombus identified. No dissection.

IMPRESSION:
    1. Significant proximal LAD disease — recommend PCI w/ DES.
    2. Moderate non-obstructive LCX and RCA disease — medical management.
    3. Preserved LV systolic function.
    4. Prior HCC in Segment 5 — stable on IR follow-up.

PLAN:
    - Continue ASA 81 mg daily, add clopidogrel 75 mg daily.
    - Optimize statin therapy (atorvastatin 80 mg qhs).
    - Schedule staged PCI of proximal LAD within 1–2 wks.
    - F/u in clinic w/ Dr. Mitchell in 4 wks.

ATTESTATION:
I have personally reviewed the images and the above report. I was present for
the key portions of the procedure and agree with the findings and plan as
documented above.

Electronically signed by: Robert J. Hansen, MD
Date/Time signed: 03/12/2024 14:37
Dictated by: R. Hansen, MD       Transcribed by: MedScribe Services
---END OF REPORT---
"""


# =============================================================================
# 2. PRESIDIO SETUP — engine + custom medical recognizers
# =============================================================================
def build_analyzer() -> AnalyzerEngine:
    """Build a Presidio analyzer with default recognizers + medical-specific ones."""
    analyzer = AnalyzerEngine()

    # --- Medical Record Number (MRN) ---
    mrn_rec = PatternRecognizer(
        supported_entity="MRN",
        patterns=[
            Pattern(name="mrn_labeled", regex=r"\bMRN[:\s#]*\d{4,10}\b", score=0.95),
        ],
        context=["mrn", "medical record"],
    )

    # --- Accession number ---
    acc_rec = PatternRecognizer(
        supported_entity="ACCESSION",
        patterns=[
            Pattern(name="acc_labeled",
                    regex=r"\bACCESSION\s*#?[:\s]*[A-Z0-9\-]+\b", score=0.9),
            Pattern(name="acc_short",
                    regex=r"\bACC\s*#?[:\s]*[A-Z0-9\-]+\b", score=0.8),
        ],
        context=["accession"],
    )

    # --- DOB (more aggressive than default DATE_TIME) ---
    dob_rec = PatternRecognizer(
        supported_entity="DOB",
        patterns=[
            Pattern(name="dob_labeled",
                    regex=r"\bDOB[:\s]*\d{1,2}[/\-.]\d{1,2}[/\-.]\d{2,4}\b", score=0.95),
        ],
        context=["dob", "date of birth"],
    )

    # --- US-style phone numbers (Presidio has this, but we reinforce) ---
    phone_rec = PatternRecognizer(
        supported_entity="PHONE_NUMBER",
        patterns=[
            Pattern(name="phone_paren",
                    regex=r"\(\d{3}\)\s*\d{3}[-.\s]?\d{4}", score=0.9),
            Pattern(name="phone_dash",
                    regex=r"\b\d{3}[-.\s]\d{3}[-.\s]\d{4}\b", score=0.85),
        ],
    )

    # --- Medical facility names (catch common institutional patterns) ---
    facility_rec = PatternRecognizer(
        supported_entity="FACILITY",
        patterns=[
            Pattern(name="med_center",
                    regex=r"\b[A-Z][A-Za-z\.]+(?:\s+[A-Z][A-Za-z\.]+){0,4}\s+"
                          r"(?:Medical Center|Hospital|Clinic|Health System|"
                          r"Medical Centre|General Hospital)\b",
                    score=0.7),
        ],
    )

    for rec in (mrn_rec, acc_rec, dob_rec, phone_rec, facility_rec):
        analyzer.registry.add_recognizer(rec)

    return analyzer


def remove_phi(text: str,
               analyzer: AnalyzerEngine,
               anonymizer: AnonymizerEngine) -> str:
    """Detect and redact PHI from text using Presidio."""
    entities = [
        "PERSON", "DATE_TIME", "PHONE_NUMBER", "EMAIL_ADDRESS",
        "LOCATION", "US_SSN", "CREDIT_CARD", "IP_ADDRESS", "URL",
        "MRN", "ACCESSION", "DOB", "FACILITY",
    ]
    results = analyzer.analyze(text=text, language="en", entities=entities)

    operators = {
        "PERSON":        OperatorConfig("replace", {"new_value": "[PERSON]"}),
        "DATE_TIME":     OperatorConfig("replace", {"new_value": "[DATE]"}),
        "PHONE_NUMBER":  OperatorConfig("replace", {"new_value": "[PHONE]"}),
        "EMAIL_ADDRESS": OperatorConfig("replace", {"new_value": "[EMAIL]"}),
        "LOCATION":      OperatorConfig("replace", {"new_value": "[LOCATION]"}),
        "US_SSN":        OperatorConfig("replace", {"new_value": "[SSN]"}),
        "MRN":           OperatorConfig("replace", {"new_value": "[MRN]"}),
        "ACCESSION":     OperatorConfig("replace", {"new_value": "[ACCESSION]"}),
        "DOB":           OperatorConfig("replace", {"new_value": "[DOB]"}),
        "FACILITY":      OperatorConfig("replace", {"new_value": "[FACILITY]"}),
        "DEFAULT":       OperatorConfig("replace", {"new_value": "[REDACTED]"}),
    }

    return anonymizer.anonymize(
        text=text, analyzer_results=results, operators=operators
    ).text


# =============================================================================
# 3. ATTESTATION / SIGNATURE REMOVAL
# =============================================================================
# Patterns that mark the *start* of a trailing attestation/signature block.
# We find the earliest match and truncate everything from there.
_ATTESTATION_HEADS = [
    r"^\s*ATTEST(?:ATION)?\b",
    r"^\s*SIGN-?OFF\b",
    r"^\s*SIGNATURE\b",
    r"\bPreliminary\s+Report\s+Electronically\s+Signed\b",
    r"\bFinal\s+Report\s+Electronically\s+Signed\b",
    r"\bReport\s+Electronically\s+Signed\b",
    r"\bElectronically\s+signed\s+by\b",
    r"\bE-?signed\s+by\b",
    r"\bDigitally\s+signed\s+by\b",
    r"\bSigned\s+by\s*[:\-]",
    r"\bDictated\s+by\s*[:\-]",
    r"\bTranscribed\s+by\s*[:\-]",
    r"\bI\s+have\s+(?:personally\s+)?(?:reviewed|examined|seen)\b",
    r"\bI\s+(?:was|have\s+been)\s+present\b",
    r"\bI\s+agree\s+with\s+the\s+(?:findings|above|report|plan)\b",
    r"\bAttending\s+(?:attestation|note|statement)\b",
    r"^\s*---\s*END\s+OF\s+REPORT\s*---\s*$",
    r"\bDate/Time\s+signed\b",
]


def remove_attestation(text: str) -> str:
    """Truncate everything from the first attestation/signature marker onward."""
    earliest = len(text)
    for pat in _ATTESTATION_HEADS:
        for m in re.finditer(pat, text, flags=re.IGNORECASE | re.MULTILINE):
            if m.start() < earliest:
                earliest = m.start()
    return text[:earliest].rstrip()


# =============================================================================
# 3b. HEADER METADATA REMOVAL
# =============================================================================
# Clinical section headers that mark the start of the *actual* report body.
# Everything before the first match (patient/MRN/DOB/accession/physician/location/
# phone metadata) is discarded — it carries no clinical signal for training.
_BODY_START_HEADERS = [
    r"PROCEDURE",
    r"INDICATIONS?",
    r"HISTORY",
    r"CLINICAL\s+HISTORY",
    r"TECHNIQUE",
    r"FINDINGS",
    r"CORONARY\s+FINDINGS",
    r"HEMODYNAMICS",
    r"IMPRESSION",
    r"COMPARISON",
    r"EXAMINATION",
]


def remove_header_metadata(text: str) -> str:
    """
    Drop the patient-info / accession / physician / location header block
    at the top of the report by locating the first clinical section header
    and keeping only content from there onward.
    """
    pattern = r"(?im)^\s*(?:" + "|".join(_BODY_START_HEADERS) + r")\s*[:\-]"
    match = re.search(pattern, text)
    if match:
        return text[match.start():].lstrip()
    # No recognized header found — return unchanged so we don't nuke the report.
    return text


# =============================================================================
# 4. MEDICAL ABBREVIATION EXPANSION
# =============================================================================
# Ordered so multi-word / compound abbreviations are handled before shorter ones.
# RULE: Every abbreviation must be fully expanded — NO abbreviated forms in output.
_ABBREVIATIONS = [
    # -------------------------------------------------------------------------
    # Anatomy / Vessels — Coronary
    # -------------------------------------------------------------------------
    (r"\bLM\b",    "left main"),
    (r"\bLAD\b",   "left anterior descending"),
    (r"\bLCX\b",   "left circumflex"),
    (r"\bRCA\b",   "right coronary artery"),
    (r"\bPDA\b",   "posterior descending artery"),
    (r"\bPLV\b",   "posterior left ventricular"),
    (r"\bOM\d*\b", "obtuse marginal"),
    (r"\bD\d+\b",  "diagonal branch"),
    (r"\bRI\b",    "ramus intermedius"),

    # -------------------------------------------------------------------------
    # Anatomy / Vessels — Non-coronary / Interventional
    # -------------------------------------------------------------------------
    (r"\bGDA\b",   "gastroduodenal artery"),
    (r"\bRAS\b",   "renal artery stenosis"),
    (r"\bSMA\b",   "superior mesenteric artery"),
    (r"\bIMA\b",   "inferior mesenteric artery"),
    (r"\bCFA\b",   "common femoral artery"),
    (r"\bSFA\b",   "superficial femoral artery"),
    (r"\bIVC\b",   "inferior vena cava"),
    (r"\bSVC\b",   "superior vena cava"),
    (r"\bAV\s+shunting\b", "arteriovenous shunting"),

    # -------------------------------------------------------------------------
    # Function / Hemodynamics
    # -------------------------------------------------------------------------
    (r"\bLVEF\b",  "left ventricular ejection fraction"),
    (r"\bLVEDP\b", "left ventricular end-diastolic pressure"),
    (r"\bLV\s+gram\b", "left ventriculogram"),
    (r"\bAo\b",    "aorta"),
    (r"\bMAP\b",   "mean arterial pressure"),
    (r"\bTIMI\b",  "thrombolysis in myocardial infarction grade"),

    # -------------------------------------------------------------------------
    # Conditions / Pathology
    # -------------------------------------------------------------------------
    (r"\bHTN\b",   "hypertension"),
    (r"\bDM2\b",   "type 2 diabetes mellitus"),
    (r"\bDM\b",    "diabetes mellitus"),
    (r"\bHLD\b",   "hyperlipidemia"),
    (r"\bCAD\b",   "coronary artery disease"),
    (r"\bCHF\b",   "congestive heart failure"),
    (r"\bMI\b",    "myocardial infarction"),
    (r"\bCP\b",    "chest pain"),
    (r"\bSOB\b",   "shortness of breath"),
    (r"\bHCC\b",   "hepatocellular carcinoma"),
    (r"\bUGI\s+bleed\b", "upper gastrointestinal bleed"),
    (r"\bAKI\b",   "acute kidney injury"),
    (r"\bESRD\b",  "end-stage renal disease"),
    (r"\bNSTEMI\b","non-ST-elevation myocardial infarction"),
    (r"\bSTEMI\b", "ST-elevation myocardial infarction"),
    (r"\bPE\b",    "pulmonary embolism"),
    (r"\bDVT\b",   "deep vein thrombosis"),

    # -------------------------------------------------------------------------
    # Procedures / Techniques
    # -------------------------------------------------------------------------
    (r"\bPCI\b",   "percutaneous coronary intervention"),
    (r"\bCABG\b",  "coronary artery bypass grafting"),
    (r"\bIR\b",    "interventional radiology"),
    (r"\bDSA\b",   "digital subtraction angiography"),
    (r"\bEGD\b",   "esophagogastroduodenoscopy"),
    (r"\bMRA\b",   "magnetic resonance angiography"),
    (r"\bTACE\b",  "transarterial chemoembolization"),
    (r"\bTARE\b",  "transarterial radioembolization"),

    # -------------------------------------------------------------------------
    # Devices / Stents
    # -------------------------------------------------------------------------
    (r"\bDES\b",   "drug-eluting stent"),
    (r"\bBMS\b",   "bare-metal stent"),

    # -------------------------------------------------------------------------
    # Drugs
    # -------------------------------------------------------------------------
    (r"\bASA\b",   "aspirin"),

    # -------------------------------------------------------------------------
    # Measurements / Units
    # -------------------------------------------------------------------------
    (r"\bEBL\b",   "estimated blood loss"),
    (r"\bmGy\b",   "milligray"),
    (r"\bcGycm",   "centigray-cm"),
    (r"\bmmHg\b",  "millimeters of mercury"),
    (r"(\d+)\s*French\b", r"\1 French (catheter size)"),
    (r"(\d+)\s*F\b(?!\w)", r"\1 French (catheter size)"),

    # -------------------------------------------------------------------------
    # Couinaud liver segmentation — make laymen-friendly
    # Single pattern (IGNORECASE handles both "Segment" and "segment").
    # We anchor with \bSegment to avoid re-matching inside already-expanded text.
    # -------------------------------------------------------------------------
    (r"\bSegment\s+([1-8])\b",    r"Couinaud liver segment \1"),
    (r"\bSeg\.?\s*([1-8])\b",     r"Couinaud liver segment \1"),

    # Standalone LV (must come after LV-compound terms above)
    (r"\bLV\b",    "left ventricular"),

    # -------------------------------------------------------------------------
    # Dictation shorthand
    # -------------------------------------------------------------------------
    (r"\bw/o\b",   "without"),           # must come before w/ to avoid partial match
    (r"\bw/(?=\s|$)", "with"),
    (r"\bh/o\b",   "history of"),
    (r"\bs/p\b",   "status post"),
    (r"\bp/w\b",   "presenting with"),
    (r"\by/?o\b",  "year old"),
    (r"\bpk-?yr\b","pack-year"),
    (r"\bhx\b",    "history"),
    (r"\bf/u\b",   "follow up"),
    (r"\bwk(s)?\b", r"week\1"),
    (r"\bqhs\b",   "at bedtime"),
    (r"\bpt\b",    "patient"),
]


def expand_abbreviations(text: str) -> str:
    """Expand common cardiology/angiography/IR abbreviations and dictation shorthand."""
    for pat, repl in _ABBREVIATIONS:
        text = re.sub(pat, repl, text, flags=re.IGNORECASE)
    return text


# =============================================================================
# 5. UNICODE, PUNCTUATION, WHITESPACE
# =============================================================================
def normalize_unicode(text: str) -> str:
    """NFKC normalization + replace common unicode punctuation with ASCII."""
    text = unicodedata.normalize("NFKC", text)
    replacements = {
        "\u2013": "-", "\u2014": "-",          # en/em dash
        "\u2018": "'", "\u2019": "'",          # smart single quotes
        "\u201C": '"', "\u201D": '"',          # smart double quotes
        "\u2026": "...",                        # ellipsis
        "\u00B0": " degrees ",                  # degree sign
        "\u00B1": "+/-",                        # plus-minus
        "\u00D7": "x",                          # multiplication
        "\u2192": "->",                         # right arrow
        "\u00A0": " ",                          # non-breaking space
    }
    for u, a in replacements.items():
        text = text.replace(u, a)
    return text


def clean_symbols_and_lists(text: str) -> str:
    """Strip stray bullets, list markers, repeated punctuation, boilerplate rules."""
    text = re.sub(r"^[\s]*[\*\u2022\-\u2013\u2014]+\s+", "", text, flags=re.MULTILINE)
    text = re.sub(r"^\s*\d+\.\s+", "", text, flags=re.MULTILINE)           # "1. " list markers
    text = re.sub(r"-{3,}.*?-{3,}", "", text)                               # --- RULES ---
    text = re.sub(r"[=_]{3,}", "", text)                                    # ===, ___
    text = re.sub(r"([!?.,;:])\1{1,}", r"\1", text)                         # collapse !!!, ..
    return text


def collapse_whitespace(text: str) -> str:
    """Trim, collapse runs of spaces/tabs, and limit consecutive blank lines to one."""
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r" *\n *", "\n", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def normalize_section_headers(text: str) -> str:
    """Standardize common section headers — keep them in BLOCK LETTERS with a colon,
    and ensure a blank line precedes each section for visual separation."""
    headers = [
        "PROCEDURE", "INDICATIONS", "HEMODYNAMICS", "CORONARY FINDINGS",
        "FINDINGS", "IMPRESSION", "PLAN", "HISTORY", "TECHNIQUE",
        "COMPARISON", "COMPLICATIONS",
    ]
    for h in headers:
        # Header alone on a line (with or without colon/dash)
        text = re.sub(
            rf"(?im)^\s*{h}\s*[:\-]?\s*$",
            f"{h}:",
            text,
        )
        # Header followed by content on the same line
        text = re.sub(
            rf"(?im)^\s*{h}\s*[:\-]\s*",
            f"{h}: ",
            text,
        )

    # Insert a blank line before each section header (unless it's already at the
    # start of the text or already preceded by a blank line).
    header_pat = "|".join(re.escape(h) for h in headers)
    text = re.sub(
        rf"(?m)(?<!\n\n)^((?:{header_pat}):)",
        r"\n\1",
        text,
    )

    return text.lstrip("\n")   # don't start with a blank line


# =============================================================================
# 6b. REMOVE HISTORY SECTION
# =============================================================================
# The recognized section headers used to detect where HISTORY ends.
_SECTION_HEADERS_PATTERN = (
    r"(?:PROCEDURE|INDICATIONS?|HEMODYNAMICS|CORONARY\s+FINDINGS|FINDINGS|"
    r"IMPRESSION|PLAN|TECHNIQUE|COMPARISON|COMPLICATIONS|EXAMINATION)\s*[:\-]"
)


def remove_history_section(text: str) -> str:
    """Remove the HISTORY section entirely (header through to the next section)."""
    pattern = re.compile(
        rf"(?ims)"
        rf"^\s*HISTORY\s*[:\-].*?"               # HISTORY header + its content
        rf"(?=^\s*{_SECTION_HEADERS_PATTERN}|\Z)",  # stop at next section or EOF
        re.MULTILINE,
    )
    return pattern.sub("", text).strip()


# =============================================================================
# 6. SENTENCE-CASE CONVERSION
# =============================================================================
def sentence_case(text: str) -> str:
    """
    Convert text to sentence case while preserving words that are fully
    capitalized (block letters, 2+ characters) — these carry domain meaning
    (e.g., grading systems, device names, units).

    Logic:
        - Words that are ALL UPPERCASE and >= 2 letters → keep as-is.
        - Everything else → lowercase.
        - Then capitalize the first letter of each sentence.
        - Restore redaction placeholders like [person] → [PERSON].
    """

    def _selective_lower(word: str) -> str:
        """Lowercase a word unless it is fully capitalized (2+ alpha chars)."""
        # Check if the word (stripping non-alpha) is all uppercase
        alpha_only = re.sub(r"[^A-Za-z]", "", word)
        if len(alpha_only) >= 2 and alpha_only.isupper():
            return word  # preserve as-is
        return word.lower()

    # Apply selective lowering word-by-word, preserving whitespace structure
    result = re.sub(r"[^\s]+", lambda m: _selective_lower(m.group(0)), text)

    # Capitalize first letter after sentence terminators or newlines
    def _cap(match: re.Match) -> str:
        return match.group(1) + match.group(2).upper()

    result = re.sub(r"^(\s*)([a-z])", _cap, result)
    result = re.sub(r"([.!?]\s+|\n+\s*)([a-z])", _cap, result)

    # Restore redaction placeholders like [person] -> [PERSON]
    result = re.sub(
        r"\[([a-z_][a-z0-9_]*)\]",
        lambda m: f"[{m.group(1).upper()}]",
        result,
    )
    return result


# =============================================================================
# 7. FULL PIPELINE
# =============================================================================
def clean_report(text: str,
                 analyzer: AnalyzerEngine,
                 anonymizer: AnonymizerEngine) -> str:
    """Apply the full cleaning pipeline in order."""
    text = normalize_unicode(text)
    text = remove_phi(text, analyzer, anonymizer)
    text = remove_attestation(text)
    text = remove_header_metadata(text)
    text = expand_abbreviations(text)
    text = clean_symbols_and_lists(text)
    text = collapse_whitespace(text)
    text = normalize_section_headers(text)
    text = sentence_case(text)
    text = collapse_whitespace(text)   # final tidy after case changes
    return text


# =============================================================================
# 7b. PARALLEL WORKER SETUP
# =============================================================================
# Presidio's AnalyzerEngine and AnonymizerEngine are heavy to construct (they
# load spaCy models) and not trivially picklable. We build them ONCE per worker
# process in an initializer and reuse them for every task that worker handles.
_WORKER_ANALYZER: AnalyzerEngine | None = None
_WORKER_ANONYMIZER: AnonymizerEngine | None = None


def _worker_init() -> None:
    """Initializer run once per worker process — builds Presidio engines."""
    global _WORKER_ANALYZER, _WORKER_ANONYMIZER
    _WORKER_ANALYZER = build_analyzer()
    _WORKER_ANONYMIZER = AnonymizerEngine()


def _clean_one(text: str) -> str:
    """Top-level (picklable) function that cleans a single report in a worker."""
    if not text:
        return ""
    # These are guaranteed non-None because _worker_init runs before any task.
    return clean_report(text, _WORKER_ANALYZER, _WORKER_ANONYMIZER)


# =============================================================================
# 8. DOCX SIDE-BY-SIDE OUTPUT
# =============================================================================
def _set_cell_shading(cell, fill_hex: str) -> None:
    """Set a table cell's background color."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tc_pr.append(shd)


def _set_cell_borders(cell, color: str = "BFBFBF", size: str = "6") -> None:
    """Thin gray borders on all 4 sides of a cell."""
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), size)
        border.set(qn("w:color"), color)
        tc_borders.append(border)
    tc_pr.append(tc_borders)


def _write_multiline(cell, text: str, mono: bool = True, size: int = 9) -> None:
    """Write text into a cell preserving newlines as separate paragraphs."""
    cell.text = ""  # clear default empty paragraph
    first = True
    for line in text.split("\n"):
        para = cell.paragraphs[0] if first else cell.add_paragraph()
        first = False
        para.paragraph_format.space_after = Pt(2)
        para.paragraph_format.space_before = Pt(0)
        run = para.add_run(line)
        run.font.size = Pt(size)
        if mono:
            run.font.name = "Consolas"
            # Also set east-asian font so Word doesn't override
            rpr = run._element.get_or_add_rPr()
            rfonts = rpr.find(qn("w:rFonts"))
            if rfonts is None:
                rfonts = OxmlElement("w:rFonts")
                rpr.append(rfonts)
            rfonts.set(qn("w:ascii"), "Consolas")
            rfonts.set(qn("w:hAnsi"), "Consolas")


def write_comparison_docx(pairs: list[tuple[str, str]],
                          out_path: str = "report_comparison.docx") -> Path:
    """Build a landscape .docx with a side-by-side table per report.

    Args:
        pairs: list of (original_text, cleaned_text) tuples.
        out_path: where to save the DOCX.
    """
    doc = Document()

    # --- Landscape US Letter with 0.7" margins ---
    section = doc.sections[0]
    section.page_width, section.page_height = Inches(11), Inches(8.5)
    section.orientation = 1  # landscape
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    # --- Document title ---
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    trun = title.add_run("Angiographic Report Cleaning — Original vs. Cleaned")
    trun.bold = True
    trun.font.size = Pt(14)
    trun.font.name = "Calibri"

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    srun = subtitle.add_run(
        "PHI redacted via Microsoft Presidio · attestation stripped · "
        "abbreviations expanded · sentence-cased"
    )
    srun.italic = True
    srun.font.size = Pt(10)
    srun.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    content_width = Inches(11 - 1.4)   # page width - L/R margins
    col_width = content_width / 2

    for idx, (original, cleaned) in tqdm(enumerate(pairs, start=1),
                                         total=len(pairs),
                                         desc="Writing DOCX",
                                         unit="report"):
        # --- Report heading ---
        doc.add_paragraph()  # spacer
        heading = doc.add_paragraph()
        hrun = heading.add_run(f"Report {idx} of {len(pairs)}")
        hrun.bold = True
        hrun.font.size = Pt(12)
        hrun.font.name = "Calibri"
        hrun.font.color.rgb = RGBColor(0x1F, 0x4E, 0x78)

        # --- Comparison table (header row + content row) ---
        table = doc.add_table(rows=2, cols=2)
        table.autofit = False
        for row in table.rows:
            for cell in row.cells:
                cell.width = col_width

        # Header row
        hdr_cells = table.rows[0].cells
        for cell, label, fill in (
            (hdr_cells[0], "Original Report", "1F4E78"),
            (hdr_cells[1], "Cleaned Report",  "2E7D32"),
        ):
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            _set_cell_shading(cell, fill)
            _set_cell_borders(cell)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(label)
            r.bold = True
            r.font.size = Pt(11)
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            r.font.name = "Calibri"

        # Content row
        body_cells = table.rows[1].cells
        _set_cell_shading(body_cells[0], "FDF6E3")
        _set_cell_shading(body_cells[1], "F1F8E9")
        for cell in body_cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            _set_cell_borders(cell)

        _write_multiline(body_cells[0], original, mono=True, size=9)
        _write_multiline(body_cells[1], cleaned,  mono=True, size=9)

        # Page break between reports (except after the last one)
        if idx < len(pairs):
            doc.add_page_break()

    out = Path(out_path).resolve()
    doc.save(out)
    return out


# =============================================================================
# 10. CSV READING UTILITIES
# =============================================================================
# Common column names for report text (case-insensitive matching).
_REPORT_COL_CANDIDATES = [
    "report", "report_text", "radrpt", "text", "report_body", "body",
    "findings", "narrative", "clinical_text", "raw_report",
    "radiology_report", "report_content",
]


def _detect_report_column(headers: list[str]) -> str:
    """Auto-detect which CSV column contains the report text."""
    lower_map = {h.lower().strip(): h for h in headers}
    for candidate in _REPORT_COL_CANDIDATES:
        if candidate in lower_map:
            return lower_map[candidate]
    # Fallback: pick the column with the longest average content
    raise ValueError(
        f"Could not auto-detect the report column. "
        f"Available columns: {headers}\n"
        f"Set REPORT_COLUMN at the top of the script to the correct column name."
    )


def read_reports_from_csv(csv_path: str,
                          column: str | None = None,
                          n: int | None = None) -> list[str]:
    """Read report texts from a CSV file.

    Args:
        csv_path: path to the CSV.
        column: column name holding report text (None = auto-detect).
        n: max number of rows to read (None = all).

    Returns:
        (rows, headers, report_column_name) where rows is a list of dicts.
    """
    path = Path(csv_path)
    if not path.exists():
        raise FileNotFoundError(f"CSV not found: {path}")

    with open(path, newline="", encoding="utf-8", errors="replace") as f:
        reader = csv.DictReader(f)
        headers = reader.fieldnames or []
        col = column if column else _detect_report_column(headers)
        if col not in headers:
            raise KeyError(
                f"Column '{col}' not found in CSV. Available: {headers}"
            )
        print(f"       Using column: '{col}'")

        rows = []
        for row in reader:
            rows.append(row)
            if n is not None and len(rows) >= n:
                break

    return rows, headers, col


def write_cleaned_csv(rows: list[dict],
                      headers: list[str],
                      report_col: str,
                      cleaned_reports: list[str],
                      out_path: str) -> Path:
    """Write a new CSV with the original columns + a 'cleaned_radrpt' column.

    Args:
        rows: original CSV rows (list of dicts).
        headers: original CSV column names.
        report_col: name of the report column.
        cleaned_reports: cleaned report texts (same length as rows).
        out_path: output CSV path.
    """
    cleaned_col = f"cleaned_{report_col}"
    out_headers = headers + [cleaned_col]

    path = Path(out_path)
    path.parent.mkdir(parents=True, exist_ok=True)

    with open(path, "w", newline="", encoding="utf-8") as f:
        writer = csv.DictWriter(f, fieldnames=out_headers)
        writer.writeheader()
        for row, cleaned in zip(rows, cleaned_reports):
            row[cleaned_col] = cleaned
            writer.writerow(row)

    return path.resolve()


# =============================================================================
# 11. ENTRYPOINT
# =============================================================================
def main() -> None:
    num_workers = NUM_WORKERS or os.cpu_count() or 1

    print(f"[1/5] Reading reports from: {CSV_PATH}")
    rows, headers, report_col = read_reports_from_csv(
        CSV_PATH, column=REPORT_COLUMN, n=NUM_REPORTS_TO_CLEAN
    )
    print(f"       Loaded {len(rows)} row(s).")

    if not rows:
        print("ERROR: No reports found. Check CSV_PATH and REPORT_COLUMN.")
        return

    # Extract raw texts up front so we preserve order and don't ship whole
    # dict rows through IPC (much smaller pickled payload).
    originals = [(row.get(report_col) or "").strip() for row in rows]

    print(f"[2/5] Spinning up {num_workers} worker process(es) "
          f"(each loads Presidio once — 10–30s)…")
    print(f"[3/5] Cleaning {len(originals)} report(s) in parallel…")

    cleaned_reports: list[str] = [""] * len(originals)
    with ProcessPoolExecutor(
        max_workers=num_workers,
        initializer=_worker_init,
    ) as executor:
        # executor.map preserves input order, which is critical for zipping
        # cleaned_reports back to rows when writing the CSV.
        results_iter = executor.map(_clean_one, originals, chunksize=CHUNK_SIZE)
        for i, cleaned in enumerate(
            tqdm(results_iter, total=len(originals),
                 desc="Cleaning reports", unit="report")
        ):
            cleaned_reports[i] = cleaned

    # --- Save full cleaned CSV (all reports) ---
    print(f"[4/5] Writing cleaned CSV → {OUTPUT_CSV}")
    csv_out = write_cleaned_csv(rows, headers, report_col, cleaned_reports, OUTPUT_CSV)
    print(f"       Saved {len(rows)} rows with '{f'cleaned_{report_col}'}' column.")

    # --- Save DOCX comparison (first NUM_REPORTS_IN_DOCX only) ---
    n_docx = min(NUM_REPORTS_IN_DOCX, len(rows))
    print(f"[5/5] Writing side-by-side DOCX (first {n_docx} reports)…")
    pairs = [(originals[i], cleaned_reports[i]) for i in range(n_docx)]
    out = write_comparison_docx(pairs, OUTPUT_DOCX)

    print(f"\nDone.")
    print(f"  CSV:  {csv_out}  ({len(rows)} reports)")
    print(f"  DOCX: {out}  ({n_docx} side-by-side comparisons)")


if __name__ == "__main__":
    main()