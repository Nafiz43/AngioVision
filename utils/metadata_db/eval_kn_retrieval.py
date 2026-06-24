#!/usr/bin/env python3
"""
K@N Cross-Validated Sequence Retrieval Evaluation — AngioVision DSA
====================================================================

Pipeline
────────
  1. Load labeled CSV; exclude 'other' rows; group by angio_run.
     Filter sequences whose required frame annotation is missing (mode-dependent)
     and save them to a sidecar CSV next to the labeled CSV.

  2. Build n-fold (--n-folds, default 10) stratified CV splits.

  3. Load embedding model (--model).

  4. Precompute ALL sequence embeddings ONCE — reused across every fold
     (10-fold CV would otherwise embed the same sequence 10 times).

  5. For each fold:
       a. Ingest TRAIN embeddings from precomputed cache → fresh in-memory ChromaDB
       b. Query each TEST sequence for all K values
       c. Majority vote → predicted sequence → Hit if label matches

  6. Aggregate across folds: pooled counts + per-fold accuracy list (mean±std).

  7. Console table + bar chart with error bars + Markdown with mean±std cells
     + Word doc with K=1 visual retrieval examples (from the last fold).

Frame modes (--frame-mode)
──────────────────────────
  best   Use only the Best_Image frame (1 frame per sequence)
  fl     Use frames from First_Diag_Image to Last_Diag_Image  [default]
  all    Use every frame (optionally capped by --max-frames)

Split modes (--split-mode)
──────────────────────────
  cv       n-fold stratified cross-validation (--n-folds)              [default]
  holdout  single stratified train/test split (--scale-down train fraction)
           — reproduces the legacy eval_k1*.py 80/20 experiments

Embedding models (--model)
──────────────────────────
  rad-dino        microsoft/rad-dino          768-dim  [default]
  vit-b16         google/vit-base-patch16-224 768-dim
  vit-l16         google/vit-large-patch16-224 1024-dim
  openclip-b32    OpenCLIP ViT-B-32 / laion2b  512-dim
  openclip-l14    OpenCLIP ViT-L-14 / laion2b  768-dim
  <any HF model>  Any HuggingFace AutoModel ID

Usage
─────
  python eval_kn_retrieval.py
  python eval_kn_retrieval.py --model vit-b16 --frame-mode fl --n-folds 10
  python eval_kn_retrieval.py --model openclip-b32 --frame-mode all --max-frames 20
  python eval_kn_retrieval.py --temporal --n-folds 5
  python eval_kn_retrieval.py --k-values 1 3 5 --workers 8
"""

import os, sys, csv, random, logging, argparse, sqlite3, datetime, itertools, io
from pathlib import Path
from collections import Counter, defaultdict
from concurrent.futures import ThreadPoolExecutor
from typing import Optional, Callable

import numpy as np
from tqdm import tqdm

logging.basicConfig(level=logging.INFO, format="%(asctime)s [%(levelname)s] %(message)s")
log = logging.getLogger(__name__)

# ── Required deps ─────────────────────────────────────────────────────────────
try:
    import pydicom
except ImportError:
    print("ERROR: pip install pydicom"); sys.exit(1)

try:
    import chromadb
except ImportError:
    print("ERROR: pip install chromadb"); sys.exit(1)

try:
    import torch
    from PIL import Image as PILImage
    from transformers import AutoImageProcessor, AutoModel
    HF_OK = True
except ImportError:
    HF_OK = False

try:
    import open_clip
    OPENCLIP_OK = True
except ImportError:
    OPENCLIP_OK = False

try:
    import matplotlib.pyplot as plt
    import matplotlib.patches as mpatches
    MATPLOTLIB_OK = True
except ImportError:
    MATPLOTLIB_OK = False

# ── Constants ─────────────────────────────────────────────────────────────────
DEFAULT_LABELED_CSV = "/data/Deep_Angiography/labeled_DSA_2023_10_24.csv"
DEFAULT_DICOM_ROOT  = "/data/Deep_Angiography/DICOM"
DEFAULT_SQLITE_DB   = "/data/Deep_Angiography/AngioVision/dicom_staging.db"
DEFAULT_OUT_PLOT    = None
DEFAULT_OUT_MD      = None
DEFAULT_K_VALUES    = [1, 3, 5, 7, 9, 11, 13, 15]
DEFAULT_WORKERS     = max(1, (os.cpu_count() or 4) // 2)
CHROMA_COLLECTION   = "eval_retrieval"
FRAME_MODES         = ("best", "fl", "all")

HF_ALIASES: dict[str, str] = {
    "rad-dino": "microsoft/rad-dino",
    "vit-b16":  "google/vit-base-patch16-224",
    "vit-l16":  "google/vit-large-patch16-224",
}
OPENCLIP_ALIASES: dict[str, tuple[str, str]] = {
    "openclip-b32": ("ViT-B-32", "laion2b_s34b_b79k"),
    "openclip-l14": ("ViT-L-14", "laion2b_s32b_b82k"),
}

CODE_MAP: dict[str, str] = {
    "intrahepatic artery":                             "IHA",
    "external iliac artery, right":                    "EIA-R",
    "celiac trunk":                                    "CT",
    "nondiagnostic":                                   "ND",
    "common hepatic artery":                           "CHA",
    "hepatic artery, right":                           "HA-R",
    "superior mesenteric artery (SMA)":                "SMA",
    "proper hepatic artery":                           "PHA",
    "lower abdominal aorta and aortic bifurcation":    "LAA",
    "hepatic artery, left":                            "HA-L",
    "splenic artery":                                  "SA",
    "internal iliac artery, left":                     "IIA-L",
    "inferior mesenteric artery (IMA)":                "IMA",
    "external iliac artery, left":                     "EIA-L",
    "internal iliac artery, right":                    "IIA-R",
    "renal artery, right":                             "RA-R",
    "unstable":                                        "UNS",
    "upper abdominal aorta":                           "UAA",
    "renal artery, left":                              "RA-L",
    "TRAS, extrenal iliac artery, right":              "TRAS-R",
    "common iliac artery, right":                      "CIA-R",
    "common iliac artery, left":                       "CIA-L",
    "common femoral artery, left":                     "CFA-L",
    "TRAS, extrenal iliac artery, left":               "TRAS-L",
}

def code(label: str) -> str:
    return CODE_MAP.get(label, label[:8])

def _auto_output_paths(
    model_alias: str, frame_mode: str, temporal: bool, split_tag: str,
) -> tuple[str, str, str]:
    """Build self-documenting filenames: kn_results_{model}_{mode}_{split_tag}[_temporal].ext"""
    safe = "".join(c if c.isalnum() or c == "-" else "-" for c in model_alias)
    while "--" in safe: safe = safe.replace("--", "-")
    safe = safe.strip("-") or "model"
    parts = ["kn_results", safe, frame_mode, split_tag]
    if temporal: parts.append("temporal")
    stem = "_".join(parts)
    return f"{stem}.md", f"{stem}.png", f"{stem}.docx"


# ═══════════════════════════════════════════════════════════════════════════════
# 1 ── CSV loading  (mode-aware filtering)
# ═══════════════════════════════════════════════════════════════════════════════

def load_labeled_csv_grouped(csv_path: Path, frame_mode: str) -> dict[str, list[dict]]:
    groups: dict[str, list[dict]] = defaultdict(list)
    seen: set[str] = set()
    missing_rows: list[dict] = []
    fieldnames: list[str] = []
    skip_other = skip_path = skip_dup = 0

    with open(csv_path, newline="", encoding="utf-8-sig") as fh:
        reader = csv.DictReader(fh)
        if not reader.fieldnames: raise ValueError(f"Empty CSV: {csv_path}")
        fieldnames = list(reader.fieldnames)
        norm = {h.strip().lower(): h for h in fieldnames}
        def col(k): return norm.get(k.lower())

        path_key  = col("file_path");  angio_key = col("angio_run")
        acc_key   = col("accession");  ser_key   = col("seriesuid")
        run_key   = col("run_type");   best_key  = col("best_image")
        first_key = col("first_diag_image"); last_key = col("last_diag_image")

        if not path_key: raise ValueError("'file_path' column not found")

        def _p1(raw):
            try:
                v = float((raw or "").strip())
                return int(v) - 1 if v >= 1 else -1
            except (ValueError, TypeError): return -1

        for row in reader:
            fp    = (row.get(path_key)  or "").strip()
            angio = (row.get(angio_key) or "").strip() if angio_key else ""
            if not fp:                skip_path  += 1; continue
            if "other" in angio.lower(): skip_other += 1; continue
            if fp in seen:            skip_dup   += 1; continue
            seen.add(fp)

            best_image_idx = _p1(row.get(best_key,  "") if best_key  else "")
            first_diag_idx = _p1(row.get(first_key, "") if first_key else "")
            last_diag_idx  = _p1(row.get(last_key,  "") if last_key  else "")
            if first_diag_idx >= 0 and last_diag_idx >= 0 and last_diag_idx < first_diag_idx:
                first_diag_idx = last_diag_idx = -1

            exclude = False
            if frame_mode == "best" and best_image_idx < 0: exclude = True
            elif frame_mode == "fl" and (first_diag_idx < 0 or last_diag_idx < 0): exclude = True
            if exclude: missing_rows.append(dict(row)); continue

            groups[angio].append({
                "accession": (row.get(acc_key) or "").strip() if acc_key else "",
                "series_uid":(row.get(ser_key) or "").strip() if ser_key else "",
                "run_type":  (row.get(run_key) or "").strip() if run_key else "",
                "angio_run": angio, "file_path": fp,
                "best_image_idx": best_image_idx,
                "first_diag_idx": first_diag_idx,
                "last_diag_idx":  last_diag_idx,
            })

    if missing_rows:
        sidecar = csv_path.parent / {"best":"missing_best_image.csv","fl":"missing_diag_window.csv"}.get(frame_mode,"missing_frames.csv")
        with open(sidecar, "w", newline="", encoding="utf-8") as fh:
            w = csv.DictWriter(fh, fieldnames=fieldnames); w.writeheader(); w.writerows(missing_rows)
        log.warning(f"  {len(missing_rows):,} sequences excluded (mode={frame_mode}) → {sidecar}")

    total = sum(len(v) for v in groups.values())
    log.info(f"CSV loaded — {len(groups)} categories, {total:,} sequences retained "
             f"(skipped other={skip_other:,}, dup={skip_dup:,}, no-path={skip_path:,}, no-annotation={len(missing_rows):,})")
    return dict(groups)


# ═══════════════════════════════════════════════════════════════════════════════
# 2 ── DICOM index
# ═══════════════════════════════════════════════════════════════════════════════

def _stem(fp: str) -> Optional[str]:
    p = (fp or "").strip()
    if not p: return None
    fname = p.rsplit("/", 1)[-1] if "/" in p else p
    return fname[:-4] if fname.lower().endswith(".dcm") else (fname or None)

def build_dicom_index(dicom_root: Path, sqlite_db: Optional[Path] = None) -> dict[str, str]:
    if sqlite_db and sqlite_db.exists():
        try:
            con = sqlite3.connect(str(sqlite_db))
            n = con.execute("SELECT COUNT(*) FROM dicom_files WHERE source_file IS NOT NULL").fetchone()[0]
            if n > 0:
                log.info(f"Building DICOM index from SQLite ({n:,} rows) …")
                idx: dict[str, str] = {}
                for sf, sp in con.execute("SELECT source_file, source_path FROM dicom_files WHERE source_file IS NOT NULL AND source_path IS NOT NULL"):
                    s = Path(sf).stem
                    if s and s not in idx: idx[s] = sp
                con.close(); log.info(f"Index ready: {len(idx):,} stems (SQLite)"); return idx
            con.close()
        except Exception as e: log.warning(f"SQLite index failed ({e}) — filesystem walk …")
    log.info(f"Walking {dicom_root} …")
    idx = {}; dups = 0
    for dp, _, fnames in os.walk(str(dicom_root)):
        for fn in fnames:
            if fn.lower().endswith(".dcm"):
                s = Path(fn).stem
                if s in idx: dups += 1
                else: idx[s] = str(Path(dp) / fn)
    log.info(f"Index ready: {len(idx):,} stems ({dups} dups ignored)"); return idx

def resolve_paths(groups: dict[str, list[dict]], idx: dict[str, str]) -> dict[str, list[dict]]:
    resolved: dict[str, list[dict]] = {}; n_missing = 0
    for label, seqs in groups.items():
        kept = []
        for seq in seqs:
            s = _stem(seq["file_path"])
            if s and s in idx: kept.append({**seq, "dicom_path": idx[s], "stem": s})
            else: n_missing += 1
        if kept: resolved[label] = kept
    log.info(f"Path resolution — found: {sum(len(v) for v in resolved.values()):,}, missing (dropped): {n_missing:,}")
    return resolved


# ═══════════════════════════════════════════════════════════════════════════════
# 3 ── Cross-validation folds
# ═══════════════════════════════════════════════════════════════════════════════

def create_cv_folds(
    groups: dict[str, list[dict]], n_folds: int, min_seqs: int, seed: int,
) -> list[dict[str, tuple[list, list]]]:
    """
    Stratified n-fold CV: sequence j → fold (j % n_folds).
    Every sequence appears in the test set exactly once across all folds.
    Categories with fewer than max(min_seqs, n_folds) sequences are excluded.
    Returns list of n_folds splits, each { label: (train_list, test_list) }.
    """
    rng = random.Random(seed)
    kept: dict[str, list[dict]] = {}
    skipped: list[tuple[str, int]] = []
    threshold = max(min_seqs, n_folds)

    for label, seqs in sorted(groups.items(), key=lambda x: -len(x[1])):
        if len(seqs) < threshold: skipped.append((label, len(seqs))); continue
        s = list(seqs); rng.shuffle(s); kept[label] = s
        log.info(f"  {code(label):8s}  total={len(s):4d}  ~train={len(s)*(n_folds-1)//n_folds:4d}  ~test={len(s)//n_folds:4d}/fold")

    if skipped:
        log.info(f"  CV skipped (< {threshold} seqs): " + ", ".join(f"{code(l)}({n})" for l, n in skipped))
    log.info(f"CV: {n_folds} folds × {len(kept)} categories ({sum(len(s) for s in kept.values()):,} sequences)")

    folds: list[dict[str, tuple[list, list]]] = []
    for fold_idx in range(n_folds):
        fold_splits: dict[str, tuple[list, list]] = {}
        for label, seqs in kept.items():
            test  = [seqs[j] for j in range(len(seqs)) if     j % n_folds == fold_idx]
            train = [seqs[j] for j in range(len(seqs)) if not j % n_folds == fold_idx]
            if test and train: fold_splits[label] = (train, test)
        folds.append(fold_splits)
    return folds


def create_holdout_split(
    groups: dict[str, list[dict]], scale_down: float, min_seqs: int, seed: int,
) -> list[dict[str, tuple[list, list]]]:
    """
    Single stratified holdout split (replicates the legacy eval_k1*.py behaviour).
    scale_down is the TRAIN fraction per category (e.g. 0.80 → 80% train / 20% test).
    Categories with fewer than min_seqs sequences are excluded.
    Returns a one-element list of { label: (train_list, test_list) } so the
    downstream fold loop / aggregation can treat it as a single "fold".
    """
    rng = random.Random(seed)
    split: dict[str, tuple[list, list]] = {}
    skipped: list[tuple[str, int]] = []

    log.info(f"Holdout split (scale_down={scale_down:.0%} train) …")
    for label, seqs in sorted(groups.items(), key=lambda x: -len(x[1])):
        n = len(seqs)
        if n < min_seqs: skipped.append((label, n)); continue
        shuffled = list(seqs); rng.shuffle(shuffled)
        n_train = max(1, round(n * scale_down))
        n_test  = n - n_train
        if n_test == 0: n_train -= 1; n_test = 1   # guarantee ≥1 test sequence
        if n_train < 1: continue
        split[label] = (shuffled[:n_train], shuffled[n_train:])
        log.info(f"  {code(label):8s}  total={n:4d}  train={n_train:4d}  test={n_test:4d}")

    if skipped:
        log.info(f"  Skipped (< {min_seqs} seqs): " + ", ".join(f"{code(l)}({n})" for l, n in skipped))
    log.info(f"Holdout: 1 split × {len(split)} categories ({sum(len(tr)+len(te) for tr,te in split.values()):,} sequences)")
    return [split]


# ═══════════════════════════════════════════════════════════════════════════════
# 4 ── Embedding model  (unified loader)
# ═══════════════════════════════════════════════════════════════════════════════

def load_embedding_model(model_name: str, device: Optional[str] = None) -> tuple[Callable, str, int]:
    """
    Load by alias or HF model ID.
    Returns (embed_fn, model_id, emb_dim).
    embed_fn: (frames: list[np.ndarray], batch_size: int) -> list[list[float]]
    """
    if device is None: device = "cuda" if torch.cuda.is_available() else "cpu"

    if model_name in OPENCLIP_ALIASES:
        if not OPENCLIP_OK: print("ERROR: pip install open-clip-torch"); sys.exit(1)
        arch, pretrained = OPENCLIP_ALIASES[model_name]
        log.info(f"Loading OpenCLIP {arch}/{pretrained} on {device} …")
        oc_model, _, preprocess = open_clip.create_model_and_transforms(arch, pretrained=pretrained)
        oc_model = oc_model.eval().to(device)
        model_id = f"openclip/{arch}/{pretrained}"
        def embed_fn(frames, batch_size=16):
            all_embs = []
            for i in range(0, len(frames), batch_size):
                batch = frames[i:i+batch_size]
                tensors = torch.stack([preprocess(PILImage.fromarray(f)) for f in batch]).to(device)
                with torch.no_grad():
                    feats = oc_model.encode_image(tensors)
                    feats = feats / feats.norm(dim=-1, keepdim=True)
                all_embs.extend(feats.cpu().float().numpy().tolist())
            return all_embs
    else:
        if not HF_OK: print("ERROR: pip install transformers torch pillow"); sys.exit(1)
        hf_id = HF_ALIASES.get(model_name, model_name); model_id = hf_id
        log.info(f"Loading HuggingFace '{hf_id}' on {device} …")
        processor = AutoImageProcessor.from_pretrained(hf_id)
        hf_model  = AutoModel.from_pretrained(hf_id).eval().to(device)
        def embed_fn(frames, batch_size=16):
            all_embs = []
            for i in range(0, len(frames), batch_size):
                batch  = frames[i:i+batch_size]
                pil    = [PILImage.fromarray(f) for f in batch]
                inputs = processor(images=pil, return_tensors="pt")
                inputs = {k: v.to(device) for k, v in inputs.items()}
                with torch.no_grad():
                    out = hf_model(**inputs)
                    cls = out.last_hidden_state[:, 0, :]
                    cls = cls / cls.norm(dim=-1, keepdim=True)
                all_embs.extend(cls.cpu().float().numpy().tolist())
            return all_embs

    test_emb = embed_fn([np.zeros((224, 224, 3), dtype=np.uint8)], batch_size=1)
    emb_dim  = len(test_emb[0])
    log.info(f"Model ready — id={model_id}  dim={emb_dim}")
    return embed_fn, model_id, emb_dim


# ═══════════════════════════════════════════════════════════════════════════════
# 4b ── Temporal sequence embedding
# ═══════════════════════════════════════════════════════════════════════════════

def compute_temporal_embedding(frame_embeddings: list[list[float]]) -> list[float]:
    """
    Aggregate N per-frame embeddings into ONE temporal sequence descriptor.

    mean pool → average spatial appearance over the run  (anatomy)
    std  pool → frame-to-frame variation                 (contrast dynamics)

    Concatenated and L2-normalised → 2×D vector.
    N=1: std=0, result = [frame_emb, 0…0] normalised (spatial only).
    """
    embs      = np.array(frame_embeddings, dtype=np.float32)
    mean_pool = embs.mean(axis=0)
    std_pool  = embs.std(axis=0)
    descriptor = np.concatenate([mean_pool, std_pool])
    norm = np.linalg.norm(descriptor)
    if norm > 0: descriptor /= norm
    return descriptor.tolist()


# ═══════════════════════════════════════════════════════════════════════════════
# 5 ── DICOM frame extraction  (mode-aware)
# ═══════════════════════════════════════════════════════════════════════════════

def _to_uint8_rgb(frame: np.ndarray, photometric: str) -> np.ndarray:
    f = frame.astype(np.float32); lo, hi = f.min(), f.max()
    f = ((f - lo) / (hi - lo) * 255.0 if hi > lo else np.zeros_like(f)).astype(np.uint8)
    if "MONOCHROME1" in photometric.upper(): f = 255 - f
    return np.stack([f, f, f], axis=-1)

def _load_pixels(path_str: str) -> tuple[Optional[np.ndarray], str]:
    try:
        ds = pydicom.dcmread(path_str, force=False)
        photometric = str(getattr(ds, "PhotometricInterpretation", "")).upper()
        pixels = ds.pixel_array
        if pixels.ndim == 2: pixels = pixels[np.newaxis]
        elif pixels.ndim == 3 and pixels.shape[2] in (3, 4): pixels = pixels[np.newaxis]
        return pixels, photometric
    except Exception as e: log.debug(f"DICOM read failed [{path_str}]: {e}"); return None, ""

def _raw_to_rgb(raw: np.ndarray, photometric: str) -> np.ndarray:
    if raw.ndim == 3:
        f = raw.astype(np.float32); lo, hi = f.min(), f.max()
        f = ((f-lo)/(hi-lo)*255.0 if hi>lo else np.zeros_like(f)).astype(np.uint8)
        return f[:,:,:3]
    return _to_uint8_rgb(raw, photometric)

def _extract_best_frame(path_str: str, frame_idx: int) -> list[np.ndarray]:
    pixels, photometric = _load_pixels(path_str)
    if pixels is None: return []
    total = pixels.shape[0]
    idx = total // 2 if frame_idx < 0 or frame_idx >= total else frame_idx
    return [_raw_to_rgb(pixels[idx], photometric)]

def _extract_frame_range(path_str: str, first_idx: int, last_idx: int) -> list[np.ndarray]:
    pixels, photometric = _load_pixels(path_str)
    if pixels is None: return []
    total = pixels.shape[0]
    lo = max(0, min(first_idx, total-1)); hi = max(0, min(last_idx, total-1))
    if hi < lo: hi = lo
    return [_raw_to_rgb(pixels[i], photometric) for i in range(lo, hi+1)]

def _extract_all_frames(path_str: str, max_frames: int = 0) -> list[np.ndarray]:
    pixels, photometric = _load_pixels(path_str)
    if pixels is None: return []
    total = pixels.shape[0]
    if max_frames > 0 and total > max_frames:
        indices = [total//2] if max_frames==1 else sorted(set(round(i*(total-1)/(max_frames-1)) for i in range(max_frames)))
    else: indices = list(range(total))
    return [_raw_to_rgb(pixels[i], photometric) for i in indices]

def extract_frames_for_seq(seq: dict, mode: str, max_frames: int = 0) -> list[np.ndarray]:
    if mode == "best": return _extract_best_frame(seq["dicom_path"], seq.get("best_image_idx", -1))
    elif mode == "fl": return _extract_frame_range(seq["dicom_path"], seq["first_diag_idx"], seq["last_diag_idx"])
    else: return _extract_all_frames(seq["dicom_path"], max_frames)

def _get_display_frame_png(seq: dict, size: int = 256) -> Optional[bytes]:
    """Load one representative frame as PNG bytes for docx thumbnails."""
    pixels, photometric = _load_pixels(seq["dicom_path"])
    if pixels is None: return None
    total = pixels.shape[0]; best_idx = seq.get("best_image_idx", -1)
    idx = best_idx if 0 <= best_idx < total else total // 2
    frame = _raw_to_rgb(pixels[idx], photometric)
    img = PILImage.fromarray(frame).resize((size, size), PILImage.LANCZOS)
    buf = io.BytesIO(); img.save(buf, format="PNG"); return buf.getvalue()


# ═══════════════════════════════════════════════════════════════════════════════
# 6 ── Threading helpers
# ═══════════════════════════════════════════════════════════════════════════════

def _chunked(iterable, size: int):
    it = iter(iterable)
    while chunk := list(itertools.islice(it, size)): yield chunk

def make_frame_reader(mode: str, max_frames: int = 0):
    def reader(seq: dict) -> tuple[dict, list]:
        return seq, extract_frames_for_seq(seq, mode, max_frames)
    return reader


# ═══════════════════════════════════════════════════════════════════════════════
# 7 ── ChromaDB  (in-memory, fresh per fold)
# ═══════════════════════════════════════════════════════════════════════════════

def create_ephemeral_collection():
    try: client = chromadb.EphemeralClient()
    except AttributeError: client = chromadb.Client()
    # Newer ChromaDB versions make EphemeralClient() a singleton — the same
    # in-memory store persists across calls in the same process.
    # Delete the collection first so each fold always starts with an empty one.
    try: client.delete_collection(CHROMA_COLLECTION)
    except (ValueError, Exception) as exc:
        log.debug(f"Collection deletion skipped (expected on first call): {exc}")
    collection = client.create_collection(name=CHROMA_COLLECTION, metadata={"hnsw:space": "cosine"})
    log.info("In-memory ChromaDB collection created (no disk I/O)")
    return collection


# ═══════════════════════════════════════════════════════════════════════════════
# 7b ── Precompute ALL embeddings once  (CV 10× speedup)
# ═══════════════════════════════════════════════════════════════════════════════

def precompute_all_embeddings(
    groups: dict[str, list[dict]], embed_fn: Callable,
    frame_mode: str, max_frames: int, embed_batch: int,
    workers: int, temporal: bool,
) -> dict[str, Optional[list]]:
    """
    Embed every sequence exactly ONCE before the fold loop.
    Returns { stem: embeddings } where embeddings is:
      temporal=False → list of N per-frame vectors
      temporal=True  → single-element list [temporal_descriptor]
      None           → unreadable
    """
    all_seqs = [seq for seqs in groups.values() for seq in seqs]
    prefetch_chunk = max(workers * 4, 32)
    reader = make_frame_reader(frame_mode, max_frames)
    mode_desc = f"mode={frame_mode}" + (" + temporal" if temporal else "")
    log.info(f"Pre-computing embeddings for {len(all_seqs):,} sequences ({mode_desc}, workers={workers}) …")

    all_embs: dict[str, Optional[list]] = {}
    skipped = 0
    with tqdm(total=len(all_seqs), unit="seq", desc="  Embedding ALL") as pbar:
        with ThreadPoolExecutor(max_workers=workers) as pool:
            for chunk in _chunked(all_seqs, prefetch_chunk):
                for seq, frames in pool.map(reader, chunk):
                    stem = seq["stem"]
                    if not frames: all_embs[stem] = None; skipped += 1
                    else:
                        embs = embed_fn(frames, embed_batch)
                        all_embs[stem] = [compute_temporal_embedding(embs)] if temporal else embs
                    pbar.update(1)
    log.info(f"Embeddings ready — {len(all_seqs)-skipped:,} ok, {skipped} skipped")
    return all_embs


def ingest_fold_from_precomputed(
    fold_splits: dict[str, tuple[list, list]],
    collection,
    all_embs:   dict[str, Optional[list]],
    frame_mode: str,
    temporal:   bool,
) -> int:
    """Add precomputed TRAIN embeddings for this fold into ChromaDB. No GPU work."""
    train_seqs = [seq for _, (train, _) in fold_splits.items() for seq in train]
    total_entries = 0; skipped = 0
    for seq in train_seqs:
        stem = seq["stem"]; embs = all_embs.get(stem)
        if not embs: skipped += 1; continue
        meta_base = {"sequence_id": stem, "angio_run": seq["angio_run"],
                     "accession": seq["accession"], "run_type": seq["run_type"]}
        if temporal:
            collection.add(embeddings=embs, ids=[stem], metadatas=[{**meta_base, "n_frames": 1}])
            total_entries += 1
        else:
            dicom_start = (max(0, seq["first_diag_idx"]) if frame_mode == "fl" and seq["first_diag_idx"] >= 0 else 0)
            collection.add(
                embeddings=embs,
                ids=[f"{stem}_f{dicom_start+i:05d}" for i in range(len(embs))],
                metadatas=[{**meta_base, "frame_idx": dicom_start+i} for i in range(len(embs))],
            )
            total_entries += len(embs)
    return total_entries


# ═══════════════════════════════════════════════════════════════════════════════
# 8 ── K@N evaluation
# ═══════════════════════════════════════════════════════════════════════════════

def evaluate_at_k(
    embedded_test: list[tuple[dict, Optional[list]]],
    splits: dict[str, tuple[list, list]],
    collection, k: int,
) -> dict[str, dict]:
    effective_k = min(k, collection.count())
    if effective_k < k: log.warning(f"  K={k} capped to {effective_k}")

    results: dict[str, dict] = {
        label: {"n_train": len(train), "n_test": len(test),
                "correct": 0, "total_evaled": 0, "skipped": 0, "predictions": []}
        for label, (train, test) in splits.items()
    }
    for seq, embs in embedded_test:
        label = seq["angio_run"]
        if label not in results: continue
        r = results[label]
        if embs is None: r["skipped"] += 1; continue

        qr = collection.query(query_embeddings=embs, n_results=effective_k, include=["metadatas"])
        seq_votes: Counter = Counter(); seq_labels: dict[str, str] = {}
        for meta_list in qr["metadatas"]:
            for m in meta_list:
                sid = m.get("sequence_id", ""); lbl = m.get("angio_run", "")
                seq_votes[sid] += 1; seq_labels[sid] = lbl
        if not seq_votes: r["skipped"] += 1; continue

        pred = seq_labels[seq_votes.most_common(1)[0][0]]
        r["total_evaled"] += 1; r["correct"] += int(pred == label)
        r["predictions"].append((label, pred))
    return results


def run_all_k_evaluations(
    embedded_test: list, splits: dict, collection, k_values: list[int],
) -> dict[int, dict[str, dict]]:
    all_results: dict[int, dict] = {}
    for k in k_values:
        log.info(f"─── Evaluating K={k} ───")
        all_results[k] = evaluate_at_k(embedded_test, splits, collection, k)
    return all_results


# ═══════════════════════════════════════════════════════════════════════════════
# 9 ── CV aggregation
# ═══════════════════════════════════════════════════════════════════════════════

def aggregate_cv_results(
    fold_results: list[dict[int, dict[str, dict]]],
    k_values: list[int], n_folds: int,
) -> dict[int, dict[str, dict]]:
    """
    Pool per-fold results:
      correct / total_evaled — summed (pooled micro accuracy)
      fold_accs              — per-fold accuracy list for mean±std
      n_train / n_test       — averaged per fold
    """
    all_labels: set[str] = set()
    for fr in fold_results:
        for k in k_values:
            if k in fr: all_labels.update(fr[k].keys())

    aggregated: dict[int, dict[str, dict]] = {}
    for k in k_values:
        aggregated[k] = {}
        for label in sorted(all_labels):
            total_correct = total_evaled = total_skip = n_train_sum = n_test_sum = 0
            fold_accs: list[float] = []
            for fr in fold_results:
                if k not in fr or label not in fr[k]: continue
                r = fr[k][label]
                total_correct += r["correct"]; total_evaled += r["total_evaled"]
                total_skip    += r["skipped"]; n_train_sum  += r["n_train"]
                n_test_sum    += r["n_test"]
                if r["total_evaled"] > 0: fold_accs.append(r["correct"] / r["total_evaled"])
            aggregated[k][label] = {
                "correct": total_correct, "total_evaled": total_evaled,
                "skipped": total_skip,
                "n_train": round(n_train_sum / n_folds),
                "n_test":  round(n_test_sum  / n_folds),
                "fold_accs": fold_accs, "predictions": [],
            }
    return aggregated


# ═══════════════════════════════════════════════════════════════════════════════
# 9b ── K=1 retrieval example collector  (for docx report)
# ═══════════════════════════════════════════════════════════════════════════════

def collect_k1_retrieval_examples(
    embedded_test: list[tuple[dict, Optional[list]]],
    splits: dict[str, tuple[list, list]],
    collection, n_per_cat: int = 5,
) -> dict[str, list[dict]]:
    """
    Run a K=1 pass and collect up to n_per_cat examples per category.
    Interleaves hits and misses so both appear at the top of each docx section.
    """
    train_index: dict[str, dict] = {
        seq["stem"]: seq for _, (train, _) in splits.items() for seq in train
    }
    by_cat: dict[str, list] = defaultdict(list)
    for seq, embs in embedded_test:
        if embs is not None: by_cat[seq["angio_run"]].append((seq, embs))

    log.info(f"Collecting K=1 retrieval examples (up to {n_per_cat}/category) …")
    examples: dict[str, list[dict]] = {}

    for label, cat_seqs in sorted(by_cat.items()):
        hits: list[dict] = []; misses: list[dict] = []
        for seq, embs in cat_seqs:
            if len(hits) >= n_per_cat and len(misses) >= n_per_cat: break
            qr = collection.query(query_embeddings=embs, n_results=1, include=["metadatas"])
            seq_votes: Counter = Counter(); seq_labels: dict[str, str] = {}
            for meta_list in qr["metadatas"]:
                for m in meta_list:
                    sid = m.get("sequence_id",""); lbl = m.get("angio_run","")
                    seq_votes[sid] += 1; seq_labels[sid] = lbl
            if not seq_votes: continue
            best_sid = seq_votes.most_common(1)[0][0]
            retr_label = seq_labels[best_sid]; hit = (retr_label == label)
            entry = {"test_seq": seq, "retr_stem": best_sid, "retr_label": retr_label,
                     "hit": hit, "train_seq": train_index.get(best_sid)}
            if hit and len(hits) < n_per_cat: hits.append(entry)
            elif not hit and len(misses) < n_per_cat: misses.append(entry)

        collected: list[dict] = []; hi = mi = 0
        while len(collected) < n_per_cat:
            added = False
            if hi < len(hits): collected.append(hits[hi]); hi += 1; added = True
            if len(collected) < n_per_cat and mi < len(misses):
                collected.append(misses[mi]); mi += 1; added = True
            if not added: break
        if collected:
            examples[label] = collected
            n_h = sum(1 for e in collected if e["hit"])
            log.info(f"  {code(label):8s}  {len(collected)} examples ({n_h} hits, {len(collected)-n_h} misses)")

    return examples


# ═══════════════════════════════════════════════════════════════════════════════
# 10 ── Results reporting
# ═══════════════════════════════════════════════════════════════════════════════

def _micro(rk: dict[str, dict]) -> tuple[int, int]:
    return (sum(r["correct"] for r in rk.values()),
            sum(r["total_evaled"] for r in rk.values()))

def _macro(rk: dict[str, dict]) -> float:
    accs = [r["correct"]/r["total_evaled"] for r in rk.values() if r["total_evaled"] > 0]
    return sum(accs)/len(accs) if accs else 0.0

def _fmt_cv(r: dict) -> str:
    """Format mean±std from fold_accs; fall back to pooled if unavailable."""
    accs = r.get("fold_accs", [])
    if len(accs) >= 2:
        mean = sum(accs)/len(accs)
        std  = (sum((a-mean)**2 for a in accs)/len(accs))**0.5
        return f"{mean:.1%}±{std:.1%}"
    if len(accs) == 1: return f"{accs[0]:.1%}"
    te = r.get("total_evaled", 0)
    return f"{r['correct']/te:.1%}" if te > 0 else "N/A"


def print_results_table(
    all_results: dict[int, dict[str, dict]],
    k_values: list[int], n_folds: int, model_id: str, frame_mode: str,
) -> None:
    r1 = all_results[k_values[0]]
    rows = sorted(
        [(lbl, code(lbl), r,
          sum(r.get("fold_accs",[]))/len(r["fold_accs"]) if r.get("fold_accs")
          else (r["correct"]/r["total_evaled"] if r["total_evaled"]>0 else None))
         for lbl, r in r1.items()],
        key=lambda x: (x[3] is None, -(x[3] or 0)),
    )
    W   = max(len(lbl) for lbl, *_ in rows) + 2
    hdr = (f"  {'Category':<{W}}  {'Code':<8}  {'~Train':>7}  {'~Test':>6}  "
           f"{'Evaled':>7}  {'K=1 Acc (mean±std)':>20}")
    sep = "─" * len(hdr)
    print()
    print(f"  AngioVision  {n_folds}-fold CV  |  model={model_id}  |  mode={frame_mode}")
    print(sep); print(hdr); print(sep)
    for lbl, cd, r, _ in rows:
        print(f"  {lbl:<{W}}  {cd:<8}  {r['n_train']:>7,}  {r['n_test']:>6,}  "
              f"{r['total_evaled']:>7,}  {_fmt_cv(r):>20}")
    print(sep)
    print()
    print(f"  Cross-validated micro & macro accuracy vs K  ({n_folds} folds):")
    print("  " + "  ".join(f"K={k:>2}" for k in k_values))
    def _fmt(val): return f"{val:>5.1%}"
    micro_r = "  "+"  ".join(_fmt(_micro(all_results[k])[0]/_micro(all_results[k])[1]) if _micro(all_results[k])[1]>0 else "  N/A" for k in k_values)
    macro_r = "  "+"  ".join(_fmt(_macro(all_results[k])) for k in k_values)
    print(f"  micro: {micro_r}"); print(f"  macro: {macro_r}"); print()


def save_bar_chart(
    all_results: dict[int, dict[str, dict]],
    k_values: list[int], n_folds: int, model_id: str, frame_mode: str, out_path: Path,
) -> None:
    if not MATPLOTLIB_OK: log.warning("matplotlib unavailable — chart skipped"); return

    r1 = all_results[k_values[0]]
    rows = []
    for lbl, r in r1.items():
        if not r["total_evaled"]: continue
        accs = r.get("fold_accs", [])
        mean = sum(accs)/len(accs) if accs else r["correct"]/r["total_evaled"]
        std  = (sum((a-mean)**2 for a in accs)/len(accs))**0.5 if len(accs)>1 else 0.0
        rows.append((code(lbl), lbl, mean, std, r["n_train"], r["n_test"]))
    rows.sort(key=lambda x: -x[2])
    if not rows: log.warning("No evaluated categories — chart skipped"); return

    colours = ["#2ecc71" if a>=0.80 else "#f39c12" if a>=0.50 else "#e74c3c" for _,_,a,*_ in rows]
    stds    = [s for _,_,_,s,*_ in rows]

    fig = plt.figure(figsize=(18, max(6, len(rows)*0.42+3)))
    ax_bar  = fig.add_axes([0.04, 0.10, 0.44, 0.82])
    ax_line = fig.add_axes([0.56, 0.10, 0.40, 0.82])

    y_pos = range(len(rows))
    ax_bar.barh(y_pos, [a for _,_,a,*_ in rows],
                xerr=stds, error_kw={"ecolor":"#555","capsize":3,"linewidth":1.2},
                color=colours, edgecolor="black", linewidth=0.5, height=0.7)
    for i, (cd,lbl,acc,std,nt,nts) in enumerate(rows):
        ax_bar.text(min(acc+std+0.015, 0.99), i, f"{acc:.1%}±{std:.1%}",
                    va="center", fontsize=7, fontweight="bold")
    ax_bar.set_yticks(list(y_pos)); ax_bar.set_yticklabels([r[0] for r in rows], fontsize=8.5)
    ax_bar.set_xlim(0, 1.22)
    ax_bar.axvline(x=0.5,  color="gray",  linewidth=1, linestyle="--", alpha=0.5)
    ax_bar.axvline(x=0.80, color="green", linewidth=1, linestyle="--", alpha=0.5)
    ax_bar.set_xlabel(f"K=1 Accuracy  (mean±std,  {n_folds} folds)", fontsize=9)
    ax_bar.set_title(f"Per-Category K=1 Accuracy  ({n_folds}-fold CV)", fontsize=10, fontweight="bold")
    ax_bar.invert_yaxis(); ax_bar.grid(axis="x", linestyle="--", alpha=0.4)
    ax_bar.legend(handles=[mpatches.Patch(color="#2ecc71",label="≥ 80%"),
                            mpatches.Patch(color="#f39c12",label="50–79%"),
                            mpatches.Patch(color="#e74c3c",label="< 50%")], fontsize=8, loc="lower right")

    micro_vals = [_micro(all_results[k])[0]/_micro(all_results[k])[1] if _micro(all_results[k])[1]>0 else 0.0 for k in k_values]
    macro_vals = [_macro(all_results[k]) for k in k_values]
    ax_line.plot(k_values, micro_vals, "o-",  color="#2980b9", linewidth=2, markersize=7, label="Micro (pooled)")
    ax_line.plot(k_values, macro_vals, "s--", color="#e67e22", linewidth=2, markersize=7, label="Macro (pooled)")
    for k, mv, av in zip(k_values, micro_vals, macro_vals):
        ax_line.annotate(f"{mv:.1%}", (k,mv), textcoords="offset points", xytext=(0,8),  ha="center", fontsize=7.5, color="#2980b9")
        ax_line.annotate(f"{av:.1%}", (k,av), textcoords="offset points", xytext=(0,-14),ha="center", fontsize=7.5, color="#e67e22")
    ax_line.set_xticks(k_values); ax_line.set_xlabel("K", fontsize=9)
    ax_line.set_ylabel("Accuracy (pooled across folds)", fontsize=9)
    ax_line.set_title(f"Micro & Macro vs K  ({n_folds}-fold CV)", fontsize=10, fontweight="bold")
    ax_line.set_ylim(0, 1.05); ax_line.axhline(y=0.80, color="green", linewidth=1, linestyle="--", alpha=0.4)
    ax_line.grid(linestyle="--", alpha=0.4); ax_line.legend(fontsize=9)
    plt.suptitle(f"AngioVision K@N  |  {n_folds}-fold CV  |  model: {model_id}  |  frames: {frame_mode}",
                 fontsize=10, fontweight="bold", y=1.01)
    plt.savefig(str(out_path), dpi=150, bbox_inches="tight"); plt.close()
    log.info(f"Chart saved → {out_path}")


def write_markdown(
    all_results: dict[int, dict[str, dict]],
    k_values: list[int], n_folds: int, seed: int,
    model_id: str, frame_mode: str, emb_dim: int, temporal: bool, out_path: Path,
) -> None:
    r1   = all_results[k_values[0]]
    cats = sorted(r1.keys(), key=lambda l: -(r1[l]["correct"]/r1[l]["total_evaled"] if r1[l]["total_evaled"]>0 else 0))
    now      = datetime.datetime.now().strftime("%Y-%m-%d %H:%M")
    khdrs    = " | ".join(f"K={k}" for k in k_values)
    emb_note = f"{emb_dim}-dim" + (f" → 2×{emb_dim}-dim temporal" if temporal else "")

    lines: list[str] = [
        "# AngioVision K@N Retrieval Evaluation  (Cross-Validated)\n\n",
        f"**Generated:** {now}  \n",
        f"**Model:** `{model_id}` ({emb_note}, L2-normalised)  \n",
        f"**Frame mode:** `{frame_mode}`  \n",
        f"**Temporal aggregation:** {'✓ mean+std pooling (one vector per sequence)' if temporal else '✗ per-frame (majority vote)'}  \n",
        f"**Validation:** {n_folds}-fold cross-validation  |  seed={seed}  \n",
        f"**Matching:** {'sequence-level cosine similarity' if temporal else 'majority vote across K × N_frames results'}  \n\n",
        f"> Cells show **mean±std** across {n_folds} folds.  Micro/Macro rows show pooled counts.\n\n",
        "## Per-Category K@N Accuracy\n\n",
        f"| Category | Code | ~Train/fold | ~Test/fold | Total Evaled | {khdrs} |\n",
        "|:---------|:----:|------------:|-----------:|-------------:"+("".join("-----------:|" for _ in k_values))+"\n",
    ]
    for label in cats:
        r = r1[label]
        cells = [_fmt_cv(all_results[k][label]) for k in k_values]
        lines.append(f"| {label} | {code(label)} | {r['n_train']:,} | {r['n_test']:,} | {r['total_evaled']:,} | "+" | ".join(cells)+" |\n")

    tc0, te0 = _micro(r1); base_micro = tc0/te0 if te0>0 else 0.0
    total_train  = sum(r1[l]["n_train"]      for l in cats)
    total_test   = sum(r1[l]["n_test"]       for l in cats)
    total_evaled = sum(r1[l]["total_evaled"] for l in cats)
    micro_cells = [f"**{_micro(all_results[k])[0]/_micro(all_results[k])[1]:.1%}**" if _micro(all_results[k])[1]>0 else "N/A" for k in k_values]
    macro_cells = [f"**{_macro(all_results[k]):.1%}**" for k in k_values]

    lines += [
        f"| **MICRO (pooled)** | | **{total_train:,}** | **{total_test:,}** | **{total_evaled:,}** | "+" | ".join(micro_cells)+" |\n",
        "| **MACRO (pooled)** | | | | | "+" | ".join(macro_cells)+" |\n\n",
        "## K-Sweep Summary\n\n",
        "| K | Micro (pooled) | Macro (pooled) | Δ vs K=1 |\n",
        "|--:|--------------:|--------------:|---------:|\n",
    ]
    for k in k_values:
        tc, te = _micro(all_results[k]); micro = tc/te if te>0 else 0.0; delta = micro-base_micro
        lines.append(f"| {k} | {micro:.1%} | {_macro(all_results[k]):.1%} | {'+' if delta>=0 else ''}{delta:.1%} |\n")

    lines += ["\n## Best K Per Category\n\n",
              "| Category | Code | ~Train/fold | ~Test/fold | Best K | Best Acc (mean) |\n",
              "|:---------|:----:|------------:|-----------:|-------:|----------------:|\n"]
    for label in cats:
        best_k, best_mean = None, -1.0
        for k in k_values:
            rk = all_results[k][label]; accs = rk.get("fold_accs",[])
            m  = sum(accs)/len(accs) if accs else (rk["correct"]/rk["total_evaled"] if rk["total_evaled"]>0 else 0.0)
            if m > best_mean: best_mean = m; best_k = k
        if best_k is not None:
            lines.append(f"| {label} | {code(label)} | {r1[label]['n_train']:,} | {r1[label]['n_test']:,} | K={best_k} | {best_mean:.1%} |\n")

    out_path.write_text("".join(lines), encoding="utf-8")
    log.info(f"Markdown saved → {out_path}")


# ═══════════════════════════════════════════════════════════════════════════════
# 11 ── Retrieval examples docx
# ═══════════════════════════════════════════════════════════════════════════════

def save_retrieval_docx(
    examples: dict[str, list[dict]], k1_results: dict[str, dict],
    out_path: Path, model_id: str, frame_mode: str, temporal: bool,
) -> None:
    """
    Word document with K=1 retrieval examples per category.
    Requires: pip install python-docx
    """
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        log.error("python-docx not installed — run: pip install python-docx"); return

    def _shade(cell, hex_fill):
        tc = cell._tc; pr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd"); shd.set(qn("w:fill"), hex_fill); shd.set(qn("w:val"), "clear"); pr.append(shd)

    def _cell_para(cell, text, bold=False, size_pt=10, color=None):
        p = cell.paragraphs[0]; p.clear(); run = p.add_run(text)
        run.bold = bold; run.font.size = Pt(size_pt)
        if color: run.font.color.rgb = color

    def _add_img(cell, png_bytes):
        p = cell.paragraphs[0]; p.clear(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if png_bytes: p.add_run().add_picture(io.BytesIO(png_bytes), width=Inches(2.3))
        else: p.add_run("[unavailable]").font.size = Pt(8)

    COL_W = [576, 3888, 3888, 1728]
    HIT_GREEN = "D5F5E3"; MISS_RED = "FADBD8"; HEADER_BG = "2C3E50"

    doc = Document()
    for sec in doc.sections:
        sec.left_margin = sec.right_margin = sec.top_margin = sec.bottom_margin = Inches(0.75)

    title = doc.add_heading("AngioVision  K@1  Retrieval Examples", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta = doc.add_paragraph(); meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"Model: {model_id}   |   Frame mode: {frame_mode}   |   "
                 f"Temporal: {'ON' if temporal else 'OFF'}   |   "
                 f"Generated: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M')}").font.size = Pt(9)
    doc.add_paragraph()

    for label in sorted(examples.keys()):
        cases = examples[label]; r1 = k1_results.get(label, {})
        n_ev = r1.get("total_evaled",0); n_cor = r1.get("correct",0)
        acc  = n_cor/n_ev if n_ev>0 else 0.0
        h = doc.add_heading(level=2); h.clear()
        h.add_run(f"{label}  ({code(label)})  —  K@1: {n_cor}/{n_ev} ({acc:.0%})  |  Train≈{r1.get('n_train','?')}  Test≈{r1.get('n_test','?')}")

        table = doc.add_table(rows=1, cols=4); table.style = "Table Grid"
        hdr = table.rows[0].cells
        for i, (cell, txt) in enumerate(zip(hdr, ["#","Query Frame","Retrieved Frame","Outcome"])):
            _cell_para(cell, txt, bold=True, size_pt=10, color=RGBColor(0xFF,0xFF,0xFF))
            _shade(cell, HEADER_BG)
            tc = cell._tc; pr = tc.get_or_add_tcPr(); w = OxmlElement("w:tcW")
            w.set(qn("w:w"), str(COL_W[i])); w.set(qn("w:type"), "dxa"); pr.append(w)

        for idx, case in enumerate(cases):
            row = table.add_row().cells; hit = case["hit"]; bg = HIT_GREEN if hit else MISS_RED
            _cell_para(row[0], str(idx+1), bold=True, size_pt=9); _shade(row[0], bg)
            q_png = _get_display_frame_png(case["test_seq"])
            _add_img(row[1], q_png); row[1].add_paragraph().add_run(f"True: {code(label)}").font.size = Pt(8)
            r_png = _get_display_frame_png(case["train_seq"]) if case.get("train_seq") else None
            _add_img(row[2], r_png); row[2].add_paragraph().add_run(f"Retrieved: {code(case['retr_label'])}").font.size = Pt(8)
            _cell_para(row[3], f"{'✓  HIT' if hit else '✗  MISS'}\n\nQuery:     {code(label)}\nRetrieved: {code(case['retr_label'])}", bold=hit, size_pt=9)
            _shade(row[3], bg)
            for i, cell in enumerate(row):
                tc = cell._tc; pr = tc.get_or_add_tcPr(); w = OxmlElement("w:tcW")
                w.set(qn("w:w"), str(COL_W[i])); w.set(qn("w:type"), "dxa"); pr.append(w)
        doc.add_paragraph()

    doc.save(str(out_path)); log.info(f"Retrieval examples docx saved → {out_path}")


# ═══════════════════════════════════════════════════════════════════════════════
# 12 ── Entry point
# ═══════════════════════════════════════════════════════════════════════════════

def main() -> None:
    parser = argparse.ArgumentParser(
        description="K@N cross-validated retrieval evaluation — AngioVision DSA",
        formatter_class=argparse.RawDescriptionHelpFormatter,
    )
    parser.add_argument("--labeled-csv", default=DEFAULT_LABELED_CSV)
    parser.add_argument("--dicom-root",  default=DEFAULT_DICOM_ROOT)
    parser.add_argument("--sqlite-db",   default=DEFAULT_SQLITE_DB)
    parser.add_argument("--out-plot",    default=None, help="Bar chart PNG (auto-named if omitted)")
    parser.add_argument("--out-md",      default=None, help="Markdown results (auto-named if omitted)")
    parser.add_argument("--out-docx",    default=None, help="Word doc examples (auto-named if omitted)")
    parser.add_argument("--model",       default="rad-dino",
        help="rad-dino | vit-b16 | vit-l16 | openclip-b32 | openclip-l14 | <HF model ID>  [default: rad-dino]")
    parser.add_argument("--frame-mode",  default="fl", choices=FRAME_MODES,
        help="best=Best_Image | fl=First→Last diag window | all=all frames  [default: fl]")
    parser.add_argument("--max-frames",  type=int, default=0,
        help="Max frames per seq when --frame-mode=all  (0=all)  [default: 0]")
    parser.add_argument("--k-values",    type=int, nargs="+", default=DEFAULT_K_VALUES,
        help="K values to evaluate  [default: 1 3 5 7 9 11 13 15]")
    parser.add_argument("--split-mode",  default="cv", choices=("cv", "holdout"),
        help="cv=n-fold cross-validation | holdout=single train/test split  [default: cv]")
    parser.add_argument("--n-folds",     type=int, default=10,
        help="Number of CV folds (split-mode=cv)  [default: 10]")
    parser.add_argument("--scale-down",  type=float, default=0.80,
        help="TRAIN fraction per category (split-mode=holdout); e.g. 0.80 = 80%% train / 20%% test  [default: 0.80]")
    parser.add_argument("--embed-batch", type=int, default=16,
        help="Frames per model forward pass  [default: 16]")
    parser.add_argument("--temporal",    action="store_true",
        help="Mean+std temporal aggregation (2×D per sequence)  [default: off]")
    parser.add_argument("--min-seqs",    type=int, default=3)
    parser.add_argument("--seed",        type=int, default=42)
    parser.add_argument("--device",      default=None)
    parser.add_argument("--workers",     type=int, default=DEFAULT_WORKERS,
        help=f"DICOM reader threads  [default: {DEFAULT_WORKERS}]")
    parser.add_argument("--limit-cats",  type=int, default=0)

    args     = parser.parse_args()
    k_values = sorted(set(args.k_values))

    # Split-mode → effective "fold" count + self-documenting filename tag
    if args.split_mode == "holdout":
        n_folds_eff = 1
        split_tag   = f"holdout{round(args.scale_down*100)}"
        split_desc  = f"holdout ({args.scale_down:.0%} train / {1-args.scale_down:.0%} test)"
    else:
        n_folds_eff = args.n_folds
        split_tag   = f"{args.n_folds}fold"
        split_desc  = f"{args.n_folds}-fold cross-validation"

    _auto_md, _auto_png, _auto_docx = _auto_output_paths(
        args.model, args.frame_mode, args.temporal, split_tag,
    )
    out_md   = Path(args.out_md   or _auto_md)
    out_plot = Path(args.out_plot or _auto_png)
    out_docx = Path(args.out_docx or _auto_docx)

    log.info("═" * 60)
    log.info("  AngioVision K@N Retrieval Evaluation")
    log.info("═" * 60)
    log.info(f"  model         : {args.model}")
    log.info(f"  frame-mode    : {args.frame_mode}" +
             (f"  (max-frames={args.max_frames})" if args.frame_mode=="all" and args.max_frames else ""))
    log.info(f"  temporal      : {'ON  (mean+std pooling)' if args.temporal else 'OFF'}")
    log.info(f"  split-mode    : {split_desc}")
    log.info(f"  K values      : {k_values}")
    log.info(f"  workers       : {args.workers}  |  embed-batch: {args.embed_batch}")
    log.info(f"  min-seqs      : {args.min_seqs}  |  seed: {args.seed}")
    log.info(f"  out-md        : {out_md}")
    log.info(f"  out-plot      : {out_plot}")
    log.info(f"  out-docx      : {out_docx}")
    log.info("═" * 60)

    # 1. CSV
    groups = load_labeled_csv_grouped(Path(args.labeled_csv), args.frame_mode)

    # 2. DICOM index
    dicom_index = build_dicom_index(Path(args.dicom_root), Path(args.sqlite_db))

    # 3. Resolve paths
    groups = resolve_paths(groups, dicom_index)
    if not groups: log.error("No sequences with resolvable DICOM paths — aborting"); sys.exit(1)

    # 4. Limit categories (smoke-test)
    if args.limit_cats > 0:
        groups = dict(list(groups.items())[:args.limit_cats])
        log.info(f"--limit-cats={args.limit_cats}: {list(groups.keys())}")

    # 5. Build splits (n-fold CV or single holdout)
    if args.split_mode == "holdout":
        log.info(f"Building holdout split (scale-down={args.scale_down:.0%}) …")
        folds = create_holdout_split(groups, args.scale_down, args.min_seqs, args.seed)
    else:
        log.info(f"Building {args.n_folds}-fold CV splits …")
        folds = create_cv_folds(groups, args.n_folds, args.min_seqs, args.seed)
    if not folds or not folds[0]: log.error("No categories remain after split filtering — aborting"); sys.exit(1)

    # 6. Load model
    embed_fn, model_id, emb_dim = load_embedding_model(args.model, args.device)

    # 7. Precompute ALL embeddings ONCE  (reused across all folds — 10× speedup)
    all_embs = precompute_all_embeddings(
        groups, embed_fn, args.frame_mode, args.max_frames,
        args.embed_batch, args.workers, temporal=args.temporal,
    )

    # 8. Evaluation loop (1 iteration for holdout, n_folds for CV)
    fold_results: list[dict[int, dict[str, dict]]] = []
    for fold_idx, fold_splits in enumerate(folds):
        log.info(f"═══ Fold {fold_idx+1}/{n_folds_eff} ═══")
        collection = create_ephemeral_collection()
        n_entries  = ingest_fold_from_precomputed(fold_splits, collection, all_embs, args.frame_mode, args.temporal)
        log.info(f"  Fold {fold_idx+1}: {n_entries:,} ChromaDB entries")
        if collection.count() == 0: log.warning(f"  Fold {fold_idx+1}: empty, skipping"); continue

        embedded_test_fold: list[tuple[dict, Optional[list]]] = [
            (seq, all_embs.get(seq["stem"]))
            for _, (_, test) in fold_splits.items() for seq in test
        ]
        fold_results.append(run_all_k_evaluations(embedded_test_fold, fold_splits, collection, k_values))

    if not fold_results: log.error("No folds completed — aborting"); sys.exit(1)

    # 9. Aggregate across folds
    all_results = aggregate_cv_results(fold_results, k_values, n_folds_eff)

    # 10. Report
    print_results_table(all_results, k_values, n_folds_eff, model_id, args.frame_mode)
    save_bar_chart(all_results, k_values, n_folds_eff, model_id, args.frame_mode, out_plot)
    write_markdown(all_results, k_values, n_folds_eff, args.seed, model_id, args.frame_mode, emb_dim, args.temporal, out_md)

    # 11. Visual examples docx — use last fold's collection
    last_splits = folds[-1]
    last_col    = create_ephemeral_collection()
    ingest_fold_from_precomputed(last_splits, last_col, all_embs, args.frame_mode, args.temporal)
    last_test = [(seq, all_embs.get(seq["stem"])) for _,(_, test) in last_splits.items() for seq in test]
    examples  = collect_k1_retrieval_examples(last_test, last_splits, last_col, n_per_cat=5)
    save_retrieval_docx(examples, all_results[k_values[0]], out_docx, model_id, args.frame_mode, args.temporal)


if __name__ == "__main__":
    main()