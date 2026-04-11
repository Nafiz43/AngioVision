#!/usr/bin/env python3
"""
generate_gt_doc.py

Generates a Word document (.docx) from a ground-truth CSV file, embedding
the mosaic.png image for each sequence sub-folder found in the test-data dir.

Folders on disk are named by SOPInstanceUID.
The CSV has columns: Accession, SOPInstanceUID, Question, Answer

Usage:
    python generate_gt_doc.py \
        --data-dir /data/Deep_Angiography/Validation_Data/test-data \
        --gt-csv   /data/Deep_Angiography/Validation_Data/test-data/gt.csv \
        --output   /data/Deep_Angiography/Validation_Data/test-data/gt_with_images.docx

All flags default to the paths above, so you can also just run:
    python generate_gt_doc.py

Dependencies:
    pip install python-docx pillow
"""

import argparse
import csv
import os
import sys
from collections import defaultdict
from datetime import datetime

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt, RGBColor

try:
    from PIL import Image as PILImage
    HAS_PIL = True
except ImportError:
    HAS_PIL = False
    print("⚠  Pillow not installed — falling back to PNG header parsing.")


# ── Constants ─────────────────────────────────────────────────────────────────
MAX_IMG_WIDTH_IN  = 6.0
MAX_IMG_HEIGHT_IN = 4.0

COLOR_HEADER_BG = "1F3864"
COLOR_HEADER_FG = "FFFFFF"
COLOR_ROW_ALT   = "EEF2F7"
COLOR_ANSWER    = "1A5276"
COLOR_TITLE     = "1F3864"
COLOR_SUBHEAD   = "2E75B6"
COLOR_WARN      = "CC0000"
COLOR_META      = "666666"

FONT_NAME = "Arial"


# ── Helpers ───────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def set_cell_borders(cell, color="CCCCCC"):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"),   "single")
        el.set(qn("w:sz"),    "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        borders.append(el)
    tcPr.append(borders)


def set_col_width(cell, width_inches):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW  = OxmlElement("w:tcW")
    tcW.set(qn("w:w"),    str(int(width_inches * 1440)))
    tcW.set(qn("w:type"), "dxa")
    tcPr.append(tcW)


def add_run(para, text, bold=False, italic=False,
            size_pt=11, color=None, font=FONT_NAME):
    run = para.add_run(text)
    run.bold        = bold
    run.italic      = italic
    run.font.name   = font
    run.font.size   = Pt(size_pt)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    return run


def heading_para(doc, text, size_pt=14, color=COLOR_TITLE, bottom_border=False):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    para.paragraph_format.space_before = Pt(10)
    para.paragraph_format.space_after  = Pt(4)
    add_run(para, text, bold=True, size_pt=size_pt, color=color)
    if bottom_border:
        pPr    = para._p.get_or_add_pPr()
        pBdr   = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"),   "single")
        bottom.set(qn("w:sz"),    "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), COLOR_SUBHEAD)
        pBdr.append(bottom)
        pPr.append(pBdr)
    return para


def scaled_image_size(img_path):
    """Return (width_in, height_in) scaled to fit within MAX constraints."""
    if HAS_PIL:
        with PILImage.open(img_path) as img:
            w_px, h_px = img.size
    else:
        import struct
        with open(img_path, "rb") as f:
            f.read(16)   # PNG sig (8) + IHDR length (4) + "IHDR" tag (4)
            w_px = struct.unpack(">I", f.read(4))[0]
            h_px = struct.unpack(">I", f.read(4))[0]

    if w_px == 0 or h_px == 0:
        return MAX_IMG_WIDTH_IN, MAX_IMG_HEIGHT_IN

    ratio = min(MAX_IMG_WIDTH_IN  / (w_px / 96),
                MAX_IMG_HEIGHT_IN / (h_px / 96),
                1.0)
    return (w_px / 96) * ratio, (h_px / 96) * ratio


def read_csv(path):
    rows = []
    with open(path, newline="", encoding="utf-8-sig") as f:
        for row in csv.DictReader(f):
            rows.append({k.strip(): v.strip() for k, v in row.items()})
    return rows


# ── Document builder ──────────────────────────────────────────────────────────

def build_document(data_dir, gt_csv, output_path):
    rows = read_csv(gt_csv)
    print(f"✅  Loaded {len(rows)} rows from gt.csv")

    # ── Group by SOPInstanceUID ───────────────────────────────────────────────
    # Each unique SOPInstanceUID corresponds to one sequence folder on disk.
    by_sop = defaultdict(list)
    for row in rows:
        sop = row.get("SOPInstanceUID") or row.get("sopinstanceuid") or ""
        by_sop[sop].append(row)

    sequences = list(by_sop.keys())
    print(f"📦  Found {len(sequences)} unique SOPInstanceUIDs\n")

    doc = Document()

    # ── Page setup ────────────────────────────────────────────────────────────
    sec = doc.sections[0]
    sec.page_width    = Inches(8.5)
    sec.page_height   = Inches(11)
    sec.top_margin    = Inches(0.75)
    sec.bottom_margin = Inches(0.75)
    sec.left_margin   = Inches(0.9)
    sec.right_margin  = Inches(0.9)

    # ── Title ─────────────────────────────────────────────────────────────────
    tp = doc.add_paragraph()
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tp.paragraph_format.space_before = Pt(24)
    tp.paragraph_format.space_after  = Pt(8)
    add_run(tp, "Ground Truth Report — Deep Angiography",
            bold=True, size_pt=20, color=COLOR_TITLE)

    mp = doc.add_paragraph()
    mp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    mp.paragraph_format.space_after = Pt(4)
    add_run(mp,
            f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}   |   "
            f"Sequences: {len(sequences)}   |   Total Q&A pairs: {len(rows)}",
            size_pt=9, color=COLOR_META)

    doc.add_page_break()

    # ── One section per SOPInstanceUID ────────────────────────────────────────
    for idx, sop_uid in enumerate(sequences):
        seq_rows    = by_sop[sop_uid]
        # Folder on disk is named exactly by SOPInstanceUID
        mosaic_path = os.path.join(data_dir, sop_uid, "mosaic.png")
        accession   = (seq_rows[0].get("Accession")
                       or seq_rows[0].get("accession")
                       or "N/A")

        print(f"[{idx + 1}/{len(sequences)}]  SOP: {sop_uid}  (Accession: {accession})")

        # --- Heading: Accession number (human-readable) ----------------------
        heading_para(doc,
                     f"Sequence {idx + 1}  |  Accession: {accession}",
                     size_pt=13, color=COLOR_TITLE, bottom_border=True)

        # --- SOP UID subtitle ------------------------------------------------
        sp = doc.add_paragraph()
        sp.paragraph_format.space_before = Pt(2)
        sp.paragraph_format.space_after  = Pt(6)
        add_run(sp, f"SOPInstanceUID: {sop_uid}",
                size_pt=8, color=COLOR_META, font="Courier New")

        # --- Mosaic image ----------------------------------------------------
        if os.path.exists(mosaic_path):
            w_in, h_in = scaled_image_size(mosaic_path)
            print(f"   🖼   mosaic → {w_in:.2f}\" × {h_in:.2f}\"")
            ip = doc.add_paragraph()
            ip.alignment = WD_ALIGN_PARAGRAPH.CENTER
            ip.paragraph_format.space_before = Pt(4)
            ip.paragraph_format.space_after  = Pt(4)
            ip.add_run().add_picture(mosaic_path,
                                     width=Inches(w_in),
                                     height=Inches(h_in))
        else:
            print(f"   ⚠   mosaic.png NOT FOUND at: {mosaic_path}")
            wp = doc.add_paragraph()
            wp.paragraph_format.space_before = Pt(4)
            wp.paragraph_format.space_after  = Pt(4)
            add_run(wp,
                    f"⚠  mosaic.png not found  ({mosaic_path})",
                    italic=True, size_pt=9, color=COLOR_WARN)

        # --- Q&A sub-heading -------------------------------------------------
        qh = doc.add_paragraph()
        qh.paragraph_format.space_before = Pt(8)
        qh.paragraph_format.space_after  = Pt(4)
        add_run(qh,
                f"Ground Truth Q&A  ({len(seq_rows)} item"
                f"{'s' if len(seq_rows) != 1 else ''})",
                bold=True, size_pt=11, color=COLOR_SUBHEAD)

        # --- Q&A table -------------------------------------------------------
        # cols: # | Accession | Question | Answer
        col_widths = [0.30, 1.50, 3.20, 2.70]
        headers    = ["#", "Accession", "Question", "Answer"]

        tbl = doc.add_table(rows=1, cols=len(headers))
        tbl.style = "Table Grid"

        # header row
        for ci, (cell, hdr) in enumerate(zip(tbl.rows[0].cells, headers)):
            set_cell_bg(cell, COLOR_HEADER_BG)
            set_cell_borders(cell, "444444")
            set_col_width(cell, col_widths[ci])
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after  = Pt(3)
            add_run(p, hdr, bold=True, size_pt=9, color=COLOR_HEADER_FG)

        # data rows
        for ri, row in enumerate(seq_rows):
            acc_val = row.get("Accession") or row.get("accession") or "—"
            q       = row.get("Question")  or row.get("question")  or "—"
            a       = row.get("Answer")    or row.get("answer")    or "—"
            bg      = "FFFFFF" if ri % 2 == 0 else COLOR_ROW_ALT

            cells = tbl.add_row().cells
            data  = [str(ri + 1), acc_val, q, a]
            for ci, (cell, val) in enumerate(zip(cells, data)):
                set_cell_bg(cell, bg)
                set_cell_borders(cell)
                set_col_width(cell, col_widths[ci])
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                p = cell.paragraphs[0]
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after  = Pt(2)
                add_run(p, val,
                        bold=(ci == 3),
                        size_pt=9,
                        color=COLOR_ANSWER if ci == 3 else None)

        # page break between sequences (not after last)
        if idx < len(sequences) - 1:
            doc.add_page_break()

    # ── Save ──────────────────────────────────────────────────────────────────
    os.makedirs(os.path.dirname(os.path.abspath(output_path)), exist_ok=True)
    doc.save(output_path)
    size_kb = os.path.getsize(output_path) / 1024
    print(f"\n✅  Done! Wrote {size_kb:.1f} KB → {output_path}")


# ── CLI ───────────────────────────────────────────────────────────────────────

def main():
    DEFAULT_DATA = "/data/Deep_Angiography/Validation_Data/test-data"

    parser = argparse.ArgumentParser(
        description="Generate a GT Word doc with mosaic images. "
                    "Folders are matched by SOPInstanceUID."
    )
    parser.add_argument("--data-dir", default=DEFAULT_DATA,
                        help="Root test-data directory")
    parser.add_argument("--gt-csv",   default=None,
                        help="Path to gt.csv (default: <data-dir>/gt.csv)")
    parser.add_argument("--output",   default=None,
                        help="Output .docx (default: <data-dir>/gt_with_images.docx)")
    args = parser.parse_args()

    data_dir = args.data_dir
    gt_csv   = args.gt_csv  or os.path.join(data_dir, "gt.csv")
    output   = args.output  or os.path.join(data_dir, "gt_with_images.docx")

    print(f"📂  Data dir : {data_dir}")
    print(f"📄  GT CSV   : {gt_csv}")
    print(f"💾  Output   : {output}\n")

    if not os.path.isfile(gt_csv):
        sys.exit(f"❌  gt.csv not found: {gt_csv}")

    build_document(data_dir, gt_csv, output)


if __name__ == "__main__":
    main()