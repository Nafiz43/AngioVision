# # #!/usr/bin/env python3
# # """
# # JSON -> DOCX with embedded mosaic images.

# # Folder layout (confirmed from your screenshot):
# #   /data/Deep_Angiography/DICOM_Sequence_Processed/<group>/<UID>/mosaic.png

# # Fixes / improvements:
# # - Robust UID normalization (strip whitespace + remove hidden chars)
# # - Robust UID extraction from JSON sequence entries (handles multiple key names)
# # - Index mosaics by UID directory name (fast)
# # - Fallback search: if UID not in index, try searching for <UID>/mosaic.png
# # - After saving, verify DOCX actually contains images by inspecting zip: word/media/*
# # - Write logs:
# #     - missing_mosaics.txt
# #     - docx_image_errors.txt
# #     - docx_without_media.txt
# # """

# # import json
# # import re
# # import zipfile
# # from pathlib import Path
# # from typing import Dict, Optional, List, Iterable, Tuple

# # from docx import Document
# # from docx.shared import Inches
# # from docx.enum.text import WD_ALIGN_PARAGRAPH


# # REPORTS_JSON_ROOT = Path("/data/Deep_Angiography/Reports/Report_List_v01_01_sequences_json")
# # DICOM_SEQ_ROOT    = Path("/data/Deep_Angiography/DICOM_Sequence_Processed")
# # OUTPUT_DOCX_ROOT  = Path("/data/Deep_Angiography/Reports/Report_List_v01_01_sequences_docx")

# # MAX_IMAGE_WIDTH_INCHES = 6.5

# # # DICOM UID: digits + dots
# # UID_RE = re.compile(r"^\d+(?:\.\d+)+$")


# # def normalize_uid(uid: str) -> str:
# #     """
# #     Make UID matching reliable:
# #     - strip whitespace
# #     - remove zero-width and other invisible chars
# #     """
# #     if uid is None:
# #         return ""
# #     uid = str(uid)
# #     uid = uid.replace("\u200b", "").replace("\ufeff", "")  # zero-width, BOM
# #     uid = uid.strip()
# #     return uid


# # def safe_get(d: dict, key: str, default=None):
# #     return d.get(key, default) if isinstance(d, dict) else default


# # def add_heading(doc: Document, text: str, level: int = 1):
# #     doc.add_heading(text, level=level)


# # def add_kv_paragraph(doc: Document, k: str, v: str):
# #     p = doc.add_paragraph()
# #     p.add_run(f"{k}: ").bold = True
# #     p.add_run(str(v))


# # def build_uid_to_mosaic_index(dicom_root: Path) -> Dict[str, Path]:
# #     """
# #     Index mosaics by UID directory name:
# #       .../<UID>/mosaic.png

# #     This matches your screenshot layout exactly.
# #     """
# #     index: Dict[str, Path] = {}
# #     count = 0
# #     for mosaic_path in dicom_root.rglob("mosaic.png"):
# #         if not mosaic_path.is_file():
# #             continue
# #         uid_dir = mosaic_path.parent.name
# #         uid_dir_n = normalize_uid(uid_dir)
# #         if UID_RE.match(uid_dir_n):
# #             index.setdefault(uid_dir_n, mosaic_path)
# #             count += 1
# #     print(f"[index] Found {count:,} mosaic.png files under UID directories")
# #     print(f"[index] Indexed {len(index):,} unique UID -> mosaic paths")
# #     return index


# # def fallback_find_mosaic(dicom_root: Path, uid: str) -> Optional[Path]:
# #     """
# #     If the index misses (because of odd paths), try direct search for:
# #       **/<UID>/mosaic.png
# #     """
# #     uid = normalize_uid(uid)
# #     if not uid:
# #         return None
# #     # This is still reasonably fast if used sparingly.
# #     for p in dicom_root.rglob("mosaic.png"):
# #         if p.is_file() and normalize_uid(p.parent.name) == uid:
# #             return p
# #     return None


# # def extract_uid_from_sequence(seq: dict) -> str:
# #     """
# #     Robustly extract UID from various possible key names.
# #     Your current JSON uses 'sequence_instance_uid', but we support others too.
# #     """
# #     candidates = [
# #         "sequence_instance_uid",
# #         "SequenceInstanceUID",
# #         "series_instance_uid",
# #         "SeriesInstanceUID",
# #         "uid",
# #         "UID",
# #     ]
# #     for k in candidates:
# #         v = seq.get(k)
# #         v = normalize_uid(v)
# #         if v:
# #             return v

# #     # Sometimes stored under nested structure
# #     # (kept conservative; only try a couple of common shapes)
# #     v = normalize_uid(seq.get("sequence", {}).get("sequence_instance_uid") if isinstance(seq.get("sequence"), dict) else "")
# #     if v:
# #         return v

# #     return ""


# # def docx_contains_media(docx_path: Path) -> bool:
# #     """
# #     Check whether the DOCX zip has any embedded images under word/media/.
# #     """
# #     try:
# #         with zipfile.ZipFile(docx_path, "r") as z:
# #             names = z.namelist()
# #             return any(n.startswith("word/media/") for n in names)
# #     except Exception:
# #         return False


# # def add_sequence_block(
# #     doc: Document,
# #     seq_number: Optional[int],
# #     uid: str,
# #     verbatim_text: str,
# #     mosaic_path: Optional[Path],
# #     missing_log: List[str],
# #     error_log: List[str],
# # ):
# #     title = f"Sequence {seq_number}" if seq_number is not None else "Sequence"
# #     add_heading(doc, title, level=2)

# #     add_kv_paragraph(doc, "SequenceInstanceUID", uid)

# #     doc.add_paragraph().add_run("Verbatim text:").bold = True
# #     doc.add_paragraph(verbatim_text or "")

# #     if mosaic_path and mosaic_path.exists():
# #         doc.add_paragraph().add_run("Mosaic:").bold = True
# #         # Insert image
# #         pic_par = doc.add_paragraph()
# #         run = pic_par.add_run()
# #         try:
# #             run.add_picture(str(mosaic_path), width=Inches(MAX_IMAGE_WIDTH_INCHES))
# #             pic_par.alignment = WD_ALIGN_PARAGRAPH.CENTER
# #         except Exception as e:
# #             error_log.append(f"[add_picture ERROR] uid={uid} file={mosaic_path} err={repr(e)}")
# #             warn = doc.add_paragraph()
# #             warn.add_run("ERROR embedding mosaic image (see docx_image_errors.txt).").bold = True
# #     else:
# #         missing_log.append(f"[missing mosaic] uid={uid}")
# #         warn = doc.add_paragraph()
# #         warn.add_run("Mosaic not found for this UID.").italic = True

# #     doc.add_page_break()


# # def json_to_docx(
# #     json_path: Path,
# #     out_docx_path: Path,
# #     uid_to_mosaic: Dict[str, Path],
# #     missing_log: List[str],
# #     error_log: List[str],
# # ):
# #     data = json.loads(json_path.read_text(encoding="utf-8", errors="replace"))
# #     doc = Document()

# #     anon_acc = safe_get(data, "Anon Acc #", safe_get(data, "Anon Acc", ""))
# #     add_heading(doc, f"Report sequences: {json_path.stem}", level=1)
# #     if anon_acc:
# #         add_kv_paragraph(doc, "Anon Acc #", anon_acc)

# #     radrpt = safe_get(data, "radrpt", "")
# #     if radrpt:
# #         add_heading(doc, "Rad report text (radrpt)", level=2)
# #         doc.add_paragraph(radrpt)

# #     sequences: List[dict] = safe_get(data, "sequences", []) or []
# #     if not sequences:
# #         doc.add_paragraph("No sequences found in this JSON.")
# #     else:
# #         add_heading(doc, "Sequences", level=1)
# #         for seq in sequences:
# #             uid = extract_uid_from_sequence(seq)
# #             verbatim = safe_get(seq, "verbatim_text", "") or ""
# #             seq_no = safe_get(seq, "sequence_number", None)

# #             mosaic_path = None
# #             if uid:
# #                 mosaic_path = uid_to_mosaic.get(uid)
# #                 if mosaic_path is None:
# #                     mosaic_path = fallback_find_mosaic(DICOM_SEQ_ROOT, uid)

# #             add_sequence_block(
# #                 doc=doc,
# #                 seq_number=seq_no,
# #                 uid=uid,
# #                 verbatim_text=verbatim,
# #                 mosaic_path=mosaic_path,
# #                 missing_log=missing_log,
# #                 error_log=error_log,
# #             )

# #     out_docx_path.parent.mkdir(parents=True, exist_ok=True)
# #     doc.save(str(out_docx_path))


# # def main():
# #     if not REPORTS_JSON_ROOT.exists():
# #         raise SystemExit(f"JSON root not found: {REPORTS_JSON_ROOT}")
# #     if not DICOM_SEQ_ROOT.exists():
# #         raise SystemExit(f"DICOM processed root not found: {DICOM_SEQ_ROOT}")

# #     OUTPUT_DOCX_ROOT.mkdir(parents=True, exist_ok=True)

# #     uid_to_mosaic = build_uid_to_mosaic_index(DICOM_SEQ_ROOT)

# #     json_files = sorted(REPORTS_JSON_ROOT.rglob("*.json"))
# #     print(f"Found {len(json_files):,} JSON files")

# #     missing_log: List[str] = []
# #     error_log: List[str] = []
# #     no_media_log: List[str] = []

# #     for jp in json_files:
# #         rel = jp.relative_to(REPORTS_JSON_ROOT)
# #         out_docx = OUTPUT_DOCX_ROOT / rel.with_suffix(".docx")

# #         try:
# #             json_to_docx(jp, out_docx, uid_to_mosaic, missing_log, error_log)

# #             # Hard verification: does the docx contain embedded media?
# #             if not docx_contains_media(out_docx):
# #                 no_media_log.append(f"[no media in docx] {out_docx}")
# #             print(f"Wrote: {out_docx}")

# #         except Exception as e:
# #             error_log.append(f"[DOCX ERROR] json={jp} err={repr(e)}")
# #             print(f"[DOCX ERROR] {jp} -> {repr(e)}")

# #     if missing_log:
# #         miss_path = OUTPUT_DOCX_ROOT / "missing_mosaics.txt"
# #         miss_path.write_text("\n".join(missing_log) + "\n", encoding="utf-8")
# #         print(f"\nWARNING: Missing mosaics for {len(missing_log):,} sequence entries. Saved: {miss_path}")

# #     if error_log:
# #         err_path = OUTPUT_DOCX_ROOT / "docx_image_errors.txt"
# #         err_path.write_text("\n".join(error_log) + "\n", encoding="utf-8")
# #         print(f"\nWARNING: Encountered {len(error_log):,} errors. Saved: {err_path}")

# #     if no_media_log:
# #         nm_path = OUTPUT_DOCX_ROOT / "docx_without_media.txt"
# #         nm_path.write_text("\n".join(no_media_log) + "\n", encoding="utf-8")
# #         print(f"\nWARNING: {len(no_media_log):,} DOCX files had no embedded media. Saved: {nm_path}")

# #     print("\nDone.")


# # if __name__ == "__main__":
# #     main()
# #!/usr/bin/env python3
# """
# JSON -> DOCX with embedded mosaic images + per-sequence Q/A blocks.

# Folder layout (confirmed from your screenshot):
#   /data/Deep_Angiography/DICOM_Sequence_Processed/<group>/<UID>/mosaic.png

# Features / improvements:
# - Robust UID normalization (strip whitespace + remove hidden chars)
# - Robust UID extraction from JSON sequence entries (handles multiple key names)
# - Index mosaics by UID directory name (fast)
# - Fallback search: if UID not in index, try searching for <UID>/mosaic.png
# - Embed per-sequence Q/A under each sequence
# - Per DOCX: print mosaics found vs not found
# - Global: print total mosaics found vs not found at end
# - After saving, verify DOCX actually contains images by inspecting zip: word/media/*
# - Write logs:
#     - missing_mosaics.txt
#     - docx_image_errors.txt
#     - docx_without_media.txt
# """

# import json
# import re
# import zipfile
# from pathlib import Path
# from typing import Dict, Optional, List, Tuple, Any

# from docx import Document
# from docx.shared import Inches
# from docx.enum.text import WD_ALIGN_PARAGRAPH


# REPORTS_JSON_ROOT = Path("/data/Deep_Angiography/Reports/Report_List_v01_01_sequences_json")
# DICOM_SEQ_ROOT    = Path("/data/Deep_Angiography/DICOM_Sequence_Processed")
# OUTPUT_DOCX_ROOT  = Path("/data/Deep_Angiography/Reports/Report_List_v01_01_sequences_docx")

# MAX_IMAGE_WIDTH_INCHES = 6.5

# # DICOM UID: digits + dots
# UID_RE = re.compile(r"^\d+(?:\.\d+)+$")


# def normalize_uid(uid: str) -> str:
#     """
#     Make UID matching reliable:
#     - strip whitespace
#     - remove zero-width and other invisible chars
#     """
#     if uid is None:
#         return ""
#     uid = str(uid)
#     uid = uid.replace("\u200b", "").replace("\ufeff", "")  # zero-width, BOM
#     uid = uid.strip()
#     return uid


# def safe_get(d: Any, key: str, default=None):
#     return d.get(key, default) if isinstance(d, dict) else default


# def add_heading(doc: Document, text: str, level: int = 1):
#     doc.add_heading(text, level=level)


# def add_kv_paragraph(doc: Document, k: str, v: Any):
#     p = doc.add_paragraph()
#     p.add_run(f"{k}: ").bold = True
#     p.add_run("" if v is None else str(v))


# def build_uid_to_mosaic_index(dicom_root: Path) -> Dict[str, Path]:
#     """
#     Index mosaics by UID directory name:
#       .../<UID>/mosaic.png
#     """
#     index: Dict[str, Path] = {}
#     count = 0
#     for mosaic_path in dicom_root.rglob("mosaic.png"):
#         if not mosaic_path.is_file():
#             continue
#         uid_dir = mosaic_path.parent.name
#         uid_dir_n = normalize_uid(uid_dir)
#         if UID_RE.match(uid_dir_n):
#             index.setdefault(uid_dir_n, mosaic_path)
#             count += 1
#     print(f"[index] Found {count:,} mosaic.png files under UID directories")
#     print(f"[index] Indexed {len(index):,} unique UID -> mosaic paths")
#     return index


# def fallback_find_mosaic(dicom_root: Path, uid: str) -> Optional[Path]:
#     """
#     If the index misses (because of odd paths), try direct search for:
#       **/<UID>/mosaic.png
#     """
#     uid = normalize_uid(uid)
#     if not uid:
#         return None
#     for p in dicom_root.rglob("mosaic.png"):
#         if p.is_file() and normalize_uid(p.parent.name) == uid:
#             return p
#     return None


# def extract_uid_from_sequence(seq: dict) -> str:
#     """
#     Robustly extract UID from various possible key names.
#     Your current JSON uses 'sequence_instance_uid', but we support others too.
#     """
#     candidates = [
#         "sequence_instance_uid",
#         "SequenceInstanceUID",
#         "series_instance_uid",
#         "SeriesInstanceUID",
#         "uid",
#         "UID",
#     ]
#     for k in candidates:
#         v = normalize_uid(seq.get(k))
#         if v:
#             return v

#     # Sometimes stored under nested structure
#     nested = seq.get("sequence")
#     if isinstance(nested, dict):
#         v = normalize_uid(nested.get("sequence_instance_uid") or nested.get("SeriesInstanceUID"))
#         if v:
#             return v

#     return ""


# def docx_contains_media(docx_path: Path) -> bool:
#     """
#     Check whether the DOCX zip has any embedded images under word/media/.
#     """
#     try:
#         with zipfile.ZipFile(docx_path, "r") as z:
#             names = z.namelist()
#             return any(n.startswith("word/media/") for n in names)
#     except Exception:
#         return False


# def add_qa_block(doc: Document, qa_items: List[dict]):
#     """
#     Add the sequence-level QA list (question, answer, confidence, evidence, notes).
#     """
#     if not qa_items:
#         return

#     doc.add_paragraph().add_run("Q/A:").bold = True

#     for i, qa in enumerate(qa_items, start=1):
#         q = safe_get(qa, "question", "")
#         a = safe_get(qa, "answer", "")
#         conf = safe_get(qa, "confidence", "")
#         evidence = safe_get(qa, "evidence", []) or []
#         notes = safe_get(qa, "notes", "")

#         p = doc.add_paragraph(style="List Number")
#         run = p.add_run(q if q else f"QA item {i}")
#         run.bold = True

#         # Answer line
#         ans_p = doc.add_paragraph(style="List Bullet")
#         ans_p.add_run("Answer: ").bold = True
#         ans_p.add_run(str(a))

#         # Confidence
#         if conf:
#             c_p = doc.add_paragraph(style="List Bullet")
#             c_p.add_run("Confidence: ").bold = True
#             c_p.add_run(str(conf))

#         # Evidence
#         if isinstance(evidence, list) and len(evidence) > 0:
#             e_p = doc.add_paragraph(style="List Bullet")
#             e_p.add_run("Evidence:").bold = True
#             for ev in evidence:
#                 doc.add_paragraph(str(ev), style="List Bullet 2")

#         # Notes
#         if notes:
#             n_p = doc.add_paragraph(style="List Bullet")
#             n_p.add_run("Notes: ").bold = True
#             n_p.add_run(str(notes))


# def add_sequence_block(
#     doc: Document,
#     seq_number: Optional[int],
#     uid: str,
#     verbatim_text: str,
#     qa_items: List[dict],
#     mosaic_path: Optional[Path],
#     missing_log: List[str],
#     error_log: List[str],
# ) -> bool:
#     """
#     Returns True if mosaic embedded successfully (or at least found & attempted),
#     False if mosaic missing (not found).
#     """
#     title = f"Sequence {seq_number}" if seq_number is not None else "Sequence"
#     add_heading(doc, title, level=2)

#     add_kv_paragraph(doc, "SequenceInstanceUID", uid)

#     doc.add_paragraph().add_run("Verbatim text:").bold = True
#     doc.add_paragraph(verbatim_text or "")

#     # Add QA block under each sequence
#     add_qa_block(doc, qa_items)

#     found = False
#     if mosaic_path and mosaic_path.exists():
#         found = True
#         doc.add_paragraph().add_run("Mosaic:").bold = True
#         pic_par = doc.add_paragraph()
#         run = pic_par.add_run()
#         try:
#             run.add_picture(str(mosaic_path), width=Inches(MAX_IMAGE_WIDTH_INCHES))
#             pic_par.alignment = WD_ALIGN_PARAGRAPH.CENTER
#         except Exception as e:
#             error_log.append(f"[add_picture ERROR] uid={uid} file={mosaic_path} err={repr(e)}")
#             warn = doc.add_paragraph()
#             warn.add_run("ERROR embedding mosaic image (see docx_image_errors.txt).").bold = True
#     else:
#         missing_log.append(f"[missing mosaic] uid={uid}")
#         warn = doc.add_paragraph()
#         warn.add_run("Mosaic not found for this UID.").italic = True

#     doc.add_page_break()
#     return found


# def json_to_docx(
#     json_path: Path,
#     out_docx_path: Path,
#     uid_to_mosaic: Dict[str, Path],
#     missing_log: List[str],
#     error_log: List[str],
# ) -> Tuple[int, int]:
#     """
#     Returns (found_count, not_found_count) for mosaics in THIS JSON file.
#     """
#     data = json.loads(json_path.read_text(encoding="utf-8", errors="replace"))
#     doc = Document()

#     anon_acc = safe_get(data, "Anon Acc #", safe_get(data, "Anon Acc", ""))
#     add_heading(doc, f"Report sequences: {json_path.stem}", level=1)
#     if anon_acc:
#         add_kv_paragraph(doc, "Anon Acc #", anon_acc)

#     radrpt = safe_get(data, "radrpt", "")
#     if radrpt:
#         add_heading(doc, "Rad report text (radrpt)", level=2)
#         doc.add_paragraph(radrpt)

#     sequences: List[dict] = safe_get(data, "sequences", []) or []

#     found_count = 0
#     not_found_count = 0

#     if not sequences:
#         doc.add_paragraph("No sequences found in this JSON.")
#     else:
#         add_heading(doc, "Sequences", level=1)
#         for seq in sequences:
#             uid = extract_uid_from_sequence(seq)
#             verbatim = safe_get(seq, "verbatim_text", "") or ""
#             seq_no = safe_get(seq, "sequence_number", None)
#             qa_items = safe_get(seq, "qa", []) or []

#             mosaic_path = None
#             if uid:
#                 mosaic_path = uid_to_mosaic.get(uid)
#                 if mosaic_path is None:
#                     mosaic_path = fallback_find_mosaic(DICOM_SEQ_ROOT, uid)

#             was_found = add_sequence_block(
#                 doc=doc,
#                 seq_number=seq_no,
#                 uid=uid,
#                 verbatim_text=verbatim,
#                 qa_items=qa_items,
#                 mosaic_path=mosaic_path,
#                 missing_log=missing_log,
#                 error_log=error_log,
#             )

#             if was_found:
#                 found_count += 1
#             else:
#                 not_found_count += 1

#     out_docx_path.parent.mkdir(parents=True, exist_ok=True)
#     doc.save(str(out_docx_path))
#     return found_count, not_found_count


# def main():
#     if not REPORTS_JSON_ROOT.exists():
#         raise SystemExit(f"JSON root not found: {REPORTS_JSON_ROOT}")
#     if not DICOM_SEQ_ROOT.exists():
#         raise SystemExit(f"DICOM processed root not found: {DICOM_SEQ_ROOT}")

#     OUTPUT_DOCX_ROOT.mkdir(parents=True, exist_ok=True)

#     uid_to_mosaic = build_uid_to_mosaic_index(DICOM_SEQ_ROOT)

#     json_files = sorted(REPORTS_JSON_ROOT.rglob("*.json"))
#     print(f"Found {len(json_files):,} JSON files")

#     missing_log: List[str] = []
#     error_log: List[str] = []
#     no_media_log: List[str] = []

#     total_found = 0
#     total_not_found = 0

#     for jp in json_files:
#         rel = jp.relative_to(REPORTS_JSON_ROOT)
#         out_docx = OUTPUT_DOCX_ROOT / rel.with_suffix(".docx")

#         try:
#             found_count, not_found_count = json_to_docx(jp, out_docx, uid_to_mosaic, missing_log, error_log)

#             total_found += found_count
#             total_not_found += not_found_count

#             # Hard verification: does the docx contain embedded media?
#             if not docx_contains_media(out_docx):
#                 no_media_log.append(f"[no media in docx] {out_docx}")

#             print(f"Wrote: {out_docx}")
#             print(f"  mosaics: found={found_count} not_found={not_found_count}")

#         except Exception as e:
#             error_log.append(f"[DOCX ERROR] json={jp} err={repr(e)}")
#             print(f"[DOCX ERROR] {jp} -> {repr(e)}")

#     if missing_log:
#         miss_path = OUTPUT_DOCX_ROOT / "missing_mosaics.txt"
#         miss_path.write_text("\n".join(missing_log) + "\n", encoding="utf-8")
#         print(f"\nWARNING: Missing mosaics for {len(missing_log):,} sequence entries. Saved: {miss_path}")

#     if error_log:
#         err_path = OUTPUT_DOCX_ROOT / "docx_image_errors.txt"
#         err_path.write_text("\n".join(error_log) + "\n", encoding="utf-8")
#         print(f"\nWARNING: Encountered {len(error_log):,} errors. Saved: {err_path}")

#     if no_media_log:
#         nm_path = OUTPUT_DOCX_ROOT / "docx_without_media.txt"
#         nm_path.write_text("\n".join(no_media_log) + "\n", encoding="utf-8")
#         print(f"\nWARNING: {len(no_media_log):,} DOCX files had no embedded media. Saved: {nm_path}")

#     print("\n=== FINAL MOSAIC COUNTS ===")
#     print(f"Total mosaics found:     {total_found}")
#     print(f"Total mosaics not found: {total_not_found}")
#     print("\nDone.")


# if __name__ == "__main__":
#     main()


#!/usr/bin/env python3
"""
json_to_docx_with_on_the_fly_mosaics.py

JSON -> DOCX with embedded mosaic images + per-sequence Q/A blocks.

Folder layout (your dataset):
  /data/Deep_Angiography/DICOM_Sequence_Processed/<group>/<UID>/frames/*
  /data/Deep_Angiography/DICOM_Sequence_Processed/<group>/<UID>/mosaic.png   (may be missing)

UPDATE (what you asked):
- If mosaic.png is NOT present, this script will generate it ON THE FLY from frames/
  and then embed it in the DOCX.

Other features:
- Robust UID normalization (strip whitespace + remove hidden chars)
- Robust UID extraction from JSON sequence entries (handles multiple key names)
- Fast indexing:
    * UID -> existing mosaic.png (if present)
    * UID -> sequence dir (by discovering <UID>/frames directories)
- Fallback search: if UID not in indexes, try searching for <UID>/mosaic.png
- Embed per-sequence Q/A under each sequence
- Per DOCX: print mosaics found vs generated vs missing
- Global: print total mosaics found vs generated vs missing at end
- After saving, verify DOCX contains images by inspecting zip: word/media/*
- Write logs:
    - missing_mosaics.txt
    - generated_mosaics.txt
    - mosaic_generation_errors.txt
    - docx_image_errors.txt
    - docx_without_media.txt
"""

import json
import math
import os
import re
import sys
import zipfile
from pathlib import Path
from typing import Dict, Optional, List, Tuple, Any

from concurrent.futures import ThreadPoolExecutor

from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from PIL import Image, ImageOps

# ----------------------------
# Paths
# ----------------------------
REPORTS_JSON_ROOT = Path("/data/Deep_Angiography/Reports/Report_List_v01_01_sequences_json")
DICOM_SEQ_ROOT    = Path("/data/Deep_Angiography/DICOM_Sequence_Processed")
OUTPUT_DOCX_ROOT  = Path("/data/Deep_Angiography/Reports/Report_List_v01_01_sequences_docx")

# ----------------------------
# Mosaic generation settings
# ----------------------------
FRAMES_SUBDIR = "frames"
MOSAIC_NAME = "mosaic.png"

MAX_FRAMES = 144       # cap how many frames to include in mosaic
STRIDE = 2             # sample every STRIDE frames before capping
TILE_SIZE = (384, 384) # (W,H) per tile
MOSAIC_MAX_COLS = 6    # auto layout max columns

FRAME_LOAD_THREADS = 4 # threads to load/decode frames for a single mosaic

# DOCX embedding image width (inches)
MAX_IMAGE_WIDTH_INCHES = 6.5

# DICOM UID: digits + dots
UID_RE = re.compile(r"^\d+(?:\.\d+)+$")

IMAGE_EXTS = {".jpg", ".jpeg", ".png", ".bmp", ".tif", ".tiff", ".webp"}


def normalize_uid(uid: str) -> str:
    """Strip whitespace + remove invisible chars to make matching reliable."""
    if uid is None:
        return ""
    uid = str(uid)
    uid = uid.replace("\u200b", "").replace("\ufeff", "")  # zero-width, BOM
    uid = uid.strip()
    return uid


def safe_get(d: Any, key: str, default=None):
    return d.get(key, default) if isinstance(d, dict) else default


def add_heading(doc: Document, text: str, level: int = 1):
    doc.add_heading(text, level=level)


def add_kv_paragraph(doc: Document, k: str, v: Any):
    p = doc.add_paragraph()
    p.add_run(f"{k}: ").bold = True
    p.add_run("" if v is None else str(v))


# ----------------------------
# Indexing helpers
# ----------------------------
def build_uid_to_mosaic_index(dicom_root: Path) -> Dict[str, Path]:
    """
    Index existing mosaics by UID directory name:
      .../<UID>/mosaic.png
    """
    index: Dict[str, Path] = {}
    count = 0
    for mosaic_path in dicom_root.rglob(MOSAIC_NAME):
        if not mosaic_path.is_file():
            continue
        uid_dir = mosaic_path.parent.name
        uid_dir_n = normalize_uid(uid_dir)
        if UID_RE.match(uid_dir_n):
            index.setdefault(uid_dir_n, mosaic_path)
            count += 1
    print(f"[index] Found {count:,} {MOSAIC_NAME} files under UID directories")
    print(f"[index] Indexed {len(index):,} unique UID -> mosaic paths")
    return index


def build_uid_to_seqdir_index_by_frames(dicom_root: Path, frames_subdir: str) -> Dict[str, Path]:
    """
    Index UID -> sequence directory by discovering directories that contain frames_subdir.
    Expected structure:
        .../<UID>/<frames_subdir>/*
    So the UID directory is frames_dir.parent.
    """
    index: Dict[str, Path] = {}
    seen = 0

    # Find all ".../frames" directories (or whatever frames_subdir is)
    for frames_dir in dicom_root.rglob(frames_subdir):
        if not frames_dir.is_dir():
            continue

        seq_dir = frames_dir.parent
        uid = normalize_uid(seq_dir.name)
        if not UID_RE.match(uid):
            continue

        # Make sure it actually has at least one image in frames/
        try:
            has_img = any(
                p.is_file() and p.suffix.lower() in IMAGE_EXTS
                for p in frames_dir.iterdir()
            )
        except PermissionError:
            continue

        if not has_img:
            continue

        index.setdefault(uid, seq_dir)
        seen += 1

    print(f"[index] Indexed {len(index):,} UID -> seq_dir by scanning '{frames_subdir}' folders (seen={seen:,})")
    return index


def fallback_find_mosaic(dicom_root: Path, uid: str) -> Optional[Path]:
    """
    Slow fallback: search for any mosaic.png whose parent directory name equals UID.
    """
    uid = normalize_uid(uid)
    if not uid:
        return None
    for p in dicom_root.rglob(MOSAIC_NAME):
        if p.is_file() and normalize_uid(p.parent.name) == uid:
            return p
    return None


# ----------------------------
# UID extraction from JSON
# ----------------------------
def extract_uid_from_sequence(seq: dict) -> str:
    """
    Robustly extract UID from various possible key names.
    Your current JSON uses 'sequence_instance_uid', but we support others too.
    """
    candidates = [
        "sequence_instance_uid",
        "SequenceInstanceUID",
        "series_instance_uid",
        "SeriesInstanceUID",
        "uid",
        "UID",
    ]
    for k in candidates:
        v = normalize_uid(seq.get(k))
        if v:
            return v

    nested = seq.get("sequence")
    if isinstance(nested, dict):
        v = normalize_uid(nested.get("sequence_instance_uid") or nested.get("SeriesInstanceUID"))
        if v:
            return v

    return ""


# ----------------------------
# DOCX verification
# ----------------------------
def docx_contains_media(docx_path: Path) -> bool:
    """Check whether the DOCX zip has any embedded images under word/media/."""
    try:
        with zipfile.ZipFile(docx_path, "r") as z:
            names = z.namelist()
            return any(n.startswith("word/media/") for n in names)
    except Exception:
        return False


# ----------------------------
# Mosaic generation (on-the-fly)
# ----------------------------
def list_frame_files(seq_dir: Path, frames_subdir: str = FRAMES_SUBDIR) -> List[Path]:
    """
    Prefer frames in seq_dir/<frames_subdir>/*.{png,jpg,...}.
    Fall back to recursive search if that folder doesn't exist or is empty.
    """
    frames_dir = seq_dir / frames_subdir

    if frames_dir.exists() and frames_dir.is_dir():
        frames = [p for p in frames_dir.iterdir() if p.is_file() and p.suffix.lower() in IMAGE_EXTS]
        frames.sort(key=lambda p: p.name)
        if frames:
            return frames

    frames = [p for p in seq_dir.rglob("*") if p.is_file() and p.suffix.lower() in IMAGE_EXTS]
    frames.sort(key=lambda p: p.as_posix())
    return frames


def pick_frames(frames: List[Path], max_frames: int, stride: int) -> List[Path]:
    if stride < 1:
        stride = 1
    sampled = frames[::stride]

    if max_frames and max_frames > 0 and len(sampled) > max_frames:
        if max_frames == 1:
            return [sampled[len(sampled) // 2]]
        step = (len(sampled) - 1) / (max_frames - 1)
        idxs = [round(i * step) for i in range(max_frames)]
        sampled = [sampled[i] for i in idxs]
    return sampled


def _open_image_rgb(path: Path) -> Optional[Image.Image]:
    try:
        img = Image.open(path)
        img = ImageOps.exif_transpose(img)
        if img.mode != "RGB":
            img = img.convert("RGB")
        return img
    except Exception:
        return None


def _fit_to_box(img: Image.Image, box: Tuple[int, int]) -> Image.Image:
    target_w, target_h = box
    w, h = img.size
    if w <= 0 or h <= 0:
        return Image.new("RGB", (target_w, target_h), (0, 0, 0))

    scale = min(target_w / w, target_h / h)
    new_w = max(1, int(round(w * scale)))
    new_h = max(1, int(round(h * scale)))
    resized = img.resize((new_w, new_h), Image.Resampling.BILINEAR)

    canvas = Image.new("RGB", (target_w, target_h), (0, 0, 0))
    off_x = (target_w - new_w) // 2
    off_y = (target_h - new_h) // 2
    canvas.paste(resized, (off_x, off_y))
    return canvas


def create_mosaic_image(
    frame_paths: List[Path],
    out_path: Path,
    tile_size: Tuple[int, int] = TILE_SIZE,
    cols: Optional[int] = None,
    max_cols: int = MOSAIC_MAX_COLS,
    threads: int = FRAME_LOAD_THREADS,
) -> Optional[Path]:
    """Create a single mosaic PNG at out_path from the provided frame_paths."""
    if not frame_paths:
        return None

    threads = max(1, int(threads))
    tiles: List[Image.Image] = []

    if threads == 1:
        for p in frame_paths:
            img = _open_image_rgb(p)
            if img is not None:
                tiles.append(img)
    else:
        with ThreadPoolExecutor(max_workers=threads) as tp:
            futs = [tp.submit(_open_image_rgb, p) for p in frame_paths]
            for f in futs:
                img = f.result()
                if img is not None:
                    tiles.append(img)

    if not tiles:
        return None

    n = len(tiles)
    if cols is None:
        cols = min(max_cols, max(1, int(math.ceil(math.sqrt(n)))))
    cols = max(1, cols)
    rows = int(math.ceil(n / cols))

    tile_w, tile_h = tile_size
    mosaic_w = cols * tile_w
    mosaic_h = rows * tile_h
    mosaic = Image.new("RGB", (mosaic_w, mosaic_h), (0, 0, 0))

    for idx, img in enumerate(tiles):
        r = idx // cols
        c = idx % cols
        tile = _fit_to_box(img, (tile_w, tile_h))
        mosaic.paste(tile, (c * tile_w, r * tile_h))

    out_path.parent.mkdir(parents=True, exist_ok=True)
    mosaic.save(out_path, format="PNG", optimize=True)
    return out_path


def ensure_mosaic_for_uid(
    uid: str,
    uid_to_mosaic: Dict[str, Path],
    uid_to_seqdir: Dict[str, Path],
    gen_log: List[str],
    gen_err_log: List[str],
) -> Optional[Path]:
    """
    Return a valid mosaic path if possible:
      1) If we already know it exists, return it.
      2) If seq_dir is known, generate mosaic.png from frames on the fly.
      3) If still not found, return None.
    """
    uid = normalize_uid(uid)
    if not uid:
        return None

    # 1) Known mosaic
    mp = uid_to_mosaic.get(uid)
    if mp and mp.exists():
        return mp

    # 2) Generate if we know the sequence dir
    seq_dir = uid_to_seqdir.get(uid)
    if seq_dir:
        mosaic_path = seq_dir / MOSAIC_NAME
        if mosaic_path.exists():
            uid_to_mosaic[uid] = mosaic_path
            return mosaic_path

        try:
            frames = list_frame_files(seq_dir, frames_subdir=FRAMES_SUBDIR)
            selected = pick_frames(frames, max_frames=MAX_FRAMES, stride=STRIDE)
            created = create_mosaic_image(
                frame_paths=selected,
                out_path=mosaic_path,
                tile_size=TILE_SIZE,
                cols=None,
                max_cols=MOSAIC_MAX_COLS,
                threads=FRAME_LOAD_THREADS,
            )
            if created and created.exists():
                uid_to_mosaic[uid] = created
                gen_log.append(f"[generated mosaic] uid={uid} path={created}")
                return created
            gen_err_log.append(
                f"[mosaic generation FAILED] uid={uid} seq_dir={seq_dir} reason=returned None or file missing"
            )
            return None
        except Exception as e:
            gen_err_log.append(f"[mosaic generation ERROR] uid={uid} seq_dir={seq_dir} err={repr(e)}")
            return None

    return None


# ----------------------------
# QA block
# ----------------------------
def add_qa_block(doc: Document, qa_items: List[dict]):
    """Add the sequence-level QA list (question, answer, confidence, evidence, notes)."""
    if not qa_items:
        return

    doc.add_paragraph().add_run("Q/A:").bold = True

    for i, qa in enumerate(qa_items, start=1):
        q = safe_get(qa, "question", "")
        a = safe_get(qa, "answer", "")
        conf = safe_get(qa, "confidence", "")
        evidence = safe_get(qa, "evidence", []) or []
        notes = safe_get(qa, "notes", "")

        p = doc.add_paragraph(style="List Number")
        run = p.add_run(q if q else f"QA item {i}")
        run.bold = True

        ans_p = doc.add_paragraph(style="List Bullet")
        ans_p.add_run("Answer: ").bold = True
        ans_p.add_run(str(a))

        if conf != "" and conf is not None:
            c_p = doc.add_paragraph(style="List Bullet")
            c_p.add_run("Confidence: ").bold = True
            c_p.add_run(str(conf))

        if isinstance(evidence, list) and len(evidence) > 0:
            e_p = doc.add_paragraph(style="List Bullet")
            e_p.add_run("Evidence:").bold = True
            for ev in evidence:
                doc.add_paragraph(str(ev), style="List Bullet 2")

        if notes:
            n_p = doc.add_paragraph(style="List Bullet")
            n_p.add_run("Notes: ").bold = True
            n_p.add_run(str(notes))


# ----------------------------
# Sequence block (embed mosaic)
# ----------------------------
def add_sequence_block(
    doc: Document,
    seq_number: Optional[int],
    uid: str,
    verbatim_text: str,
    qa_items: List[dict],
    mosaic_path: Optional[Path],
    missing_log: List[str],
    error_log: List[str],
) -> bool:
    """
    Returns True if mosaic embedded successfully (or at least found & attempted),
    False if mosaic missing (not found).
    """
    title = f"Sequence {seq_number}" if seq_number is not None else "Sequence"
    add_heading(doc, title, level=2)

    add_kv_paragraph(doc, "SequenceInstanceUID", uid)

    doc.add_paragraph().add_run("Verbatim text:").bold = True
    doc.add_paragraph(verbatim_text or "")

    add_qa_block(doc, qa_items)

    found = False
    if mosaic_path and mosaic_path.exists():
        found = True
        doc.add_paragraph().add_run("Mosaic:").bold = True
        pic_par = doc.add_paragraph()
        run = pic_par.add_run()
        try:
            run.add_picture(str(mosaic_path), width=Inches(MAX_IMAGE_WIDTH_INCHES))
            pic_par.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception as e:
            error_log.append(f"[add_picture ERROR] uid={uid} file={mosaic_path} err={repr(e)}")
            warn = doc.add_paragraph()
            warn.add_run("ERROR embedding mosaic image (see docx_image_errors.txt).").bold = True
    else:
        missing_log.append(f"[missing mosaic] uid={uid}")
        warn = doc.add_paragraph()
        warn.add_run("Mosaic not found for this UID (even after on-the-fly generation).").italic = True

    doc.add_page_break()
    return found


# ----------------------------
# JSON -> DOCX
# ----------------------------
def json_to_docx(
    json_path: Path,
    out_docx_path: Path,
    uid_to_mosaic: Dict[str, Path],
    uid_to_seqdir: Dict[str, Path],
    missing_log: List[str],
    generated_log: List[str],
    gen_err_log: List[str],
    error_log: List[str],
) -> Tuple[int, int, int]:
    """
    Returns (found_count, generated_count, not_found_count) for mosaics in THIS JSON file.
    """
    data = json.loads(json_path.read_text(encoding="utf-8", errors="replace"))
    doc = Document()

    anon_acc = safe_get(data, "Anon Acc #", safe_get(data, "Anon Acc", ""))
    add_heading(doc, f"Report sequences: {json_path.stem}", level=1)
    if anon_acc:
        add_kv_paragraph(doc, "Anon Acc #", anon_acc)

    radrpt = safe_get(data, "radrpt", "")
    if radrpt:
        add_heading(doc, "Rad report text (radrpt)", level=2)
        doc.add_paragraph(radrpt)

    sequences: List[dict] = safe_get(data, "sequences", []) or []

    found_count = 0
    generated_count = 0
    not_found_count = 0

    if not sequences:
        doc.add_paragraph("No sequences found in this JSON.")
    else:
        add_heading(doc, "Sequences", level=1)
        for seq in sequences:
            uid = extract_uid_from_sequence(seq)
            verbatim = safe_get(seq, "verbatim_text", "") or ""
            seq_no = safe_get(seq, "sequence_number", None)
            qa_items = safe_get(seq, "qa", []) or []

            mosaic_path = None
            was_generated = False

            if uid:
                # First: known mosaic index
                mp = uid_to_mosaic.get(uid)
                if mp and mp.exists():
                    mosaic_path = mp
                else:
                    # Try on-the-fly generation using UID->seq_dir index
                    before = len(generated_log)
                    mosaic_path = ensure_mosaic_for_uid(
                        uid=uid,
                        uid_to_mosaic=uid_to_mosaic,
                        uid_to_seqdir=uid_to_seqdir,
                        gen_log=generated_log,
                        gen_err_log=gen_err_log,
                    )
                    was_generated = (len(generated_log) > before) and (mosaic_path is not None)

                    # Final fallback: slow search (in case the seq_dir index missed it)
                    if mosaic_path is None:
                        mosaic_path = fallback_find_mosaic(DICOM_SEQ_ROOT, uid)

            was_found = add_sequence_block(
                doc=doc,
                seq_number=seq_no,
                uid=uid,
                verbatim_text=verbatim,
                qa_items=qa_items,
                mosaic_path=mosaic_path,
                missing_log=missing_log,
                error_log=error_log,
            )

            if was_found:
                found_count += 1
                if was_generated:
                    generated_count += 1
            else:
                not_found_count += 1

    out_docx_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_docx_path))
    return found_count, generated_count, not_found_count


# ----------------------------
# Main
# ----------------------------
def main():
    if not REPORTS_JSON_ROOT.exists():
        raise SystemExit(f"JSON root not found: {REPORTS_JSON_ROOT}")
    if not DICOM_SEQ_ROOT.exists():
        raise SystemExit(f"DICOM processed root not found: {DICOM_SEQ_ROOT}")

    OUTPUT_DOCX_ROOT.mkdir(parents=True, exist_ok=True)

    # Index: existing mosaics + UID->seq_dir (by scanning frames/)
    uid_to_mosaic = build_uid_to_mosaic_index(DICOM_SEQ_ROOT)
    uid_to_seqdir = build_uid_to_seqdir_index_by_frames(DICOM_SEQ_ROOT, frames_subdir=FRAMES_SUBDIR)

    json_files = sorted(REPORTS_JSON_ROOT.rglob("*.json"))
    print(f"Found {len(json_files):,} JSON files")

    missing_log: List[str] = []
    generated_log: List[str] = []
    gen_err_log: List[str] = []
    error_log: List[str] = []
    no_media_log: List[str] = []

    total_found = 0
    total_generated = 0
    total_not_found = 0

    for jp in json_files:
        rel = jp.relative_to(REPORTS_JSON_ROOT)
        out_docx = OUTPUT_DOCX_ROOT / rel.with_suffix(".docx")

        try:
            found_count, generated_count, not_found_count = json_to_docx(
                json_path=jp,
                out_docx_path=out_docx,
                uid_to_mosaic=uid_to_mosaic,
                uid_to_seqdir=uid_to_seqdir,
                missing_log=missing_log,
                generated_log=generated_log,
                gen_err_log=gen_err_log,
                error_log=error_log,
            )

            total_found += found_count
            total_generated += generated_count
            total_not_found += not_found_count

            if not docx_contains_media(out_docx):
                no_media_log.append(f"[no media in docx] {out_docx}")

            print(f"Wrote: {out_docx}")
            print(f"  mosaics: found={found_count} generated={generated_count} not_found={not_found_count}")

        except Exception as e:
            error_log.append(f"[DOCX ERROR] json={jp} err={repr(e)}")
            print(f"[DOCX ERROR] {jp} -> {repr(e)}")

    # Logs
    if missing_log:
        miss_path = OUTPUT_DOCX_ROOT / "missing_mosaics.txt"
        miss_path.write_text("\n".join(missing_log) + "\n", encoding="utf-8")
        print(f"\nWARNING: Missing mosaics for {len(missing_log):,} sequence entries. Saved: {miss_path}")

    if generated_log:
        gen_path = OUTPUT_DOCX_ROOT / "generated_mosaics.txt"
        gen_path.write_text("\n".join(generated_log) + "\n", encoding="utf-8")
        print(f"\nINFO: Generated {len(generated_log):,} mosaics on-the-fly. Saved: {gen_path}")

    if gen_err_log:
        ge_path = OUTPUT_DOCX_ROOT / "mosaic_generation_errors.txt"
        ge_path.write_text("\n".join(gen_err_log) + "\n", encoding="utf-8")
        print(f"\nWARNING: Mosaic generation errors: {len(gen_err_log):,}. Saved: {ge_path}")

    if error_log:
        err_path = OUTPUT_DOCX_ROOT / "docx_image_errors.txt"
        err_path.write_text("\n".join(error_log) + "\n", encoding="utf-8")
        print(f"\nWARNING: Encountered {len(error_log):,} DOCX/image errors. Saved: {err_path}")

    if no_media_log:
        nm_path = OUTPUT_DOCX_ROOT / "docx_without_media.txt"
        nm_path.write_text("\n".join(no_media_log) + "\n", encoding="utf-8")
        print(f"\nWARNING: {len(no_media_log):,} DOCX files had no embedded media. Saved: {nm_path}")

    print("\n=== FINAL MOSAIC COUNTS ===")
    print(f"Total mosaics embedded:       {total_found}")
    print(f"Total mosaics generated:      {total_generated}")
    print(f"Total mosaics still missing:  {total_not_found}")
    print("\nDone.")


if __name__ == "__main__":
    try:
        main()
    except KeyboardInterrupt:
        print("\nInterrupted — partial outputs are preserved.", file=sys.stderr)
        raise
