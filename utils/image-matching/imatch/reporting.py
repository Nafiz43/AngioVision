"""Result outputs: console table, bar chart PNG, Markdown report, examples docx.

matplotlib and python-docx are optional — each output degrades to a logged
warning when its dependency is missing.
"""

import csv
import datetime
import io
import logging
from pathlib import Path

from .config import code
from .evaluation import fmt_cv, macro, micro

log = logging.getLogger(__name__)


def print_results_table(
    all_results: dict[int, dict[str, dict]],
    k_values: list[int], n_folds: int, model_id: str, frame_mode: str,
) -> None:
    r1 = all_results[k_values[0]]
    rows = sorted(
        [(lbl, code(lbl), r,
          sum(r.get("fold_accs", [])) / len(r["fold_accs"]) if r.get("fold_accs")
          else (r["correct"] / r["total_evaled"] if r["total_evaled"] > 0 else None))
         for lbl, r in r1.items()],
        key=lambda x: (x[3] is None, -(x[3] or 0)),
    )
    W   = max(len(lbl) for lbl, *_ in rows) + 2
    hdr = (f"  {'Category':<{W}}  {'Code':<8}  {'~Train':>7}  {'~Test':>6}  "
           f"{'Evaled':>7}  {'K=1 Acc (mean±std)':>20}")
    sep = "─" * len(hdr)
    print()
    print(f"  AngioVision  {n_folds}-fold CV  |  model={model_id}  |  mode={frame_mode}")
    print(sep); print(hdr); print(sep)
    for lbl, cd, r, _ in rows:
        print(f"  {lbl:<{W}}  {cd:<8}  {r['n_train']:>7,}  {r['n_test']:>6,}  "
              f"{r['total_evaled']:>7,}  {fmt_cv(r):>20}")
    print(sep)
    print()
    print(f"  Cross-validated micro & macro accuracy vs K  ({n_folds} folds):")
    print("  " + "  ".join(f"K={k:>2}" for k in k_values))

    def _fmt(val):
        return f"{val:>5.1%}"

    micro_r = "  " + "  ".join(
        _fmt(micro(all_results[k])[0] / micro(all_results[k])[1])
        if micro(all_results[k])[1] > 0 else "  N/A" for k in k_values)
    macro_r = "  " + "  ".join(_fmt(macro(all_results[k])) for k in k_values)
    print(f"  micro: {micro_r}"); print(f"  macro: {macro_r}"); print()


def save_bar_chart(
    all_results: dict[int, dict[str, dict]],
    k_values: list[int], n_folds: int, model_id: str, frame_mode: str, out_path: Path,
) -> None:
    try:
        import matplotlib.pyplot as plt
        import matplotlib.patches as mpatches
    except ImportError:
        log.warning("matplotlib unavailable — chart skipped"); return

    r1 = all_results[k_values[0]]
    rows = []
    for lbl, r in r1.items():
        if not r["total_evaled"]:
            continue
        accs = r.get("fold_accs", [])
        mean = sum(accs) / len(accs) if accs else r["correct"] / r["total_evaled"]
        std  = (sum((a - mean) ** 2 for a in accs) / len(accs)) ** 0.5 if len(accs) > 1 else 0.0
        rows.append((code(lbl), lbl, mean, std, r["n_train"], r["n_test"]))
    rows.sort(key=lambda x: -x[2])
    if not rows:
        log.warning("No evaluated categories — chart skipped"); return

    colours = ["#2ecc71" if a >= 0.80 else "#f39c12" if a >= 0.50 else "#e74c3c" for _, _, a, *_ in rows]
    stds    = [s for _, _, _, s, *_ in rows]

    fig = plt.figure(figsize=(18, max(6, len(rows) * 0.42 + 3)))
    ax_bar  = fig.add_axes([0.04, 0.10, 0.44, 0.82])
    ax_line = fig.add_axes([0.56, 0.10, 0.40, 0.82])

    y_pos = range(len(rows))
    ax_bar.barh(y_pos, [a for _, _, a, *_ in rows],
                xerr=stds, error_kw={"ecolor": "#555", "capsize": 3, "linewidth": 1.2},
                color=colours, edgecolor="black", linewidth=0.5, height=0.7)
    for i, (cd, lbl, acc, std, nt, nts) in enumerate(rows):
        ax_bar.text(min(acc + std + 0.015, 0.99), i, f"{acc:.1%}±{std:.1%}",
                    va="center", fontsize=7, fontweight="bold")
    ax_bar.set_yticks(list(y_pos)); ax_bar.set_yticklabels([r[0] for r in rows], fontsize=8.5)
    ax_bar.set_xlim(0, 1.22)
    ax_bar.axvline(x=0.5,  color="gray",  linewidth=1, linestyle="--", alpha=0.5)
    ax_bar.axvline(x=0.80, color="green", linewidth=1, linestyle="--", alpha=0.5)
    ax_bar.set_xlabel(f"K=1 Accuracy  (mean±std,  {n_folds} folds)", fontsize=9)
    ax_bar.set_title(f"Per-Category K=1 Accuracy  ({n_folds}-fold CV)", fontsize=10, fontweight="bold")
    ax_bar.invert_yaxis(); ax_bar.grid(axis="x", linestyle="--", alpha=0.4)
    ax_bar.legend(handles=[mpatches.Patch(color="#2ecc71", label="≥ 80%"),
                           mpatches.Patch(color="#f39c12", label="50–79%"),
                           mpatches.Patch(color="#e74c3c", label="< 50%")],
                  fontsize=8, loc="lower right")

    micro_vals = [micro(all_results[k])[0] / micro(all_results[k])[1]
                  if micro(all_results[k])[1] > 0 else 0.0 for k in k_values]
    macro_vals = [macro(all_results[k]) for k in k_values]
    ax_line.plot(k_values, micro_vals, "o-",  color="#2980b9", linewidth=2, markersize=7, label="Micro (pooled)")
    ax_line.plot(k_values, macro_vals, "s--", color="#e67e22", linewidth=2, markersize=7, label="Macro (pooled)")
    for k, mv, av in zip(k_values, micro_vals, macro_vals):
        ax_line.annotate(f"{mv:.1%}", (k, mv), textcoords="offset points", xytext=(0, 8),   ha="center", fontsize=7.5, color="#2980b9")
        ax_line.annotate(f"{av:.1%}", (k, av), textcoords="offset points", xytext=(0, -14), ha="center", fontsize=7.5, color="#e67e22")
    ax_line.set_xticks(k_values); ax_line.set_xlabel("K", fontsize=9)
    ax_line.set_ylabel("Accuracy (pooled across folds)", fontsize=9)
    ax_line.set_title(f"Micro & Macro vs K  ({n_folds}-fold CV)", fontsize=10, fontweight="bold")
    ax_line.set_ylim(0, 1.05); ax_line.axhline(y=0.80, color="green", linewidth=1, linestyle="--", alpha=0.4)
    ax_line.grid(linestyle="--", alpha=0.4); ax_line.legend(fontsize=9)
    plt.suptitle(f"AngioVision K@N  |  {n_folds}-fold CV  |  model: {model_id}  |  frames: {frame_mode}",
                 fontsize=10, fontweight="bold", y=1.01)
    plt.savefig(str(out_path), dpi=150, bbox_inches="tight"); plt.close()
    log.info(f"Chart saved → {out_path}")


def write_markdown(
    all_results: dict[int, dict[str, dict]],
    k_values: list[int], n_folds: int, seed: int,
    model_id: str, frame_mode: str, emb_dim: int, temporal: bool, out_path: Path,
) -> None:
    r1   = all_results[k_values[0]]
    cats = sorted(r1.keys(), key=lambda l: -(r1[l]["correct"] / r1[l]["total_evaled"]
                                             if r1[l]["total_evaled"] > 0 else 0))
    now      = datetime.datetime.now().strftime("%Y-%m-%d %H:%M")
    khdrs    = " | ".join(f"K={k}" for k in k_values)
    emb_note = f"{emb_dim}-dim" + (f" → 2×{emb_dim}-dim temporal" if temporal else "")

    lines: list[str] = [
        "# AngioVision K@N Retrieval Evaluation  (Cross-Validated)\n\n",
        f"**Generated:** {now}  \n",
        f"**Model:** `{model_id}` ({emb_note}, L2-normalised)  \n",
        f"**Frame mode:** `{frame_mode}`  \n",
        f"**Temporal aggregation:** {'✓ mean+std pooling (one vector per sequence)' if temporal else '✗ per-frame (majority vote)'}  \n",
        f"**Validation:** {n_folds}-fold cross-validation  |  seed={seed}  \n",
        f"**Matching:** {'sequence-level cosine similarity' if temporal else 'majority vote across K × N_frames results'}  \n\n",
        f"> Cells show **mean±std** across {n_folds} folds.  Micro/Macro rows show pooled counts.\n\n",
        "## Per-Category K@N Accuracy\n\n",
        f"| Category | Code | ~Train/fold | ~Test/fold | Total Evaled | {khdrs} |\n",
        "|:---------|:----:|------------:|-----------:|-------------:" + ("".join("-----------:|" for _ in k_values)) + "\n",
    ]
    for label in cats:
        r = r1[label]
        cells = [fmt_cv(all_results[k][label]) for k in k_values]
        lines.append(f"| {label} | {code(label)} | {r['n_train']:,} | {r['n_test']:,} | "
                     f"{r['total_evaled']:,} | " + " | ".join(cells) + " |\n")

    tc0, te0 = micro(r1); base_micro = tc0 / te0 if te0 > 0 else 0.0
    total_train  = sum(r1[l]["n_train"]      for l in cats)
    total_test   = sum(r1[l]["n_test"]       for l in cats)
    total_evaled = sum(r1[l]["total_evaled"] for l in cats)
    micro_cells = [f"**{micro(all_results[k])[0]/micro(all_results[k])[1]:.1%}**"
                   if micro(all_results[k])[1] > 0 else "N/A" for k in k_values]
    macro_cells = [f"**{macro(all_results[k]):.1%}**" for k in k_values]

    lines += [
        f"| **MICRO (pooled)** | | **{total_train:,}** | **{total_test:,}** | **{total_evaled:,}** | "
        + " | ".join(micro_cells) + " |\n",
        "| **MACRO (pooled)** | | | | | " + " | ".join(macro_cells) + " |\n\n",
        "## K-Sweep Summary\n\n",
        "| K | Micro (pooled) | Macro (pooled) | Δ vs K=1 |\n",
        "|--:|--------------:|--------------:|---------:|\n",
    ]
    for k in k_values:
        tc, te = micro(all_results[k]); m = tc / te if te > 0 else 0.0; delta = m - base_micro
        lines.append(f"| {k} | {m:.1%} | {macro(all_results[k]):.1%} | "
                     f"{'+' if delta >= 0 else ''}{delta:.1%} |\n")

    lines += ["\n## Best K Per Category\n\n",
              "| Category | Code | ~Train/fold | ~Test/fold | Best K | Best Acc (mean) |\n",
              "|:---------|:----:|------------:|-----------:|-------:|----------------:|\n"]
    for label in cats:
        best_k, best_mean = None, -1.0
        for k in k_values:
            rk = all_results[k][label]; accs = rk.get("fold_accs", [])
            m  = (sum(accs) / len(accs) if accs
                  else (rk["correct"] / rk["total_evaled"] if rk["total_evaled"] > 0 else 0.0))
            if m > best_mean:
                best_mean = m; best_k = k
        if best_k is not None:
            lines.append(f"| {label} | {code(label)} | {r1[label]['n_train']:,} | "
                         f"{r1[label]['n_test']:,} | K={best_k} | {best_mean:.1%} |\n")

    out_path.write_text("".join(lines), encoding="utf-8")
    log.info(f"Markdown saved → {out_path}")


def write_csv(
    all_results: dict[int, dict[str, dict]],
    k_values: list[int], n_folds: int, seed: int,
    model_id: str, frame_mode: str, temporal: bool, out_path: Path,
) -> None:
    """Machine-readable results: one row per category × K, plus MICRO/MACRO rows per K.

    Long format so it loads straight into pandas/R without header parsing.
    mean_acc/std_acc are across folds; pooled_acc = correct/total_evaled pooled
    over all folds (identical to mean_acc when split-mode=holdout).
    """
    fields = ["model", "frame_mode", "temporal", "split_folds", "seed",
              "category", "code", "k", "n_train_per_fold", "n_test_per_fold",
              "total_evaled", "skipped", "correct", "pooled_acc",
              "mean_acc", "std_acc", "n_folds_scored"]
    base = {"model": model_id, "frame_mode": frame_mode,
            "temporal": int(temporal), "split_folds": n_folds, "seed": seed}

    with open(out_path, "w", newline="", encoding="utf-8") as fh:
        w = csv.DictWriter(fh, fieldnames=fields)
        w.writeheader()
        for k in k_values:
            rk = all_results[k]
            for label in sorted(rk.keys()):
                r = rk[label]
                accs = r.get("fold_accs", [])
                mean = sum(accs) / len(accs) if accs else None
                std  = ((sum((a - mean) ** 2 for a in accs) / len(accs)) ** 0.5
                        if accs and mean is not None and len(accs) > 1 else (0.0 if accs else None))
                pooled = r["correct"] / r["total_evaled"] if r["total_evaled"] > 0 else None
                w.writerow({**base, "category": label, "code": code(label), "k": k,
                            "n_train_per_fold": r["n_train"], "n_test_per_fold": r["n_test"],
                            "total_evaled": r["total_evaled"], "skipped": r["skipped"],
                            "correct": r["correct"],
                            "pooled_acc": f"{pooled:.6f}" if pooled is not None else "",
                            "mean_acc": f"{mean:.6f}" if mean is not None else "",
                            "std_acc": f"{std:.6f}" if std is not None else "",
                            "n_folds_scored": len(accs)})
            tc, te = micro(rk)
            w.writerow({**base, "category": "MICRO", "code": "", "k": k,
                        "n_train_per_fold": sum(r["n_train"] for r in rk.values()),
                        "n_test_per_fold": sum(r["n_test"] for r in rk.values()),
                        "total_evaled": te, "skipped": sum(r["skipped"] for r in rk.values()),
                        "correct": tc,
                        "pooled_acc": f"{tc/te:.6f}" if te > 0 else "",
                        "mean_acc": "", "std_acc": "", "n_folds_scored": ""})
            w.writerow({**base, "category": "MACRO", "code": "", "k": k,
                        "n_train_per_fold": "", "n_test_per_fold": "",
                        "total_evaled": "", "skipped": "", "correct": "",
                        "pooled_acc": f"{macro(rk):.6f}",
                        "mean_acc": "", "std_acc": "", "n_folds_scored": ""})
    log.info(f"CSV saved → {out_path}")


def save_retrieval_docx(
    examples: dict[str, list[dict]], k1_results: dict[str, dict],
    out_path: Path, model_id: str, frame_mode: str, temporal: bool,
) -> None:
    """
    Word document with K=1 retrieval examples per category.
    Requires: pip install python-docx
    """
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        log.error("python-docx not installed — run: pip install python-docx"); return

    from .frames import get_display_frame_png

    def _shade(cell, hex_fill):
        tc = cell._tc; pr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd"); shd.set(qn("w:fill"), hex_fill); shd.set(qn("w:val"), "clear")
        pr.append(shd)

    def _cell_para(cell, text, bold=False, size_pt=10, color=None):
        p = cell.paragraphs[0]; p.clear(); run = p.add_run(text)
        run.bold = bold; run.font.size = Pt(size_pt)
        if color:
            run.font.color.rgb = color

    def _add_img(cell, png_bytes):
        p = cell.paragraphs[0]; p.clear(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if png_bytes:
            p.add_run().add_picture(io.BytesIO(png_bytes), width=Inches(2.3))
        else:
            p.add_run("[unavailable]").font.size = Pt(8)

    COL_W = [576, 3888, 3888, 1728]
    HIT_GREEN = "D5F5E3"; MISS_RED = "FADBD8"; HEADER_BG = "2C3E50"

    doc = Document()
    for sec in doc.sections:
        sec.left_margin = sec.right_margin = sec.top_margin = sec.bottom_margin = Inches(0.75)

    title = doc.add_heading("AngioVision  K@1  Retrieval Examples", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta = doc.add_paragraph(); meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"Model: {model_id}   |   Frame mode: {frame_mode}   |   "
                 f"Temporal: {'ON' if temporal else 'OFF'}   |   "
                 f"Generated: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M')}").font.size = Pt(9)
    doc.add_paragraph()

    for label in sorted(examples.keys()):
        cases = examples[label]; r1 = k1_results.get(label, {})
        n_ev = r1.get("total_evaled", 0); n_cor = r1.get("correct", 0)
        acc  = n_cor / n_ev if n_ev > 0 else 0.0
        h = doc.add_heading(level=2); h.clear()
        h.add_run(f"{label}  ({code(label)})  —  K@1: {n_cor}/{n_ev} ({acc:.0%})  |  "
                  f"Train≈{r1.get('n_train', '?')}  Test≈{r1.get('n_test', '?')}")

        table = doc.add_table(rows=1, cols=4); table.style = "Table Grid"
        hdr = table.rows[0].cells
        for i, (cell, txt) in enumerate(zip(hdr, ["#", "Query Frame", "Retrieved Frame", "Outcome"])):
            _cell_para(cell, txt, bold=True, size_pt=10, color=RGBColor(0xFF, 0xFF, 0xFF))
            _shade(cell, HEADER_BG)
            tc = cell._tc; pr = tc.get_or_add_tcPr(); w = OxmlElement("w:tcW")
            w.set(qn("w:w"), str(COL_W[i])); w.set(qn("w:type"), "dxa"); pr.append(w)

        for idx, case in enumerate(cases):
            row = table.add_row().cells; hit = case["hit"]; bg = HIT_GREEN if hit else MISS_RED
            _cell_para(row[0], str(idx + 1), bold=True, size_pt=9); _shade(row[0], bg)
            q_png = get_display_frame_png(case["test_seq"])
            _add_img(row[1], q_png)
            row[1].add_paragraph().add_run(f"True: {code(label)}").font.size = Pt(8)
            r_png = get_display_frame_png(case["train_seq"]) if case.get("train_seq") else None
            _add_img(row[2], r_png)
            row[2].add_paragraph().add_run(f"Retrieved: {code(case['retr_label'])}").font.size = Pt(8)
            _cell_para(row[3], f"{'✓  HIT' if hit else '✗  MISS'}\n\nQuery:     {code(label)}\n"
                               f"Retrieved: {code(case['retr_label'])}", bold=hit, size_pt=9)
            _shade(row[3], bg)
            for i, cell in enumerate(row):
                tc = cell._tc; pr = tc.get_or_add_tcPr(); w = OxmlElement("w:tcW")
                w.set(qn("w:w"), str(COL_W[i])); w.set(qn("w:type"), "dxa"); pr.append(w)
        doc.add_paragraph()

    doc.save(str(out_path)); log.info(f"Retrieval examples docx saved → {out_path}")
